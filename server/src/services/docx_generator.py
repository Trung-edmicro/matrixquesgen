import json
from pathlib import Path
from typing import Dict, List, Union, Optional, Any
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class DocxGenerator:
    """Class tạo file DOCX từ dữ liệu JSON"""
    
    def __init__(self, verbose: bool = False):
        self.document = None
        self.output_path = None
        self.verbose = verbose
        
        # Màu nền theo level
        self.level_colors = {
            'NB': 'ADFF2F',  # Bright green
            'TH': '40E0D0',  # Turquoise
            'VD': 'FFFF00'   # Yellow
        }
    
    def create_new_document(self):
        """Tạo document mới"""
        self.document = Document()
        if self.verbose:
            print("✓ Đã tạo document mới")
    
    def set_document_style(self, 
                          font_name: str = "Times New Roman",
                          font_size: int = 12):
        """
        Thiết lập style mặc định cho document
        
        Args:
            font_name (str): Tên font chữ
            font_size (int): Kích thước font
        """
        if self.document is None:
            self.create_new_document()
        
        # Thiết lập style cho Normal
        style = self.document.styles['Normal']
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)
        
        if self.verbose:
            print(f"✓ Đã thiết lập style: {font_name}, {font_size}pt")
    
    def add_heading(self, 
                   text: str, 
                   level: int = 1,
                   alignment: str = "left"):
        """
        Thêm tiêu đề
        
        Args:
            text (str): Nội dung tiêu đề
            level (int): Cấp độ tiêu đề (1-9)
            alignment (str): Căn lề (left, center, right)
        """
        if self.document is None:
            self.create_new_document()
        
        heading = self.document.add_heading(text, level=level)
        
        if alignment == "center":
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "right":
            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        return heading
    
    def add_paragraph(self, 
                     text: str,
                     bold: bool = False,
                     italic: bool = False,
                     alignment: str = "left"):
        """
        Thêm đoạn văn
        
        Args:
            text (str): Nội dung
            bold (bool): In đậm
            italic (bool): In nghiêng
            alignment (str): Căn lề
        """
        if self.document is None:
            self.create_new_document()
        
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        
        if alignment == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == "justify":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        return paragraph
    
    def add_table_from_dict(self, 
                           data: List[Dict],
                           headers: Optional[List[str]] = None,
                           auto_headers: bool = True):
        """
        Thêm bảng từ dữ liệu dictionary
        
        Args:
            data (List[Dict]): Dữ liệu dạng list of dict
            headers (List[str], optional): Tiêu đề cột
            auto_headers (bool): Tự động lấy headers từ dict keys
        """
        if self.document is None:
            self.create_new_document()
        
        if not data:
            if self.verbose:
                print("⚠ Không có dữ liệu để tạo bảng")
            return
        
        # Lấy headers
        if headers is None and auto_headers:
            headers = list(data[0].keys())
        
        # Tạo bảng
        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # Thêm header
        header_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            header_cells[idx].text = str(header)
            # Format header
            for paragraph in header_cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Thêm dữ liệu
        for item in data:
            row_cells = table.add_row().cells
            for idx, header in enumerate(headers):
                value = item.get(header, "")
                row_cells[idx].text = str(value)
        
        if self.verbose:
            print(f"✓ Đã thêm bảng {len(data)} hàng x {len(headers)} cột")
        return table
    
    def add_page_break(self):
        """Thêm ngắt trang"""
        if self.document is None:
            self.create_new_document()
        
        self.document.add_page_break()
    
    def _set_run_background(self, run, level: str):
        """Tô màu nền cho run dựa vào level"""
        color = self.level_colors.get(level, 'FFFFFF')
        
        # Tạo shading element
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        run._element.get_or_add_rPr().append(shading_elm)
    
    def generate_from_json(self, 
                          json_data: Union[Dict, str],
                          template: Optional[Dict] = None):
        """
        Tạo document từ dữ liệu JSON
        
        Args:
            json_data (Union[Dict, str]): Dữ liệu JSON hoặc đường dẫn file JSON
            template (Dict, optional): Template định nghĩa cấu trúc document
        """
        if self.document is None:
            self.create_new_document()
        
        # Đọc JSON nếu là đường dẫn
        if isinstance(json_data, str):
            with open(json_data, 'r', encoding='utf-8') as f:
                json_data = json.load(f)
        
        # Nếu có template, sử dụng template
        if template:
            self._apply_template(json_data, template)
        else:
            # Mặc định: tạo document đơn giản từ JSON
            self._generate_default(json_data)
    
    def _apply_template(self, data: Dict, template: Dict):
        """Áp dụng template để tạo document"""
        for section in template.get('sections', []):
            section_type = section.get('type')
            
            if section_type == 'heading':
                self.add_heading(
                    text=section.get('text', '').format(**data),
                    level=section.get('level', 1),
                    alignment=section.get('alignment', 'left')
                )
            
            elif section_type == 'paragraph':
                self.add_paragraph(
                    text=section.get('text', '').format(**data),
                    bold=section.get('bold', False),
                    italic=section.get('italic', False),
                    alignment=section.get('alignment', 'left')
                )
            
            elif section_type == 'table':
                data_key = section.get('data_key')
                if data_key and data_key in data:
                    self.add_table_from_dict(
                        data=data[data_key],
                        headers=section.get('headers')
                    )
            
            elif section_type == 'page_break':
                self.add_page_break()
    
    def _generate_default(self, data: Dict):
        """Tạo document mặc định từ JSON"""
        for key, value in data.items():
            # Thêm key làm heading
            self.add_heading(str(key).replace('_', ' ').title(), level=2)
            
            # Thêm value
            if isinstance(value, list) and value and isinstance(value[0], dict):
                # Nếu là list of dict, tạo bảng
                self.add_table_from_dict(value)
            elif isinstance(value, dict):
                # Nếu là dict, hiển thị từng cặp key-value
                for k, v in value.items():
                    self.add_paragraph(f"{k}: {v}")
            else:
                # Giá trị đơn giản
                self.add_paragraph(str(value))
            
            self.add_paragraph("")  # Thêm dòng trống
    
    def save(self, output_path: str):
        """
        Lưu document
        
        Args:
            output_path (str): Đường dẫn file output
        """
        if self.document is None:
            print("✗ Chưa có document để lưu")
            return
        
        try:
            # Tạo thư mục nếu chưa tồn tại
            output_dir = Path(output_path).parent
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # Lưu file
            self.document.save(output_path)
            self.output_path = output_path
            
            if self.verbose:
                print(f"✓ Đã lưu file: {output_path}")
        
        except Exception as e:
            print(f"✗ Lỗi khi lưu file: {str(e)}")
            raise
    
    def generate_questions_document(self, 
                                   json_data: Union[Dict, str],
                                   output_path: str):
        """
        Tạo document câu hỏi từ dữ liệu JSON
        
        Args:
            json_data (Union[Dict, str]): Dữ liệu JSON hoặc đường dẫn file JSON
            output_path (str): Đường dẫn file DOCX output
        """
        # Đọc JSON nếu là đường dẫn
        if isinstance(json_data, str):
            with open(json_data, 'r', encoding='utf-8') as f:
                json_data = json.load(f)
        
        # Tạo document mới
        self.create_new_document()
        self.set_document_style()
        
        # Thêm tiêu đề
        metadata = json_data.get('metadata', {})
        self.add_heading('ĐỀ THI TRẮC NGHIỆM', level=1, alignment='center')
        self.add_paragraph(f"Ma trận: {metadata.get('matrix_file', 'N/A')}", alignment='center')
        self.add_paragraph(f"Ngày tạo: {metadata.get('generated_at', 'N/A')[:10]}", alignment='center')
        self.add_paragraph(f"Tổng số câu: {metadata.get('total_questions', 0)}", alignment='center')
        self.add_paragraph("")
        
        questions_data = json_data.get('questions', {})
        
        # Xuất câu hỏi TN
        tn_questions = questions_data.get('TN', [])
        if tn_questions:
            tn_questions_sorted = sorted(tn_questions, key=lambda x: int(x.get('question_code', 'C0')[1:]))
            
            self.add_heading('PHẦN I. Thí sinh trả lời từ câu 1 đến câu 24. Mỗi câu hỏi thí sinh chỉ chọn một phương án.', level=2)
            
            for idx, q in enumerate(tn_questions_sorted, 1):
                self._add_tn_question(q, idx)
            
            self.add_paragraph("")
        
        # Xuất câu hỏi DS
        ds_questions = questions_data.get('DS', [])
        if ds_questions:
            ds_questions_sorted = sorted(ds_questions, key=lambda x: int(x.get('question_code', 'C0')[1:]))
            
            self.add_heading('PHẦN II. Thí sinh trả lời từ câu 1 đến 4. Trong mỗi ý a), b), c), d) ở mỗi câu, thí sinh chọn đúng hoặc sai.', level=2)
            
            for idx, q in enumerate(ds_questions_sorted, 1):
                self._add_ds_question(q, idx)
            
            self.add_paragraph("")
        
        # Thêm phần đáp án
        self.add_page_break()
        self.add_heading('ĐÁP ÁN VÀ GIẢI THÍCH', level=1, alignment='center')
        self.add_paragraph("")
        
        # Đáp án TN
        if tn_questions:
            self.add_heading('PHẦN I: TRẮC NGHIỆM', level=2)
            
            for idx, q in enumerate(tn_questions_sorted, 1):
                self._add_tn_answer(q, idx)
        
        # Đáp án DS
        if ds_questions:
            self.add_heading('PHẦN II: ĐÚNG/SAI', level=2)
            
            for idx, q in enumerate(ds_questions_sorted, 1):
                self._add_ds_answer(q, idx)
        
        # Lưu file
        self.save(output_path)
    
    def _add_tn_question(self, question: Dict, number: int):
        """Thêm câu hỏi trắc nghiệm"""
        # Câu hỏi
        para = self.document.add_paragraph()
        
        level = question.get('level', 'NB')
        
        run = para.add_run(f"Câu {number} ")
        run.bold = True
        run = para.add_run(f"({level})")
        run.bold = True
        self._set_run_background(run, level)
        
        para.add_run(". ")
        para.add_run(question.get('question_stem', ''))
        
        # Các lựa chọn
        options = question.get('options', {})
        correct_answer = question.get('correct_answer', '')
        
        for key in ['A', 'B', 'C', 'D']:
            if key in options:
                para = self.document.add_paragraph()
                
                # Tô đỏ nếu là đáp án đúng
                if key == correct_answer:
                    run = para.add_run(f"{key}. {options[key]}")
                    run.font.color.rgb = RGBColor(255, 0, 0)
                else:
                    para.add_run(f"{key}. {options[key]}")
    
    def _add_ds_question(self, question: Dict, number: int):
        """Thêm câu hỏi đúng/sai"""
        # Tiêu đề câu
        para = self.document.add_paragraph()
        para.add_run(f"Câu {number}. ").bold = True
        para.add_run("Cho đoạn tư liệu sau:")
        
        # Tư liệu
        self.add_paragraph(question.get('source_text', ''), italic=True)
        
        statements = question.get('statements', {})
        for label in ['a', 'b', 'c', 'd']:
            if label in statements:
                stmt = statements[label]
                level = stmt.get('level', 'NB')
                text = stmt.get('text', '')
                is_correct = stmt.get('correct_answer', False)
                
                para = self.document.add_paragraph()
                
                # "(level)" - tô màu nền
                run_level = para.add_run(f"({level})")
                self._set_run_background(run_level, level)
                
                # " label. text" - tô đỏ nếu đúng
                run_content = para.add_run(f" {label}. {text}")
                if is_correct:
                    run_content.font.color.rgb = RGBColor(255, 0, 0)
    
    def _add_tn_answer(self, question: Dict, number: int):
        """Thêm đáp án trắc nghiệm"""
        para = self.document.add_paragraph()
        para.add_run(f"Câu {number}. ").bold = True
        
        # Đáp án đúng tô đỏ
        run = para.add_run(f"Đáp án: {question.get('correct_answer', 'N/A')}")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)
        
        # Giải thích
        explanation = question.get('explanation', '')
        if explanation:
            self.add_paragraph(f"Giải thích: {explanation}")
        
        self.add_paragraph("")
    
    def _add_ds_answer(self, question: Dict, number: int):
        """Thêm đáp án đúng/sai"""
        para = self.document.add_paragraph()
        para.add_run(f"Câu {number}. ").bold = True
        
        # Đáp án từng mệnh đề
        statements = question.get('statements', {})
        answer_parts = []
        for label in ['a', 'b', 'c', 'd']:
            if label in statements:
                stmt = statements[label]
                is_correct = stmt.get('correct_answer', False)
                answer_text = "Đ" if is_correct else "S"
                answer_parts.append(f"{label}. {answer_text}")
        
        # Tô đỏ đáp án
        run = para.add_run(", ".join(answer_parts))
        run.font.color.rgb = RGBColor(255, 0, 0)
        
        # Giải thích
        explanations = question.get('explanation', {})
        if explanations:
            self.add_paragraph("Giải thích:")
            for label in ['a', 'b', 'c', 'd']:
                if label in explanations:
                    self.add_paragraph(f"{label}. {explanations[label]}")
        
        self.add_paragraph("")
    
    def create_from_json_file(self, 
                             json_path: str,
                             output_path: str,
                             template: Optional[Dict] = None):
        """
        Tạo DOCX từ file JSON
        
        Args:
            json_path (str): Đường dẫn file JSON
            output_path (str): Đường dẫn file DOCX output
            template (Dict, optional): Template
        """
        try:
            # Đọc JSON
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Tạo document
            self.create_new_document()
            self.set_document_style()
            self.generate_from_json(data, template)
            
            # Lưu file
            self.save(output_path)
            
            if self.verbose:
                print(f"✓ Đã tạo DOCX từ JSON: {json_path} -> {output_path}")
        
        except Exception as e:
            print(f"✗ Lỗi khi tạo DOCX từ JSON: {str(e)}")
            raise
