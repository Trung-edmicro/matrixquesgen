import json
from pathlib import Path
from typing import Dict, List, Union, Optional, Any
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import tempfile
import os
import re
import base64
from io import BytesIO
from PIL import Image
import threading
try:
    from playwright.sync_api import sync_playwright
    PLAYWRIGHT_AVAILABLE = True
    print("✓ Playwright available for chart rendering")
except ImportError:
    PLAYWRIGHT_AVAILABLE = False
    print("⚠️  Playwright not available. Charts will be rendered as text placeholders.")

import sys
sys.path.insert(0, str(Path(__file__).parent.parent.parent))
from config.settings import Config
from services.generators.helpers.chart_generation_helper import apply_layout


class DocxGenerator:
    """Class tạo file DOCX từ dữ liệu JSON"""
    
    def __init__(self, verbose: bool = False):
        self.document = None
        self.output_path = None
        self.verbose = verbose
        
        # Màu nền theo level
        self.level_colors = {
            'NB': 'ADFF2F',  # Bright green
            'TH': '40E0D0',  # Turquoise
            'VD': 'FFFF00'   # Yellow
        }
    
    def create_new_document(self):
        """Tạo document mới"""
        self.document = Document()
        if self.verbose:
            print("✓ Đã tạo document mới")
    
    def set_document_style(self, 
                          font_name: str = "Times New Roman",
                          font_size: int = 12):
        """
        Thiết lập style mặc định cho document
        
        Args:
            font_name (str): Tên font chữ
            font_size (int): Kích thước font
        """
        if self.document is None:
            self.create_new_document()
        
        # Thiết lập style cho Normal
        style = self.document.styles['Normal']
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)
        
        if self.verbose:
            print(f"✓ Đã thiết lập style: {font_name}, {font_size}pt")
    
    def add_heading(self, 
                   text: str, 
                   level: int = 1,
                   alignment: str = "left"):
        """
        Thêm tiêu đề
        
        Args:
            text (str): Nội dung tiêu đề
            level (int): Cấp độ tiêu đề (1-9)
            alignment (str): Căn lề (left, center, right)
        """
        if self.document is None:
            self.create_new_document()
        
        heading = self.document.add_heading(text, level=level)
        
        if alignment == "center":
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "right":
            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        return heading
    
    def add_paragraph(self, 
                     text: str,
                     bold: bool = False,
                     italic: bool = False,
                     alignment: str = "left"):
        """
        Thêm đoạn văn
        
        Args:
            text (str): Nội dung
            bold (bool): In đậm
            italic (bool): In nghiêng
            alignment (str): Căn lề
        """
        if self.document is None:
            self.create_new_document()
        
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        
        if alignment == "center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == "justify":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        return paragraph
    
    def add_table_from_dict(self, 
                           data: List[Dict],
                           headers: Optional[List[str]] = None,
                           auto_headers: bool = True):
        """
        Thêm bảng từ dữ liệu dictionary
        
        Args:
            data (List[Dict]): Dữ liệu dạng list of dict
            headers (List[str], optional): Tiêu đề cột
            auto_headers (bool): Tự động lấy headers từ dict keys
        """
        if self.document is None:
            self.create_new_document()
        
        if not data:
            if self.verbose:
                print("⚠ Không có dữ liệu để tạo bảng")
            return
        
        # Lấy headers
        if headers is None and auto_headers:
            headers = list(data[0].keys())
        
        # Tạo bảng với border đơn giản, không tô màu
        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'  # Border đơn giản, không màu
        
        # Thêm header
        header_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            cell_para = header_cells[idx].paragraphs[0]
            cell_para.clear()
            self._add_text_with_math(cell_para, str(header), bold=True)
            for run in cell_para.runs:
                run.bold = True
                run.font.size = Pt(11)
        
        # Thêm dữ liệu
        for item in data:
            row_cells = table.add_row().cells
            for idx, header in enumerate(headers):
                value = item.get(header, "")
                cell_para = row_cells[idx].paragraphs[0]
                cell_para.clear()
                self._add_text_with_math(cell_para, str(value))
        
        if self.verbose:
            print(f"✓ Đã thêm bảng {len(data)} hàng x {len(headers)} cột")
        return table
    
    def add_page_break(self):
        """Thêm ngắt trang"""
        if self.document is None:
            self.create_new_document()
        
        self.document.add_page_break()
    
    def _set_run_background(self, run, level: str):
        """Tô màu nền cho run dựa vào level"""
        color = self.level_colors.get(level, 'FFFFFF')
        
        # Tạo shading element
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        run._element.get_or_add_rPr().append(shading_elm)

    # ─────────────────────── Math / LaTeX helpers ───────────────────────

    def _find_mml2omml_xsl(self) -> Optional[str]:
        """Locate MML2OMML.XSL - checks bundled package first, then Microsoft Office."""
        candidates = []

        # 1. PyInstaller bundle (_MEIPASS) - highest priority when running as .exe
        if hasattr(sys, '_MEIPASS'):
            candidates.append(os.path.join(sys._MEIPASS, 'mml2omml', 'MML2OMML.XSL'))

        # 2. assets/ folder in project root (dev mode: copied from Office by build.ps1)
        _here = os.path.dirname(os.path.abspath(__file__))
        # Walk up from server/src/services/exporters/ to project root
        _root = _here
        for _ in range(5):
            _candidate = os.path.join(_root, 'assets', 'MML2OMML.XSL')
            if os.path.exists(_candidate):
                candidates.append(_candidate)
                break
            _root = os.path.dirname(_root)

        # 3. Microsoft Office fallback (for dev machines with Office installed)
        candidates += [
            r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL",
            r"C:\Program Files\Microsoft Office\root\Office15\MML2OMML.XSL",
            r"C:\Program Files (x86)\Microsoft Office\root\Office16\MML2OMML.XSL",
            r"C:\Program Files\Microsoft Office\Office16\MML2OMML.XSL",
            r"C:\Program Files (x86)\Microsoft Office\Office16\MML2OMML.XSL",
        ]

        for path in candidates:
            if os.path.exists(path):
                if self.verbose:
                    print(f"✓ MML2OMML.XSL found: {path}")
                return path
        return None

    def _init_math_engine(self):
        """Lazy-initialize the LaTeX → OMML conversion engine."""
        if hasattr(self, '_math_ready'):
            return
        self._math_ready = False
        self._math_xslt = None
        self._latex2mml_fn = None
        try:
            from latex2mathml.converter import convert as _latex2mml
            import lxml.etree as _etree
            xsl_path = self._find_mml2omml_xsl()
            if xsl_path:
                xsl_tree = _etree.parse(xsl_path)
                self._math_xslt = _etree.XSLT(xsl_tree)
                self._lxml_etree = _etree
                self._latex2mml_fn = _latex2mml
                self._math_ready = True
                if self.verbose:
                    print(f"✓ Math engine: latex2mathml + {xsl_path}")
            else:
                if self.verbose:
                    print("⚠️  MML2OMML.XSL not found — math will render as plain text")
        except ImportError as exc:
            if self.verbose:
                print(f"⚠️  Math engine unavailable ({exc}) — math will render as plain text")

    def _latex_to_omml_node(self, latex_str: str):
        """
        Convert a LaTeX math string (without surrounding $) to an OMML
        lxml Element ready for insertion into a paragraph's XML tree.
        Returns None if conversion is unavailable or fails.
        """
        self._init_math_engine()
        if not self._math_ready:
            return None
        try:
            etree = self._lxml_etree
            mml_str = self._latex2mml_fn(latex_str)
            mml_tree = etree.fromstring(mml_str.encode('utf-8'))
            omml_result = self._math_xslt(mml_tree)
            return omml_result.getroot()
        except Exception as exc:
            if self.verbose:
                print(f"⚠️  LaTeX→OMML failed for '{latex_str[:40]}': {exc}")
            return None

    def _add_text_with_math(self, para, text: str,
                            bold: bool = False,
                            italic: bool = False,
                            color_rgb: Optional[RGBColor] = None):
        """
        Add *text* to *para*, rendering any ``$...$ `` spans as native Word
        equations (OMML) via the latex2mathml + MML2OMML pipeline.
        Falls back to plain ``$...$`` text if the engine is unavailable.
        """
        if not text:
            return
        # Split on un-escaped $...$ (non-greedy, single-line content)
        parts = re.split(r'(?<!\\)\$(.+?)(?<!\\)\$', str(text), flags=re.DOTALL)
        for i, part in enumerate(parts):
            if i % 2 == 0:
                # Plain-text segment
                part = part.replace('\\$', '$')  # unescape \$
                if part:
                    run = para.add_run(part)
                    run.bold = bold
                    run.italic = italic
                    if color_rgb:
                        run.font.color.rgb = color_rgb
            else:
                # LaTeX math segment
                omml_node = self._latex_to_omml_node(part)
                if omml_node is not None:
                    para._p.append(omml_node)
                else:
                    # Fallback: show raw LaTeX
                    run = para.add_run(f"${part}$")
                    run.bold = bold
                    run.italic = italic
                    if color_rgb:
                        run.font.color.rgb = color_rgb

    # ────────────────────────────────────────────────────────────────────

    def _add_mixed_content(self, para, content: List):
        """Xử lý mixed content: text + table + chart"""
        for item in content:
            if isinstance(item, str):
                # Text content
                para.add_run(item)
            elif isinstance(item, dict):
                item_type = item.get('type', '')
                
                if item_type == 'table':
                    # Add line break before table
                    para.add_run('\n')
                    self._add_inline_table(item.get('content', {}), item.get('metadata', {}))
                    
                elif item_type == 'chart':
                    # Add line break before chart
                    para.add_run('\n')
                    # ✨ NEW: Look up Base64 image for this chart
                    chart_key = f"{self.current_question_code}-{self.current_chart_index}"
                    image_base64 = self.chart_images.get(chart_key)
                    self._add_inline_chart(item.get('content', {}), image_base64=image_base64)
                    self.current_chart_index += 1
                    
                elif item_type == 'image':
                    # Add line break before image
                    para.add_run('\n')
                    caption = item.get('metadata', {}).get('caption', '')
                    if caption:
                        para.add_run(f"[Hình ảnh: {caption}]")
                    else:
                        para.add_run(f"[Hình ảnh: {item.get('content', '')}]")
                    para.add_run('\n')
                    
                elif item_type == 'chart' and 'chart_id' in item:
                    # Chart reference (không có echarts data, chỉ có chart_id)
                    para.add_run('\n')
                    para.add_run(f"[Biểu đồ {item.get('chart_id')}]")
                    para.add_run('\n')
    
    def _render_question_stem(self, para, question_stem):
        """
        Render question_stem (dict or plain string) into the document.

        *para* receives the first inline text segment exactly as typed in the
        question-number line.  Every subsequent text segment and every embedded
        table / chart / image is written as a new standalone element so the
        order text → table/chart → closing-text is always preserved correctly
        for ALL stem types: text, table, chart, mixed.
        """
        if not isinstance(question_stem, dict):
            self._add_text_with_math(para, str(question_stem))
            return

        stem_type = question_stem.get('type', 'text')
        content = question_stem.get('content', '')

        # Simple text stem
        if stem_type == 'text':
            self._add_text_with_math(para, str(content) if not isinstance(content, list) else '')
            return

        # List-based stems: table / chart / mixed – all handled identically
        if isinstance(content, list):
            first_text_done = False
            for item in content:
                if isinstance(item, str):
                    if not first_text_done:
                        self._add_text_with_math(para, item)
                        first_text_done = True
                    else:
                        new_para = self.document.add_paragraph()
                        self._add_text_with_math(new_para, item)
                elif isinstance(item, dict):
                    item_type = item.get('type', '')
                    if item_type == 'table':
                        self._add_inline_table(item.get('content', {}), item.get('metadata', {}))
                    elif item_type == 'chart':
                        # ✨ NEW: Look up Base64 image for this chart
                        chart_key = f"{self.current_question_code}-{self.current_chart_index}"
                        image_base64 = self.chart_images.get(chart_key)
                        self._add_inline_chart(item.get('content', {}), image_base64=image_base64)
                        self.current_chart_index += 1
                    elif item_type == 'image':
                        caption = item.get('metadata', {}).get('caption', '')
                        para_img = self.document.add_paragraph()
                        text = f"[Hình ảnh: {caption}]" if caption else f"[Hình ảnh: {item.get('content', '')}]"
                        para_img.add_run(text).italic = True
        else:
            # Fallback for non-list content
            self._add_text_with_math(para, str(content))

    def _add_inline_table(self, table_data: Dict, metadata: Dict = None):
        """Thêm bảng inline vào document"""
        if not table_data:
            return
        
        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])
        
        if not headers or not rows:
            return
        
        # Caption trên bảng
        meta = metadata if metadata else table_data.get('metadata', {})
        caption = meta.get('caption', '')
        source = meta.get('source', '')
        
        if caption:
            para = self.document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(caption)
            run.bold = True
            run.font.size = Pt(10)
        
        # Tạo bảng với border đơn giản, không tô màu
        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'  # Border đơn giản, không màu
        
        # Thêm header
        header_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            # Clear default empty run, then render with math support
            cell_para = header_cells[idx].paragraphs[0]
            cell_para.clear()
            self._add_text_with_math(cell_para, str(header), bold=True)
            # Apply font size to all resulting runs
            for run in cell_para.runs:
                run.bold = True
                run.font.size = Pt(11)
        
        # Thêm dữ liệu
        for row_data in rows:
            row_cells = table.add_row().cells
            for idx, cell_value in enumerate(row_data):
                if idx < len(row_cells):
                    # Render cell text with LaTeX math support
                    cell_para = row_cells[idx].paragraphs[0]
                    cell_para.clear()
                    self._add_text_with_math(cell_para, str(cell_value))
                    # Apply font size to all resulting runs
                    for run in cell_para.runs:
                        run.font.size = Pt(11)
        
        # Source ở dưới bảng, căn phải
        if source:
            para = self.document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = para.add_run(f"(Nguồn: {source})")
            run.italic = True
            run.font.size = Pt(10)
    
    def _add_inline_chart(self, chart_data: Dict, image_base64: Optional[str] = None):
        """
        ✨ SIMPLIFIED: Thêm chart vào document từ Base64 image ONLY
        
        Playwright rendering loại bỏ hoàn toàn.
        Chart image PHẢI được cung cấp bởi frontend (từ canvas.toDataURL()).
        
        Args:
            chart_data: Dict chứa chartType, echarts config, etc.
            image_base64: Base64 PNG từ client canvas (REQUIRED từ frontend)
        """
        if not chart_data:
            if self.verbose:
                print("⚠️  Chart data is empty")
            return
        
        chart_type = chart_data.get('chartType', 'bar')
        echarts = chart_data.get('echarts', {})
        
        # Extract title from echarts config
        title_data = echarts.get('title', {})
        if isinstance(title_data, list):
            title = title_data[0].get('text', 'Biểu đồ') if title_data else 'Biểu đồ'
        elif isinstance(title_data, dict):
            title = title_data.get('text', 'Biểu đồ')
        else:
            title = 'Biểu đồ'
        
        # ✨ ONLY USE Base64 from client - NO PLAYWRIGHT FALLBACK
        if not image_base64:
            # Chart image not provided by frontend
            if self.verbose:
                print(f"⚠️  [SKIP] No chart image from client for: {title}")
            return  # Skip chart if no image provided
        
        try:
            if self.verbose:
                print(f"📊 [{chart_type.upper()}] Processing chart from Base64: {title}")
            
            # Strip data URL prefix if present
            if ',' in image_base64:
                image_base64 = image_base64.split(',')[1]
            
            # Decode Base64 → bytes
            image_bytes = base64.b64decode(image_base64)
            
            # Save temporarily
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
                temp_image_path = f.name
                f.write(image_bytes)
            
            # Insert image into DOCX with appropriate size
            para = self.document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            
            # Adaptive width based on chart type
            if chart_type.lower() == 'pie':
                # Pie: square aspect ratio (1:1)
                run.add_picture(temp_image_path, width=Inches(4.0))
            else:
                # Bar/Line/Area: landscape aspect ratio (1.7:1)
                run.add_picture(temp_image_path, width=Inches(5.5))
            
            # Cleanup temp file
            try:
                os.remove(temp_image_path)
            except:
                pass
            
            if self.verbose:
                print(f"✅ Chart embedded: {title} ({len(image_bytes)/1024:.1f}KB)")
        
        except Exception as e:
            if self.verbose:
                print(f"❌ Failed to embed chart image: {e}")
                import traceback
                traceback.print_exc()
    
    def _resolve_chart_placeholders_to_js(self, echarts_config: Dict) -> str:
        """
        Convert echarts config với placeholders thành JavaScript code
        
        Args:
            echarts_config: ECharts config dict có chứa placeholders
            
        Returns:
            JavaScript code string defining the option variable
        """
        import re
        
        # Deep copy để không modify original
        config_copy = json.loads(json.dumps(echarts_config))
        
        # Tăng font-size cho xuất DOCX (tăng 4 size để đủ lớn với deviceScaleFactor=2)
        font_changes = []
        def increase_font_size(obj, path=''):
            if isinstance(obj, dict):
                for key, value in obj.items():
                    if key == 'fontSize' and isinstance(value, (int, float)):
                        old_val = obj[key]
                        obj[key] = value + 4  # Tăng 4 thay vì 2
                        font_changes.append(f"{path}.{key}: {old_val} → {obj[key]}")
                    elif isinstance(value, (dict, list)):
                        increase_font_size(value, f"{path}.{key}" if path else key)
            elif isinstance(obj, list):
                for i, item in enumerate(obj):
                    increase_font_size(item, f"{path}[{i}]")
        
        increase_font_size(config_copy)
        
        if self.verbose and font_changes:
            print(f"  Font size increases: {len(font_changes)} changes")
            for change in font_changes[:5]:  # Show first 5
                print(f"    {change}")
        
        # Process _patternType - có thể ở series level (bar/line) hoặc data level (pie)
        pattern_replacements = []
        if 'series' in config_copy:
            for idx, series in enumerate(config_copy['series']):
                # Case 1: _patternType ở series level (Bar, Line, Area)
                if '_patternType' in series:
                    pattern_type = series['_patternType']
                    # Ensure itemStyle exists
                    if 'itemStyle' not in series:
                        series['itemStyle'] = {}
                    # Set color as placeholder that will be replaced
                    series['itemStyle']['color'] = f'PATTERN_PLACEHOLDER_{pattern_type}'
                    pattern_replacements.append(f"series[{idx}]: {pattern_type}")
                    # Remove the meta field
                    del series['_patternType']
                
                # Case 2: _patternType ở data item level (Pie)
                if 'data' in series and isinstance(series['data'], list):
                    for item_idx, item in enumerate(series['data']):
                        if isinstance(item, dict) and '_patternType' in item:
                            pattern_type = item['_patternType']
                            # Ensure itemStyle exists
                            if 'itemStyle' not in item:
                                item['itemStyle'] = {}
                            # Set color as placeholder that will be replaced
                            item['itemStyle']['color'] = f'PATTERN_PLACEHOLDER_{pattern_type}'
                            pattern_replacements.append(f"series[{idx}].data[{item_idx}]: {pattern_type}")
                            # Remove the meta field
                            del item['_patternType']
        
        if self.verbose and pattern_replacements:
            print(f"  Pattern replacements: {len(pattern_replacements)}")
            for pr in pattern_replacements:
                print(f"    {pr}")
        
        # Convert thành JSON string
        json_str = json.dumps(config_copy, ensure_ascii=False, indent=2)
        
        # Get yAxis max value for formatter (nếu có)
        y_axis_max = config_copy.get('yAxis', {})
        if isinstance(y_axis_max, list):
            y_axis_max = y_axis_max[0].get('max') if y_axis_max else None
        else:
            y_axis_max = y_axis_max.get('max')
        
        # Replace pattern placeholders
        # PATTERN_PLACEHOLDER_dots → {type: 'pattern', image: createPattern('dots'), repeat: 'repeat'}
        pattern_count = len(re.findall(r'"PATTERN_PLACEHOLDER_\w+"', json_str))
        json_str = re.sub(
            r'"PATTERN_PLACEHOLDER_(\w+)"',
            lambda m: '{type: "pattern", image: createPattern("' + m.group(1) + '"), repeat: "repeat"}',
            json_str
        )
        
        if self.verbose and pattern_count > 0:
            print(f"  Replaced {pattern_count} pattern placeholders in JSON")
        
        # Replace formatter placeholders
        replacements = {
            '"FORMATTER_PLACEHOLDER"': f'''function(val) {{
                if ({y_axis_max} && val > {y_axis_max} - 1) return '';
                return '{{value|' + val.toFixed(1).replace('.', ',') + '}}{{tick|}}';
            }}''',
            
            '"FORMATTER_LABEL_PLACEHOLDER_BAR"': '''function(params) {
                let val = params.value;
                if (typeof val === 'number') {
                    if (val > 1000000) return (val / 1000000).toFixed(1).replace('.', ',') + 'M';
                    if (val > 1000) return (val / 1000).toFixed(1).replace('.', ',') + 'K';
                    return val.toFixed(1).replace('.', ',');
                }
                return val?.toString() || '';
            }''',
            
            '"FORMATTER_LABEL_PLACEHOLDER_PIE"': '''function(params) {
                if (params.value !== null && params.value !== undefined) {
                    if (typeof params.value === 'number') {
                        return params.value.toFixed(1).replace('.', ',');
                    }
                    return params.value.toString();
                }
                return '';
            }''',
            
            '"FORMATTER_LABEL_PLACEHOLDER"': '''function(params) {
                // Bỏ label cho điểm đầu tiên (tiếp điểm với trục tung)
                if (params.dataIndex === 0) {
                    return '';
                }
                let val = params.value;
                if (typeof val === 'number') {
                    if (val > 1000000) return (val / 1000000).toFixed(1).replace('.', ',') + 'M';
                    if (val > 1000) return (val / 1000).toFixed(1).replace('.', ',') + 'K';
                    return val.toFixed(1).replace('.', ',');
                }
                return val?.toString() || '';
            }''',
            
            '"FORMATTER_SCATTER_LABEL_PLACEHOLDER"': '''function(params) {
                // params.value là array [x_index, y_value] cho scatter series
                // Lấy realValue từ data nếu có, nếu không thì dùng y_value
                if (params.data && params.data.realValue !== undefined) {
                    let val = params.data.realValue;
                    if (typeof val === 'number') {
                        return val.toFixed(1).replace('.', ',');
                    }
                    return val?.toString() || '';
                }
                // Fallback: lấy từ value array
                if (Array.isArray(params.value)) {
                    let val = params.value[1];
                    if (typeof val === 'number') {
                        return val.toFixed(1).replace('.', ',');
                    }
                    return val?.toString() || '';
                }
                return '';
            }''',
            
            '"FORMATTER_X_PLACEHOLDER"': '''function(value) {
                if (typeof value === 'string' && value.startsWith('_')) {
                    return '';
                }
                return value.toString();
            }''',
            
            '"FORMATTER_X_INTERVAL_PLACEHOLDER"': '''function(index, value) {
                if (typeof value === 'string' && value.startsWith('_')) {
                    return false;
                }
                return true;
            }'''
        }
        
        for placeholder, js_func in replacements.items():
            json_str = json_str.replace(placeholder, js_func)
        
        return f"var option = {json_str};"
    
    def _render_chart_to_image(self, echarts_config: Dict) -> Optional[str]:
        """
        Render ECharts config thành PNG image sử dụng Playwright
        
        Args:
            echarts_config: ECharts configuration dict
            
        Returns:
            Path đến file PNG tạm, hoặc None nếu thất bại
        """
        if not PLAYWRIGHT_AVAILABLE:
            return None
        
        try:
            # Apply layout resolution (tính grid.bottom nếu chưa có)
            echarts_config = apply_layout(echarts_config)
            
            # Tắt animation để đảm bảo screenshot đầy đủ
            if 'animation' not in echarts_config:
                echarts_config['animation'] = False
            
            # Resolve local echarts.min.js (bundled with app to avoid CDN dependency)
            _here = Path(getattr(sys, '_MEIPASS', Path(__file__).parent))
            _echarts_candidates = [
                Path(__file__).parent / 'echarts.min.js',       # dev: same dir
                _here / 'server' / 'src' / 'services' / 'exporters' / 'echarts.min.js',  # frozen
            ]
            _echarts_script_tag = '<script src="https://cdn.jsdelivr.net/npm/echarts@6.0.0/dist/echarts.min.js"></script>'
            for _p in _echarts_candidates:
                if _p.exists():
                    with open(_p, 'r', encoding='utf-8') as _f:
                        _echarts_js = _f.read()
                    _echarts_script_tag = f'<script>{_echarts_js}</script>'
                    break

            # Resolve placeholders thành JavaScript code
            option_js = self._resolve_chart_placeholders_to_js(echarts_config)
            
            # Import createPattern function từ utils
            from services.utils.chart.utils import create_pattern_js_function
            pattern_function = create_pattern_js_function()

            # Tạo HTML chứa ECharts với animation tắt
            # Chart dimensions: 1200x900 để có tỷ lệ 4:3 phù hợp với DOCX
            html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <link href="https://fonts.googleapis.com/css2?family=Roboto:wght@400;700&display=swap" rel="stylesheet">
    {_echarts_script_tag}
    <style>
        body {{ 
            margin: 0; 
            padding: 0; 
            background: white; 
            font-family: 'Roboto', sans-serif;
            display: flex;
            justify-content: center;
            align-items: center;
            width: 1200px;
            height: 900px;
        }}
        #chart {{ 
            width: 1200px; 
            height: 900px; 
        }}
    </style>
</head>
<body>
    <div id="chart"></div>
    <script>
        {pattern_function}
        
        var chartDom = document.getElementById('chart');
        var myChart = echarts.init(chartDom);
        {option_js}
        
        // Set option với animation tắt
        myChart.setOption(option);
        
        // Đợi render hoàn tất
        setTimeout(function() {{
            window.chartReady = true;
        }}, 500);
    </script>
</body>
</html>
"""
            
            # Tạo file PNG tạm
            with tempfile.NamedTemporaryFile(mode='w', suffix='.png', delete=False) as f:
                png_path = f.name
            
            # Result container for thread
            result = {'path': None, 'error': None}
            
            def render_in_thread():
                try:
                    # Sử dụng Playwright để render (trong thread riêng để tránh xung đột với asyncio)
                    with sync_playwright() as p:
                        # Thử theo thứ tự ưu tiên:
                        # 1. Playwright Chromium (cần PLAYWRIGHT_BROWSERS_PATH đúng)
                        # 2. System Chrome (phổ biến trên máy user)
                        # 3. System Edge (Windows built-in)
                        browser = None
                        last_err = None
                        for launch_kwargs in [
                            {},                          # playwright's own chromium
                            {'channel': 'chrome'},       # system Chrome
                            {'channel': 'msedge'},       # system Edge
                        ]:
                            try:
                                browser = p.chromium.launch(headless=True, **launch_kwargs)
                                break
                            except Exception as e:
                                last_err = e
                                continue
                        
                        if browser is None:
                            raise Exception(f"Cannot launch any browser. Last error: {last_err}")
                        
                        # Viewport phù hợp với chart 1200x900 (tỷ lệ 4:3)
                        # Set deviceScaleFactor = 2 để ảnh sắc nét hơn (Retina-quality)
                        page = browser.new_page(
                            viewport={'width': 1200, 'height': 900},
                            device_scale_factor=2
                        )
                        
                        # Load HTML content directly (better than file:// for security)
                        page.set_content(html_content)
                        
                        # Đợi chart render xong
                        page.wait_for_function('window.chartReady === true', timeout=10000)
                        
                        # Đợi thêm để đảm bảo chart render hoàn toàn (tăng từ 1s lên 2s)
                        page.wait_for_timeout(2000)
                        
                        # Screenshot full page thay vì chỉ element để tránh lệch
                        page.screenshot(path=png_path, type='png', full_page=False)
                        
                        browser.close()
                        result['path'] = png_path
                except Exception as e:
                    result['error'] = str(e)
            
            try:
                # Chạy render trong thread riêng (tăng timeout cho lần đầu download playwright)
                thread = threading.Thread(target=render_in_thread)
                thread.start()
                thread.join(timeout=30)  # Wait max 30 seconds (lần đầu cần download browser)
                
                if result['error']:
                    raise Exception(result['error'])
                
                if result['path']:
                    return result['path']
                else:
                    # Timeout or no result
                    try:
                        os.remove(png_path)
                    except:
                        pass
                    return None
            except Exception as e:
                # Cleanup
                try:
                    os.remove(png_path)
                except:
                    pass
                raise
                
        except Exception as e:
            if self.verbose:
                print(f"✗ Lỗi khi render chart: {e}")
            return None
    
    def generate_from_json(self, 
                          json_data: Union[Dict, str],
                          template: Optional[Dict] = None):
        """
        Tạo document từ dữ liệu JSON
        
        Args:
            json_data (Union[Dict, str]): Dữ liệu JSON hoặc đường dẫn file JSON
            template (Dict, optional): Template định nghĩa cấu trúc document
        """
        if self.document is None:
            self.create_new_document()
        
        # Đọc JSON nếu là đường dẫn
        if isinstance(json_data, str):
            with open(json_data, 'r', encoding='utf-8') as f:
                json_data = json.load(f)
        
        # Nếu có template, sử dụng template
        if template:
            self._apply_template(json_data, template)
        else:
            # Mặc định: tạo document đơn giản từ JSON
            self._generate_default(json_data)
    
    def _apply_template(self, data: Dict, template: Dict):
        """Áp dụng template để tạo document"""
        for section in template.get('sections', []):
            section_type = section.get('type')
            
            if section_type == 'heading':
                self.add_heading(
                    text=section.get('text', '').format(**data),
                    level=section.get('level', 1),
                    alignment=section.get('alignment', 'left')
                )
            
            elif section_type == 'paragraph':
                self.add_paragraph(
                    text=section.get('text', '').format(**data),
                    bold=section.get('bold', False),
                    italic=section.get('italic', False),
                    alignment=section.get('alignment', 'left')
                )
            
            elif section_type == 'table':
                data_key = section.get('data_key')
                if data_key and data_key in data:
                    self.add_table_from_dict(
                        data=data[data_key],
                        headers=section.get('headers')
                    )
            
            elif section_type == 'page_break':
                self.add_page_break()
    
    def _generate_default(self, data: Dict):
        """Tạo document mặc định từ JSON"""
        for key, value in data.items():
            # Thêm key làm heading
            self.add_heading(str(key).replace('_', ' ').title(), level=2)
            
            # Thêm value
            if isinstance(value, list) and value and isinstance(value[0], dict):
                # Nếu là list of dict, tạo bảng
                self.add_table_from_dict(value)
            elif isinstance(value, dict):
                # Nếu là dict, hiển thị từng cặp key-value
                for k, v in value.items():
                    self.add_paragraph(f"{k}: {v}")
            else:
                # Giá trị đơn giản
                self.add_paragraph(str(value))
            
            self.add_paragraph("")  # Thêm dòng trống
    
    def save(self, output_path: str):
        """
        Lưu document
        
        Args:
            output_path (str): Đường dẫn file output
        """
        if self.document is None:
            print("✗ Chưa có document để lưu")
            return
        
        try:
            # Tạo thư mục nếu chưa tồn tại
            output_dir = Path(output_path).parent
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # Lưu file
            self.document.save(output_path)
            self.output_path = output_path
            
            if self.verbose:
                print(f"✓ Đã lưu file: {output_path}")
        
        except Exception as e:
            print(f"✗ Lỗi khi lưu file: {str(e)}")
            raise
    
    def generate_questions_document(self, 
                                   json_data: Union[Dict, str],
                                   output_path: str,
                                   chart_images: Dict[str, str] = None):
        """
        Tạo document câu hỏi từ dữ liệu JSON
        
        Args:
            json_data (Union[Dict, str]): Dữ liệu JSON hoặc đường dẫn file JSON
            output_path (str): Đường dẫn file DOCX output
            chart_images (Dict[str, str]): Optional dict mapping "question_code-chart_index" to Base64 images
        """
        # Đọc JSON nếu là đường dẫn
        if isinstance(json_data, str):
            with open(json_data, 'r', encoding='utf-8') as f:
                json_data = json.load(f)
        
        # ✨ NEW: Store chart images for use during rendering
        self.chart_images = chart_images or {}
        self.current_question_code = None  # Track current question for chart lookup
        self.current_chart_index = 0  # Track chart index within question
        
        # Tạo document mới
        self.create_new_document()
        self.set_document_style()
        
        # Thêm tiêu đề
        metadata = json_data.get('metadata', {})
        self.add_heading('ĐỀ THI TRẮC NGHIỆM', level=1, alignment='center')
        
        matrix_file = metadata.get('matrix_file', 'N/A')
        generated_at = metadata.get('generated_at', 'N/A')
        if isinstance(generated_at, str) and len(generated_at) >= 10:
            generated_at = generated_at[:10]
        total_questions = metadata.get('total_questions', 0)
        
        self.add_paragraph(f"Ma trận: {matrix_file}", alignment='center')
        self.add_paragraph(f"Ngày tạo: {generated_at}", alignment='center')
        self.add_paragraph(f"Tổng số câu: {total_questions}", alignment='center')
        self.add_paragraph("")
        
        questions_data = json_data.get('questions', {})
        
        # Lấy subject từ metadata
        subject = metadata.get('subject', '')
        
        # Xuất câu hỏi TN
        tn_questions = questions_data.get('TN', [])
        if tn_questions:
            try:
                tn_questions_sorted = sorted(tn_questions, key=lambda x: int(x.get('question_code', 'C0')[1:]))
            except (ValueError, TypeError, AttributeError) as e:
                tn_questions_sorted = tn_questions
            
            self.add_heading('PHẦN I. Thí sinh trả lời từ câu 1 đến câu 24. Mỗi câu hỏi thí sinh chỉ chọn một phương án.', level=2)
            
            for idx, q in enumerate(tn_questions_sorted, 1):
                self.current_question_code = q.get('question_code', f'TN{idx}')
                self.current_chart_index = 0  # Reset chart index for new question
                self._add_tn_question(q, idx)
            
            self.add_paragraph("")
        
        # Xuất câu hỏi DS
        ds_questions = questions_data.get('DS', [])
        if ds_questions:
            try:
                ds_questions_sorted = sorted(ds_questions, key=lambda x: int(x.get('question_code', 'C0')[1:]))
            except (ValueError, TypeError, AttributeError) as e:
                ds_questions_sorted = ds_questions
            
            self.add_heading('PHẦN II. Thí sinh trả lời từ câu 1 đến 4. Trong mỗi ý a), b), c), d) ở mỗi câu, thí sinh chọn đúng hoặc sai.', level=2)
            
            for idx, q in enumerate(ds_questions_sorted, 1):
                self.current_question_code = q.get('question_code', f'DS{idx}')
                self.current_chart_index = 0  # Reset chart index for new question
                self._add_ds_question(q, idx, subject)
            
            self.add_paragraph("")
        
        # Xuất câu hỏi TLN
        tln_questions = questions_data.get('TLN', [])
        if tln_questions:
            try:
                tln_questions_sorted = sorted(tln_questions, key=lambda x: int(x.get('question_code', 'C0')[1:]))
            except (ValueError, TypeError, AttributeError) as e:
                tln_questions_sorted = tln_questions
            
            self.add_heading('PHẦN III. Câu trả lời ngắn. Thí sinh viết đáp án vào chỗ trống.', level=2)
            
            for idx, q in enumerate(tln_questions_sorted, 1):
                self.current_question_code = q.get('question_code', f'TLN{idx}')
                self.current_chart_index = 0  # Reset chart index for new question
                self._add_tln_question(q, idx)
            
            self.add_paragraph("")
        
        # Xuất câu hỏi TL
        tl_questions = questions_data.get('TL', [])
        if tl_questions:
            try:
                tl_questions_sorted = sorted(tl_questions, key=lambda x: int(x.get('question_code', 'C0')[1:]))
            except (ValueError, TypeError, AttributeError) as e:
                tl_questions_sorted = tl_questions
            
            self.add_heading('PHẦN IV. Câu tự luận. Thí sinh trình bày lập luận đầy đủ.', level=2)
            
            for idx, q in enumerate(tl_questions_sorted, 1):
                self._add_tl_question(q, idx)
            
            self.add_paragraph("")
        
        # Thêm phần đáp án
        self.add_page_break()
        self.add_heading('ĐÁP ÁN VÀ GIẢI THÍCH', level=1, alignment='center')
        self.add_paragraph("")
        
        # Đáp án TN
        if tn_questions:
            self.add_heading('PHẦN I: TRẮC NGHIỆM', level=2)
            
            for idx, q in enumerate(tn_questions_sorted, 1):
                self._add_tn_answer(q, idx)
        
        # Đáp án DS
        if ds_questions:
            self.add_heading('PHẦN II: ĐÚNG/SAI', level=2)
            
            for idx, q in enumerate(ds_questions_sorted, 1):
                self._add_ds_answer(q, idx)
        
        # Đáp án TLN
        if tln_questions:
            self.add_heading('PHẦN III: TRẢ LỜI NGẮN', level=2)
            
            for idx, q in enumerate(tln_questions_sorted, 1):
                self._add_tln_answer(q, idx)
        
        # Đáp án TL
        if tl_questions:
            self.add_heading('PHẦN IV: TỰ LUẬN', level=2)
            
            for idx, q in enumerate(tl_questions_sorted, 1):
                self._add_tl_answer(q, idx)
        
        # Lưu file
        self.save(output_path)
    
    def _add_tn_question(self, question: Dict, number: int):
        """Thêm câu hỏi trắc nghiệm"""
        try:
            # Câu hỏi
            para = self.document.add_paragraph()
            
            level = question.get('level', 'NB')
            
            run = para.add_run(f"Câu {number} ")
            run.bold = True
            run = para.add_run(f"({level})")
            run.bold = True
            self._set_run_background(run, level)
            
            para.add_run(". ")
            
            # Handle question_stem - renders text/table/chart in correct order
            question_stem = question.get('question_stem', '')
            self._render_question_stem(para, question_stem)
            
            # Các lựa chọn
            options = question.get('options', {})
            correct_answer = question.get('correct_answer', '')
            
            # Ensure options is a dict
            if not isinstance(options, dict):
                options = {}
            
            for key in ['A', 'B', 'C', 'D']:
                if key in options:
                    option_text = options[key]
                    if not isinstance(option_text, str):
                        option_text = str(option_text)
                    
                    para = self.document.add_paragraph()
                    
                    # Tô đỏ nếu là đáp án đúng
                    if key == correct_answer:
                        run_key = para.add_run(f"{key}. ")
                        run_key.font.color.rgb = RGBColor(255, 0, 0)
                        self._add_text_with_math(para, option_text, color_rgb=RGBColor(255, 0, 0))
                    else:
                        para.add_run(f"{key}. ")
                        self._add_text_with_math(para, option_text)
        except Exception as e:
            print(f"Error in _add_tn_question {number}: {e}")
            print(f"Question data: {question}")
            raise
    
    def _add_ds_question(self, question: Dict, number: int, subject: str = ''):
        """Thêm câu hỏi đúng/sai"""
        try:
            # Tiêu đề câu
            para = self.document.add_paragraph()
            para.add_run(f"Câu {number}. ").bold = True
            para.add_run("Cho đoạn tư liệu sau:")
            
            # Always display source_text.content for DS questions
            source_text = question.get('source_text', '')
            source_metadata = None
            if isinstance(source_text, dict):
                source_metadata = source_text.get('metadata', {})
                source_text = source_text.get('content', '')
            
            # Handle mixed content or plain text
            if isinstance(source_text, list):
                # Mixed content: process each item separately to preserve order
                for item in source_text:
                    if isinstance(item, str):
                        # Plain text - create new paragraph
                        para = self.document.add_paragraph()
                        self._add_text_with_math(para, item, italic=True)
                    elif isinstance(item, dict):
                        item_type = item.get('type', '')
                        if item_type == 'table':
                            self._add_inline_table(item.get('content', {}), item.get('metadata', {}))
                        elif item_type == 'image':
                            caption = item.get('metadata', {}).get('caption', '')
                            para = self.document.add_paragraph()
                            if caption:
                                run = para.add_run(f"[Hình ảnh: {caption}]")
                            else:
                                run = para.add_run(f"[Hình ảnh: {item.get('content', '')}]")
                            run.italic = True
                        elif item_type == 'chart':
                            self._add_inline_chart(item.get('content', {}))
                        else:
                            para = self.document.add_paragraph()
                            run = para.add_run(str(item))
                            run.italic = True
            else:
                # Plain text - use _add_text_with_math to render LaTeX equations
                para = self.document.add_paragraph()
                self._add_text_with_math(para, source_text, italic=True)
            
            # Only display source citation if subject needs source display
            should_display_source = Config.should_display_source(subject)
            if should_display_source:
                # Thêm nguồn tư liệu (căn lề phải, bọc trong ngoặc đơn)
                # Ưu tiên source_citation, nếu không có thì dùng source từ metadata
                source_citation = question.get('source_citation', '')
                if not source_citation and source_metadata:
                    source_citation = source_metadata.get('source', '')
                
                if source_citation:
                    para_source = self.document.add_paragraph()
                    para_source.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    para_source.add_run(f"({source_citation})").italic = True
            
            statements = question.get('statements', {})
            
            # Ensure statements is a dict
            if not isinstance(statements, dict):
                statements = {}
            
            for label in ['a', 'b', 'c', 'd']:
                if label in statements:
                    stmt = statements[label]
                    level = stmt.get('level', 'NB')
                    text = stmt.get('text', '')
                    if not isinstance(text, str):
                        text = str(text)
                    is_correct = stmt.get('correct_answer', False)
                    
                    para = self.document.add_paragraph()
                    
                    # "(level)" - tô màu nền
                    run_level = para.add_run(f"({level})")
                    self._set_run_background(run_level, level)
                    
                    # " label. text" - tô đỏ nếu đúng
                    run_label = para.add_run(f" {label}) ")
                    if is_correct:
                        run_label.font.color.rgb = RGBColor(255, 0, 0)
                    self._add_text_with_math(para, text, color_rgb=RGBColor(255, 0, 0) if is_correct else None)
        except Exception as e:
            print(f"Error in _add_ds_question {number}: {e}")
            print(f"Question data: {question}")
            raise
    
    def _add_tn_answer(self, question: Dict, number: int):
        """Thêm đáp án trắc nghiệm"""
        para = self.document.add_paragraph()
        para.add_run(f"Câu {number}. ").bold = True
        
        # Đáp án đúng tô đỏ
        run = para.add_run(f"Đáp án: {question.get('correct_answer', 'N/A')}")
        run.bold = True
        run.font.color.rgb = RGBColor(255, 0, 0)
        
        # Giải thích
        explanation = question.get('explanation', '')
        if explanation:
            para_exp = self.document.add_paragraph()
            self._add_text_with_math(para_exp, str(explanation))
        
        self.add_paragraph("")
    
    def _add_ds_answer(self, question: Dict, number: int):
        """Thêm đáp án đúng/sai"""
        para = self.document.add_paragraph()
        para.add_run(f"Câu {number}. ").bold = True
        
        # Đáp án từng mệnh đề
        statements = question.get('statements', {})
        answer_parts = []
        for label in ['a', 'b', 'c', 'd']:
            if label in statements:
                stmt = statements[label]
                is_correct = stmt.get('correct_answer', False)
                answer_text = "Đ" if is_correct else "S"
                answer_parts.append(f"{label}. {answer_text}")
        
        # Tô đỏ đáp án
        run = para.add_run(", ".join(answer_parts))
        run.font.color.rgb = RGBColor(255, 0, 0)
        
        # Giải thích
        explanations = question.get('explanation', {})
        
        # Ensure explanations is a dict
        if not isinstance(explanations, dict):
            explanations = {}
        
        if explanations:
            self.add_paragraph("Giải thích:")
            for label in ['a', 'b', 'c', 'd']:
                if label in explanations:
                    explanation_text = explanations[label]
                    if not isinstance(explanation_text, str):
                        explanation_text = str(explanation_text)
                    para_exp = self.document.add_paragraph()
                    para_exp.add_run(f"{label}. ")
                    self._add_text_with_math(para_exp, explanation_text)
        
        self.add_paragraph("")
    
    def _add_tl_question(self, question: Dict, number: int):
        """Thêm câu hỏi tự luận"""
        try:
            # Câu hỏi
            para = self.document.add_paragraph()
            
            level = question.get('level', 'NB')
            
            run = para.add_run(f"Câu {number} ")
            run.bold = True
            run = para.add_run(f"({level})")
            run.bold = True
            self._set_run_background(run, level)
            
            para.add_run(". ")
            
            # Handle question_stem - renders text/table/chart in correct order
            question_stem = question.get('question_stem', '')
            self._render_question_stem(para, question_stem)
            
            # Sub-questions (a, b, ...)
            explanation = question.get('explanation', {})
            sub_questions = []
            if isinstance(explanation, dict):
                sub_questions = explanation.get('sub_questions', [])
            if sub_questions:
                for sub in sub_questions:
                    label = sub.get('label', '')
                    sub_para = self.document.add_paragraph()
                    run_label = sub_para.add_run(f"{label}) ")
                    run_label.bold = True
                    self._render_question_stem(sub_para, sub.get('question_stem', ''))
            
        except Exception as e:
            print(f"Error in _add_tl_question {number}: {e}")
            print(f"Question data: {question}")
            raise
    
    def _add_tl_answer(self, question: Dict, number: int):
        """Thêm đáp án câu hỏi TL"""
        para = self.document.add_paragraph()
        para.add_run(f"Câu {number}. ").bold = True
        
        # Đáp án mẫu
        correct_answer = question.get('correct_answer', '')
        if correct_answer:
            run_label = para.add_run("Đáp án mẫu: ")
            run_label.bold = True
            run_label.font.color.rgb = RGBColor(255, 0, 0)
            para_ans = self.document.add_paragraph()
            self._add_text_with_math(para_ans, str(correct_answer))
        
        # Hướng dẫn chấm điểm
        explanation = question.get('explanation', '')
        if explanation:
            if isinstance(explanation, dict):
                sub_questions = explanation.get('sub_questions', [])
                if sub_questions:
                    # Multi-sub-question answer
                    para_hd = self.document.add_paragraph()
                    para_hd.add_run('Hướng dẫn chấm điểm:').bold = True
                    for sub in sub_questions:
                        label = sub.get('label', '')
                        para_sub = self.document.add_paragraph()
                        run_sub_label = para_sub.add_run(f'{label}) ')
                        run_sub_label.bold = True
                        as_dict = sub.get('answer_structure', {})
                        if isinstance(as_dict, dict):
                            parts = []
                            if as_dict.get('intro'):
                                parts.append(f"Mở đầu: {as_dict['intro']}")
                            if as_dict.get('body'):
                                parts.append(f"Nội dung: {as_dict['body']}")
                            if as_dict.get('conclusion'):
                                parts.append(f"Kết luận: {as_dict['conclusion']}")
                            for part in parts:
                                p = self.document.add_paragraph(style='List Bullet')
                                self._add_text_with_math(p, part)
                        sub_exp = sub.get('explanation', '')
                        if sub_exp:
                            p_exp = self.document.add_paragraph()
                            p_exp.add_run('Giải thích: ').italic = True
                            self._add_text_with_math(p_exp, str(sub_exp))
                else:
                    # Single answer_structure
                    para_exp = self.document.add_paragraph()
                    para_exp.add_run('Hướng dẫn chấm điểm:').bold = True
                    as_dict = explanation.get('answer_structure', {})
                    if isinstance(as_dict, dict):
                        parts = []
                        if as_dict.get('intro'):
                            parts.append(f"Mở đầu: {as_dict['intro']}")
                        if as_dict.get('body'):
                            parts.append(f"Nội dung: {as_dict['body']}")
                        if as_dict.get('conclusion'):
                            parts.append(f"Kết luận: {as_dict['conclusion']}")
                        for part in parts:
                            p = self.document.add_paragraph(style='List Bullet')
                            self._add_text_with_math(p, part)
            else:
                para_exp = self.document.add_paragraph()
                run_exp = para_exp.add_run('Hướng dẫn chấm điểm: ')
                run_exp.bold = True
                para_exp2 = self.document.add_paragraph()
                self._add_text_with_math(para_exp2, str(explanation))
        
        self.add_paragraph("")    
    def _add_tln_question(self, question: Dict, number: int):
        """Thêm câu hỏi trắc nghiệm luận (trả lời ngắn)"""
        try:
            # Câu hỏi
            para = self.document.add_paragraph()
            
            level = question.get('level', 'NB')
            
            run = para.add_run(f"Câu {number} ")
            run.bold = True
            run = para.add_run(f"({level})")
            run.bold = True
            self._set_run_background(run, level)
            
            para.add_run(". ")
            
            # Handle question_stem - renders text/table/chart in correct order
            question_stem = question.get('question_stem', '')
            self._render_question_stem(para, question_stem)
            
        except Exception as e:
            print(f"Error in _add_tln_question {number}: {e}")
            print(f"Question data: {question}")
            raise
    
    def _add_tln_answer(self, question: Dict, number: int):
        """Thêm đáp án câu hỏi TLN"""
        para = self.document.add_paragraph()
        para.add_run(f"Câu {number}. ").bold = True
        
        # Đáp án tô đỏ
        run_label = para.add_run("Đáp án: ")
        run_label.bold = True
        run_label.font.color.rgb = RGBColor(255, 0, 0)
        self._add_text_with_math(para, str(question.get('correct_answer', 'N/A')),
                                 bold=True, color_rgb=RGBColor(255, 0, 0))
        
        # Giải thích
        explanation = question.get('explanation', '')
        if explanation:
            para_exp = self.document.add_paragraph()
            self._add_text_with_math(para_exp, str(explanation))
        
        self.add_paragraph("")
    
    def create_from_json_file(self, 
                             json_path: str,
                             output_path: str,
                             template: Optional[Dict] = None):
        """
        Tạo DOCX từ file JSON
        
        Args:
            json_path (str): Đường dẫn file JSON
            output_path (str): Đường dẫn file DOCX output
            template (Dict, optional): Template
        """
        try:
            # Đọc JSON
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Tạo document
            self.create_new_document()
            self.set_document_style()
            self.generate_from_json(data, template)
            
            # Lưu file
            self.save(output_path)
            
            if self.verbose:
                print(f"✓ Đã tạo DOCX: {json_path} -> {output_path}")
        
        except Exception as e:
            print(f"✗ Lỗi khi tạo DOCX: {str(e)}")
            raise
