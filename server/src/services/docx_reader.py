from pathlib import Path
from typing import Dict, List, Optional, Any
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

class DocxReader:
    """Class đọc và trích xuất nội dung từ file DOCX"""
    
    def __init__(self, verbose: bool = False):
        self.document = None
        self.file_path = None
        self.verbose = verbose
    
    def load_document(self, file_path: str) -> bool:
        """
        Tải file docx
        
        Args:
            file_path (str): Đường dẫn đến file docx
            
        Returns:
            bool: True nếu tải thành công
        """
        try:
            self.file_path = Path(file_path)
            if not self.file_path.exists():
                raise FileNotFoundError(f"File không tồn tại: {file_path}")
            
            if not self.file_path.suffix.lower() == '.docx':
                raise ValueError(f"File không phải là docx: {file_path}")
            
            self.document = Document(str(self.file_path))
            
            if self.verbose:
                print(f"✓ Đã tải file: {self.file_path.name}")
                print(f"  Số đoạn văn: {len(self.document.paragraphs)}")
                print(f"  Số bảng: {len(self.document.tables)}")
            
            return True
            
        except Exception as e:
            if self.verbose:
                print(f"✗ Lỗi khi tải file: {str(e)}")
            raise
    
    def get_paragraphs(self) -> List[Dict[str, Any]]:
        """
        Lấy tất cả đoạn văn trong document
        
        Returns:
            List[Dict]: Danh sách các đoạn văn với thông tin chi tiết
        """
        if self.document is None:
            raise ValueError("Chưa tải document. Gọi load_document() trước.")
        
        paragraphs = []
        for idx, para in enumerate(self.document.paragraphs):
            para_info = {
                'index': idx,
                'text': para.text,
                'style': para.style.name if para.style else 'Normal',
                'alignment': str(para.alignment) if para.alignment else None,
                'runs': []
            }
            
            # Lấy thông tin từng run (đoạn text với format riêng)
            for run in para.runs:
                run_info = {
                    'text': run.text,
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name if run.font.name else None,
                    'font_size': run.font.size.pt if run.font.size else None,
                }
                para_info['runs'].append(run_info)
            
            paragraphs.append(para_info)
        
        if self.verbose:
            print(f"✓ Đã trích xuất {len(paragraphs)} đoạn văn")
        
        return paragraphs
    
    def get_text(self) -> str:
        """
        Lấy toàn bộ text trong document
        
        Returns:
            str: Nội dung text
        """
        if self.document is None:
            raise ValueError("Chưa tải document. Gọi load_document() trước.")
        
        text = '\n'.join([para.text for para in self.document.paragraphs])
        
        if self.verbose:
            print(f"✓ Đã trích xuất text ({len(text)} ký tự)")
        
        return text
    
    def get_tables(self) -> List[Dict[str, Any]]:
        """
        Lấy tất cả bảng trong document
        
        Returns:
            List[Dict]: Danh sách các bảng với nội dung
        """
        if self.document is None:
            raise ValueError("Chưa tải document. Gọi load_document() trước.")
        
        tables_data = []
        for table_idx, table in enumerate(self.document.tables):
            table_info = {
                'index': table_idx,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'data': []
            }
            
            # Lấy dữ liệu từng cell
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell_idx, cell in enumerate(row.cells):
                    cell_data = {
                        'row': row_idx,
                        'col': cell_idx,
                        'text': cell.text.strip()
                    }
                    row_data.append(cell_data)
                table_info['data'].append(row_data)
            
            tables_data.append(table_info)
        
        if self.verbose:
            print(f"✓ Đã trích xuất {len(tables_data)} bảng")
        
        return tables_data
    
    def get_structure(self) -> Dict[str, Any]:
        """
        Lấy cấu trúc tổng quát của document
        
        Returns:
            Dict: Thông tin cấu trúc document
        """
        if self.document is None:
            raise ValueError("Chưa tải document. Gọi load_document() trước.")
        
        # Phân tích heading structure
        headings = []
        for idx, para in enumerate(self.document.paragraphs):
            if para.style and 'Heading' in para.style.name:
                headings.append({
                    'index': idx,
                    'level': para.style.name,
                    'text': para.text
                })
        
        structure = {
            'file_name': self.file_path.name if self.file_path else None,
            'total_paragraphs': len(self.document.paragraphs),
            'total_tables': len(self.document.tables),
            'headings': headings,
            'sections': len(self.document.sections),
        }
        
        if self.verbose:
            print(f"✓ Cấu trúc document:")
            print(f"  - Đoạn văn: {structure['total_paragraphs']}")
            print(f"  - Bảng: {structure['total_tables']}")
            print(f"  - Tiêu đề: {len(headings)}")
            print(f"  - Sections: {structure['sections']}")
        
        return structure
    
    def extract_all(self) -> Dict[str, Any]:
        """
        Trích xuất toàn bộ nội dung document
        
        Returns:
            Dict: Tất cả thông tin từ document
        """
        if self.document is None:
            raise ValueError("Chưa tải document. Gọi load_document() trước.")
        
        result = {
            'structure': self.get_structure(),
            'text': self.get_text(),
            'paragraphs': self.get_paragraphs(),
            'tables': self.get_tables()
        }
        
        if self.verbose:
            print("✓ Đã trích xuất toàn bộ nội dung document")
        
        return result
    
    def search_text(self, keyword: str, case_sensitive: bool = False) -> List[Dict[str, Any]]:
        """
        Tìm kiếm text trong document
        
        Args:
            keyword (str): Từ khóa cần tìm
            case_sensitive (bool): Phân biệt hoa thường
            
        Returns:
            List[Dict]: Danh sách các đoạn văn chứa từ khóa
        """
        if self.document is None:
            raise ValueError("Chưa tải document. Gọi load_document() trước.")
        
        results = []
        search_keyword = keyword if case_sensitive else keyword.lower()
        
        for idx, para in enumerate(self.document.paragraphs):
            text = para.text if case_sensitive else para.text.lower()
            if search_keyword in text:
                results.append({
                    'index': idx,
                    'text': para.text,
                    'style': para.style.name if para.style else 'Normal'
                })
        
        if self.verbose:
            print(f"✓ Tìm thấy {len(results)} đoạn văn chứa '{keyword}'")
        
        return results


# Hàm tiện ích để sử dụng nhanh
def read_docx(file_path: str, verbose: bool = False) -> Dict[str, Any]:
    """
    Hàm tiện ích để đọc file docx
    
    Args:
        file_path (str): Đường dẫn file
        verbose (bool): Hiển thị thông tin chi tiết
        
    Returns:
        Dict: Toàn bộ nội dung document
    """
    reader = DocxReader(verbose=verbose)
    reader.load_document(file_path)
    return reader.extract_all()


def read_docx_text(file_path: str) -> str:
    """
    Hàm tiện ích để chỉ lấy text từ docx
    
    Args:
        file_path (str): Đường dẫn file
        
    Returns:
        str: Nội dung text
    """
    reader = DocxReader()
    reader.load_document(file_path)
    return reader.get_text()
