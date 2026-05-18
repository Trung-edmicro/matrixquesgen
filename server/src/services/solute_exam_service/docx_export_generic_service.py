import re
import base64
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
import latex2mathml.converter
import mathml2omml
from lxml import etree
import pypandoc
import re
import tempfile
import zipfile
import os

class DocxExportService:
    @staticmethod
    def clean_latex(latex_raw):
        if not latex_raw: return ""
        latex_raw = str(latex_raw).strip()
        latex_raw = re.sub(r'^\[\s*\$(.*?)\$\s*\]$', r'\1', latex_raw)
        if not (latex_raw.startswith('$') and latex_raw.endswith('$')):
            latex_raw = f"${latex_raw}$"
        return latex_raw


    # def latex_to_omml(latex_math):
    #     try:
    #         raw = latex_math[1:-1] if latex_math.startswith('$') else latex_math
    #         mathml = latex2mathml.converter.convert(raw)
    #         omml = mathml2omml.convert(mathml)
    #         return omml
    #     except: return None
    def latex_to_omml(self, latex_math):

        try:

            raw = latex_math.strip()

            if raw.startswith("$$") and raw.endswith("$$"):
                raw = raw[2:-2]

            elif raw.startswith("$") and raw.endswith("$"):
                raw = raw[1:-1]

            raw = raw.strip()

            if not raw:
                return None

            # tạo latex content
            latex_content = f"""
            \\[
            {raw}
            \\]
            """

            # file temp
            with tempfile.TemporaryDirectory() as tmpdir:

                tex_file = os.path.join(tmpdir, "math.tex")
                docx_file = os.path.join(tmpdir, "math.docx")

                # write latex
                with open(tex_file, "w", encoding="utf-8") as f:
                    f.write(latex_content)

                # convert bằng pandoc
                pypandoc.convert_file(
                    tex_file,
                    to="docx",
                    format="latex",
                    outputfile=docx_file
                )

                # unzip docx
                with zipfile.ZipFile(docx_file, "r") as zip_ref:

                    xml_content = zip_ref.read(
                        "word/document.xml"
                    ).decode("utf-8")

                # extract omml
                match = re.search(
                    r'(<m:oMathPara.*?</m:oMathPara>|<m:oMath.*?</m:oMath>)',
                    xml_content,
                    flags=re.DOTALL
                )

                if not match:
                    print("NO OMML FOUND")
                    print(xml_content)
                    return None

                omml = match.group(1)

                # normalize về inline math
                if "<m:oMathPara" in omml:

                    inner = re.search(
                        r'<m:oMath>(.*?)</m:oMath>',
                        omml,
                        flags=re.DOTALL
                    )

                    if inner:
                        omml = f"<m:oMath>{inner.group(1)}</m:oMath>"

                return omml

        except Exception as e:

            print("===================")
            print("LATEX:", latex_math)
            print("ERROR:", e)

            return None

    # def insert_text_with_math(self, paragraph, text, bold=False, italic=False, underline=False):
    #     """Xử lý trộn lẫn văn bản, thẻ HTML (b, i, u) và $latex$"""
    #     if not text: return
        
    #     # Tách tags HTML và Latex
    #     # parts = re.split(r'(<b>|</b>|<i>|</i>|<u>|</u>|\$[^$]*\$)', str(text))
    #     parts = re.split(
    #                 r'(<b>|</b>|<i>|</i>|<u>|</u>|\$\$.*?\$\$|\$.*?\$)',
    #                 str(text),
    #                 flags=re.DOTALL
    #             )
        
    #     curr_bold = bold
    #     curr_italic = italic
    #     curr_underline = underline

    #     for part in parts:
    #         if not part: continue
    #         if part == "<b>": curr_bold = True
    #         elif part == "</b>": curr_bold = bold
    #         elif part == "<i>": curr_italic = True
    #         elif part == "</i>": curr_italic = italic
    #         elif part == "<u>": curr_underline = True
    #         elif part == "</u>": curr_underline = underline
    #         elif part.startswith('$') and part.endswith('$'):
    #             omml = self.latex_to_omml(self.clean_latex(part))
    #             if omml:
    #                 try:
    #                     # omml_elm = parse_xml(omml)
    #                     # run = paragraph.add_run()
    #                     # run._r.append(omml_elm)
    #                     omml_elm = etree.fromstring(
    #                         omml.encode("utf-8")
    #                     )

    #                     run = paragraph.add_run()
    #                     run._r.append(omml_elm)
    #                     continue
    #                 except: pass
    #             paragraph.add_run(part)
    #         else:
    #             run = paragraph.add_run(part)
    #             run.bold = curr_bold
    #             run.italic = curr_italic
    #             run.underline = curr_underline

    def insert_text_with_math(
        self,
        paragraph,
        text,
        bold=False,
        italic=False,
        underline=False
    ):

        if not text:
            return

        parts = re.split(
            r'(<b>|</b>|<i>|</i>|<u>|</u>|\$\$.*?\$\$|\$.*?\$)',
            str(text),
            flags=re.DOTALL
        )

        curr_bold = bold
        curr_italic = italic
        curr_underline = underline

        for part in parts:

            if not part:
                continue

            if part == "<b>":
                curr_bold = True

            elif part == "</b>":
                curr_bold = bold

            elif part == "<i>":
                curr_italic = True

            elif part == "</i>":
                curr_italic = italic

            elif part == "<u>":
                curr_underline = True

            elif part == "</u>":
                curr_underline = underline

            elif (
                (part.startswith("$$") and part.endswith("$$"))
                or
                (part.startswith("$") and part.endswith("$"))
            ):

                omml = self.latex_to_omml(part)

                if omml:
                    try:

                        # omml_elm = etree.fromstring(
                        #     omml.encode("utf-8")
                        # )

                        # run = paragraph.add_run()
                        # run._r.append(omml_elm)
                        xml = f"""
                        <w:r
                            xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                            xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">

                            {omml}

                        </w:r>
                        """

                        root = etree.fromstring(
                            xml.encode("utf-8")
                        )

                        run = paragraph.add_run()

                        # append toàn bộ child
                        for child in root:
                            run._r.append(child)

                        continue

                    except Exception as e:
                        print("OMML ERROR:", e)

                paragraph.add_run(part)

            else:
                run = paragraph.add_run(part)
                run.bold = curr_bold
                run.italic = curr_italic
                run.underline = curr_underline

    def _add_table_to_doc(self, doc, table_data):
        """Logic xử lý bảng đầy đủ: Tiêu đề, đơn vị, dữ liệu, ghi chú"""
        if not table_data: return

        # 1. Tiêu đề bảng
        if table_data.get("title"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.insert_text_with_math(p, table_data["title"], bold=True)

        # 2. Đơn vị
        unit = table_data.get("unit")
        if unit:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            unit_text = unit if unit.startswith("(") else f"(Đơn vị: {unit})"
            self.insert_text_with_math(p, unit_text, italic=True)

        # 3. Tạo bảng
        headers = table_data.get("headers", [])
        rows = table_data.get("rows", [])
        num_cols = len(headers) if headers else (len(rows[0]) if rows else 0)
        if num_cols == 0: return

        table = doc.add_table(rows=0, cols=num_cols)
        table.style = 'Table Grid'

        # Thêm Header
        if headers:
            hdr_cells = table.add_row().cells
            for i, h in enumerate(headers):
                self.insert_text_with_math(hdr_cells[i].paragraphs[0], h, bold=True)

        # Thêm Dữ liệu
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, val in enumerate(row_data):
                if i < num_cols:
                    self.insert_text_with_math(row_cells[i].paragraphs[0], val)

        # 4. Ghi chú dưới bảng
        if table_data.get("notes"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self.insert_text_with_math(p, f"{table_data['notes']}", italic=True)

    def _add_image_to_doc(self, doc, media_data):
        """Xử lý hình ảnh từ base64"""
        img_b64 = media_data.get("image_base64") or media_data.get("source")
        if not img_b64: return
        try:
            if "base64," in img_b64: img_b64 = img_b64.split("base64,")[1]
            image_bytes = base64.b64decode(img_b64)
            doc.add_picture(BytesIO(image_bytes), width=Inches(4))
            if media_data.get("image_caption"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self.insert_text_with_math(p, media_data["image_caption"], italic=True)
        except: pass

    # def create_standard_docx(self, data, file_path):
    #     doc = Document()
    #     # Header Đề thi
    #     title_p = doc.add_paragraph()
    #     title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #     self.insert_text_with_math(title_p, data.get("exam_title", "ĐỀ THI"), bold=True)

    #     for section in data.get("sections", []):
    #         # 1. Section Title
    #         if section.get("section_title"):
    #             doc.add_paragraph() # Khoảng trống
    #             sec_p = doc.add_paragraph()
    #             self.insert_text_with_math(sec_p, section["section_title"].upper(), bold=True)

    #         for q in section.get("questions", []):
    #             media_list = q.get("media") or []
                
    #             # Render Media "before_question_content"
    #             for m in media_list:
    #                 if m.get("position") == "before_question_content":
    #                     if m["type"] == "table": self._add_table_to_doc(doc, m)
    #                     elif m["type"] == "image": self._add_image_to_doc(doc, m)

    #             # 2. Câu hỏi
    #             p_q = doc.add_paragraph()
    #             p_q.add_run(f"Câu {q.get('question_number')}. ").bold = True
    #             if q.get("question_title"):
    #                 self.insert_text_with_math(p_q, f"{q['question_title']}: ", bold=True)
    #             self.insert_text_with_math(p_q, q.get("question_content", ""))

    #             # Render Media "after_question_content"
    #             for m in media_list:
    #                 if m.get("position") != "before_question_content":
    #                     if m["type"] == "table": self._add_table_to_doc(doc, m)
    #                     elif m["type"] == "image": self._add_image_to_doc(doc, m)

    #             # 3. Handle Types
    #             q_type = q.get("type")
    #             if q_type == "multiple_choice": self._handle_multiple_choice(doc, q)
    #             elif q_type == "true_false": self._handle_true_false(doc, q)
    #             elif q_type == "essay": self._handle_essay(doc, q)

    #             # 4. Note
    #             if q.get("note"):
    #                 p_n = doc.add_paragraph()
    #                 self.insert_text_with_math(p_n, f"Lưu ý: {q['note']}", italic=True)

    #             # 5. Lời giải
    #             self._add_explanation_section(doc, q)

    #     doc.save(file_path)
    def create_standard_docx(self, data, file_path):
        doc = Document()
        if not data:
            return
            
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.insert_text_with_math(title_p, data.get("exam_title") or "ĐỀ THI", bold=True)

        # Sửa lỗi: (data.get("sections") or []) để tránh None
        for section in (data.get("sections") or []):
            if section.get("section_title"):
                doc.add_paragraph() 
                sec_p = doc.add_paragraph()
                self.insert_text_with_math(sec_p, section["section_title"].upper(), bold=True)

            # Sửa lỗi: (section.get("questions") or [])
            for q in (section.get("questions") or []):
                # Xử lý Media an toàn
                media_list = (q.get("media") or [])
                
                # Nội dung câu hỏi
                p_q = doc.add_paragraph()
                p_q.add_run(f"Câu {q.get('question_number') or ''}. ").bold = True
                if q.get("question_title"):
                        self.insert_text_with_math(
                            p_q,
                            f"{q['question_title']} ",
                            bold=True
                        )
                # self.insert_text_with_math(p_q, q.get("question_content") or "")

                # for m in media_list:
                #         if m["type"] == "table": self._add_table_to_doc(doc, m)
                #         elif m["type"] == "image": self._add_image_to_doc(doc, m)
                for m in media_list:
                    if not m:
                        continue

                    if m.get("position") == "before_question_content":
                        if m["type"] == "table":
                            self._add_table_to_doc(doc, m)

                        elif m["type"] == "image":
                            self._add_image_to_doc(doc, m)
                
                question_content = q.get("question_content")

                # if question_content:
                #     p_content = doc.add_paragraph()

                #     self.insert_text_with_math(
                #         p_content,
                #         question_content
                #     )

                has_after_media = any(
                    m and m.get("position") == "after_question_content"
                    for m in media_list
                )

                if question_content:
                    # Nếu có after_question_content
                    # -> gộp content vào cùng paragraph với "Câu X"
                    if has_after_media:
                        self.insert_text_with_math(
                            p_q,
                            question_content
                        )

                    # Logic cũ giữ nguyên
                    else:
                        p_content = doc.add_paragraph()

                        self.insert_text_with_math(
                            p_content,
                            question_content
                        )

                # AFTER -> thêm mới
                for m in media_list:
                    if not m:
                        continue

                    if m.get("position") == "after_question_content":
                        if m["type"] == "table":
                            self._add_table_to_doc(doc, m)

                        elif m["type"] == "image":
                            self._add_image_to_doc(doc, m)

                # Xử lý các loại câu hỏi
                q_type = q.get("type")
                if q_type == "multiple_choice": self._handle_multiple_choice(doc, q)
                elif q_type == "true_false": self._handle_true_false(doc, q)
                elif q_type == "essay": self._handle_essay(doc, q)

                # Ghi chú
                if q.get("note"):
                    p_n = doc.add_paragraph()
                    self.insert_text_with_math(p_n, f"Lưu ý: {q['note']}")

                # Lời giải
                self._add_explanation_section(doc, q)

        doc.save(file_path)

    def _handle_multiple_choice(self, doc, q):
        if not q.get("options"): return
        opt = q["options"][0]
        for label, key in [("A", "option_a"), ("B", "option_b"), ("C", "option_c"), ("D", "option_d")]:
            if opt.get(key):
                p = doc.add_paragraph('')
                p.add_run(f"{label}. ").bold = True
                self.insert_text_with_math(p, opt[key])

    def _handle_true_false(self, doc, q):
        for opt in q.get("options", []):
            p = doc.add_paragraph()
            lbl = opt.get("label", "").lower()
            if lbl: self.insert_text_with_math(p, f"{lbl}) ")
            self.insert_text_with_math(p, opt.get("content", ""))

    # def _handle_essay(self, doc, q):
    #     for pas in q.get("passage_data", []):
    #         if pas.get("passage_title"):
    #             p = doc.add_paragraph()
    #             self.insert_text_with_math(p, pas["passage_title"], bold=True)
    #         if pas.get("passage_content"):
    #             self.insert_text_with_math(doc.add_paragraph(), pas["passage_content"])
    def _handle_essay(self, doc, q):
        # Sửa lỗi None ở passage_data
        for pas in (q.get("passage_data") or []):
            if pas.get("passage_title"):
                p = doc.add_paragraph()
                self.insert_text_with_math(p, pas["passage_title"], bold=True)
            if pas.get("passage_content"):
                self.insert_text_with_math(doc.add_paragraph(), pas["passage_content"])
        
        # Sửa lỗi None ở images

    def _add_explanation_section(self, doc, q):
        doc.add_paragraph().add_run("Lời giải").bold = True
        
        # Đáp án ngắn gọn
        if q.get("type") == "multiple_choice":
            ans = q["options"][0].get("answer") if q.get("options") else ""
            if ans: doc.add_paragraph(f"Chọn {ans}")
        elif q.get("type") == "true_false":
            binary = "".join(["1" if o.get("is_correct") else "0" for o in q.get("options", [])])
            if binary: doc.add_paragraph(binary)
        elif q.get("type") == "short_answer":
            p = doc.add_paragraph("Đáp án: ")
            self.insert_text_with_math(p, f"[[{q.get('correct_answer', '')}]]", bold=True)

        doc.add_paragraph("####").runs[0].bold = True
        if q.get("type") == "true_false":

            for opt in (q.get("options") or []):

                # Dòng nội dung + ĐÚNG/SAI
                p = doc.add_paragraph()

                label = opt.get("label", "")
                content = opt.get("content", "")

                self.insert_text_with_math(p, content)

                p.add_run(" → ")

                result_run = p.add_run(
                    "ĐÚNG" if opt.get("is_correct") else "SAI"
                )
                result_run.bold = True

                # Explanation
                explanation = opt.get("explanation")

                if explanation:
                    for line in str(explanation).split("\n"):
                        self.insert_text_with_math(
                            doc.add_paragraph(),
                            line
                        )

            return

        if q.get("explanation"):
            for line in str(q["explanation"]).split("\n"):
                self.insert_text_with_math(doc.add_paragraph(), line)
        
        if q.get("conclusion"):
            p = doc.add_paragraph()
            self.insert_text_with_math(p, q["conclusion"], bold=True)

    def create_literature_docx(self, data, file_path):
        """Xử lý Ngữ Văn: Đoạn trích, nguồn, và cấu trúc Writing 200/600"""
        doc = Document()
        schema = data.get("exam_data_schema", data)
        
        head_p = doc.add_paragraph()
        head_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.insert_text_with_math(head_p, schema.get("exam_title", "ĐỀ THI NGỮ VĂN"), bold=True)

        for sec in schema.get("sections", []):
            if sec.get("section_title"):
                doc.add_paragraph()
                self.insert_text_with_math(doc.add_paragraph(), sec["section_title"].upper(), bold=True)
            
            passage = sec.get("reading_passage", {})
            if passage:
                if passage.get("intro_text"):
                    p = doc.add_paragraph()
                    self.insert_text_with_math(p, passage["intro_text"], bold=True)
                if passage.get("content"):
                    for line in passage["content"].split("\n"):
                        self.insert_text_with_math(doc.add_paragraph(), line)
                if passage.get("source"):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    self.insert_text_with_math(p, f"({passage['source']})")

            for q in sec.get("questions", []):
                p_q = doc.add_paragraph()
                p_q.add_run(f"Câu {q.get('number')}. ").bold = True
                self.insert_text_with_math(p_q, q.get("question_content", ""))

                sol = q.get("solution", {})
                doc.add_paragraph("Lời giải:").runs[0].bold = True
                
                if "WRITING" in str(q.get("question_type", "")):
                    self._handle_lit_writing_structure(doc, sol)
                else:
                    if sol.get("explanation"):
                        for line in sol["explanation"].split("\\n"):
                            self.insert_text_with_math(doc.add_paragraph(), line)
        doc.save(file_path)

    def _handle_lit_writing_structure(self, doc, solution):
        struct = solution.get("structured_content", {})
        
        # a. Yêu cầu chung
        gen = struct.get("a_general_requirements", {})
        if gen:
            doc.add_paragraph().add_run("a. Yêu cầu chung:").bold = True
            for k, key in [("Vấn đề", "issue"), ("Hình thức", "form"), ("Dung lượng", "length"), ("Bằng chứng", "evidence")]:
                if gen.get(key):
                    p = doc.add_paragraph()
                    self.insert_text_with_math(p, f"- {k}: {gen[key]}")

        # b. Yêu cầu cụ thể
        spec = struct.get("b_specific_requirements", {})
        if spec:
            doc.add_paragraph().add_run("b. Yêu cầu cụ thể:").bold = True
            steps = spec.get("steps", {})
            for skey in sorted(steps.keys()):
                sdata = steps[skey]
                if not sdata: continue
                p = doc.add_paragraph()
                p.add_run(f"{skey}. {sdata.get('name', '')}: ").italic = True
                if sdata.get("content"):
                    for line in sdata["content"].split("\n"):
                        p_line = doc.add_paragraph()
                        self.insert_text_with_math(p_line, line)