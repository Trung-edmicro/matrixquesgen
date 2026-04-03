from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from bs4 import BeautifulSoup

def export_soluted_english_exam_from_data(data: dict, output_path: str):
    # Lấy dữ liệu từ results hoặc data
    results = data.get("results") or data.get("data", [])
    if isinstance(results, list) and len(results) > 0 and isinstance(results[0], list):
        results = results[0]

    if not results:
        raise Exception("Không tìm thấy dữ liệu để xuất file")

    doc = Document()
    _apply_default_style(doc)

    i = 0
    n = len(results)

    while i < n:
        # Gom nhóm các câu có cùng type và titleQuestion
        group, next_i = _collect_same_group(results, i)
        res_type = group[0].get("type")
        
        # 1. Thêm tiêu đề nhóm (titleQuestion) - render BOLD
        title_text = group[0].get("titleQuestion")
        if title_text:
            _add_group_title(doc, title_text)

        # 2. Điều hướng render theo loại block
        if res_type in ("RC", "CLOZE", "GAP"):
            _render_passage_based_group(doc, group)
        elif res_type == "ARRANGE":
            _render_arrange_group(doc, group)
        elif res_type == "PRONUNCIATION_STRESS":
            _render_pronunciation_group(doc, group)
        elif res_type == "ERROR_IDENTIFICATION":
            _render_error_group(doc, group)
        elif res_type == "SYNONYM_ANTONYM":
            _render_synonym_antonym_group(doc, group)
        elif res_type == "SENTENCE_COMPLETION":
            _render_sentence_completion_group(doc, group)
        elif res_type == "SENTENCE_TRANSFORMATION":
            _render_transformation_group(doc, group)
        elif res_type == "DIALOGUE_COMPLETION":
            _render_dialogue_group(doc, group)
        elif res_type == "LOGICAL_THINKING":
            _render_logical_group(doc, group)
        elif res_type == "WORD_REORDERING":
            _render_reordering_group(doc, group)
        else:
            _render_simple_question_group(doc, group)

        i = next_i

    doc.save(output_path)
    return output_path

def export_soluted_standard_english_exam_from_data(data: dict, output_path: str):
    # Lấy dữ liệu từ results hoặc data
    results = data.get("results") or data.get("data", [])
    if isinstance(results, list) and len(results) > 0 and isinstance(results[0], list):
        results = results[0]

    if not results:
        raise Exception("Không tìm thấy dữ liệu để xuất file")

    doc = Document()
    _apply_default_style(doc)

    i = 0
    n = len(results)

    while i < n:
        # Gom nhóm các câu có cùng type và titleQuestion
        group, next_i = _collect_same_group(results, i)
        res_type = group[0].get("type")
        
        # 1. Thêm tiêu đề nhóm (titleQuestion) - render BOLD
        title_text = group[0].get("titleQuestion")
        if title_text:
            _add_group_title(doc, title_text)

        # 2. Điều hướng render theo loại block
        if res_type in ("RC", "CLOZE", "GAP"):
            _render_standard_passage_based_group(doc, group)
        elif res_type == "ARRANGE":
            _render_standard_arrange_group(doc, group)
        elif res_type == "PRONUNCIATION_STRESS":
            _render_standard_pronunciation_group(doc, group)
        elif res_type == "ERROR_IDENTIFICATION":
            _render_standard_error_group(doc, group)
        elif res_type == "SYNONYM_ANTONYM":
            _render_standard_synonym_antonym_group(doc, group)
        elif res_type == "SENTENCE_COMPLETION":
            _render_standard_sentence_completion_group(doc, group)
        elif res_type == "SENTENCE_TRANSFORMATION":
            _render_standard_transformation_group(doc, group)
        elif res_type == "DIALOGUE_COMPLETION":
            _render_standard_dialogue_group(doc, group)
        elif res_type == "LOGICAL_THINKING":
            _render_standard_logical_group(doc, group)
        elif res_type == "WORD_REORDERING":
            _render_standard_reordering_group(doc, group)
        else:
            _render_standard_simple_question_group(doc, group)

        i = next_i

    doc.save(output_path)
    return output_path



# ============================
# HELPERS (STYLE, GROUPING, HTML)
# ============================

def _apply_default_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

def _collect_same_group(results, start_index):
    first = results[start_index]
    group = []
    i = start_index
    while i < len(results):
        res = results[i]
        if (res.get("type") == first.get("type") and 
            res.get("titleQuestion") == first.get("titleQuestion")):
            group.append(res)
            i += 1
        else:
            break
    return group, i

def _add_group_title(doc, title):
    p = doc.add_paragraph()
    run = p.add_run(str(title))
    run.bold = True

def _render_html_text(paragraph, html_content):
    """Xử lý render HTML (b, u, strong, i) cho nội dung"""
    if not html_content: return
    # Làm sạch chuỗi
    html_content = str(html_content).replace("&nbsp;", " ").replace("\r\n", "\n")
    soup = BeautifulSoup(html_content, "html.parser")
    
    for element in soup.descendants:
        if element.name is None and element.string:
            run = paragraph.add_run(element.string)
            parents = [t.name for t in element.parents]
            if 'b' in parents or 'strong' in parents: run.bold = True
            if 'u' in parents: run.underline = True
            if 'i' in parents: run.italic = True

def _render_solution_block(doc, q, res_type=None):
    p_ans = doc.add_paragraph()
    p_ans.add_run("Lời giải").bold = True
    
    doc.add_paragraph(f"Chọn {q.get('answer', '')}")
    
    if q.get("explanation"):
        p_exp = doc.add_paragraph()
        
        explanation = str(q["explanation"])
        
        # 1. Convert markdown **bold** → HTML <b>
        explanation = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", explanation)
        
        # 2. Render bằng HTML engine
        _render_html_text(p_exp, explanation)

    # if q.get("quote"):
    #     p_quote = doc.add_paragraph()
    #     p_quote.add_run("Trích bài: ").bold = True
        
    #     quote_html = str(q["quote"])
    #     _render_html_text(p_quote, quote_html) 

    if q.get("quote"):
        p_quote = doc.add_paragraph()

        label = "Thông tin: " if res_type == "RC" else "Trích bài: "
        p_quote.add_run(label).bold = True
        
        _render_html_text(p_quote, str(q["quote"]))

                
    if q.get("translation"):
        p_trans = doc.add_paragraph()
        p_trans.add_run("Tạm dịch: ").bold = True
        p_trans.add_run(str(q['translation']))


def _render_solution__reading_comprehensive_block(doc, q):
    p_ans = doc.add_paragraph()
    p_ans.add_run("Lời giải").bold = True
    
    doc.add_paragraph(f"Chọn {q.get('answer', '')}")
    
    if q.get("explanation"):
        p_exp = doc.add_paragraph()
        
        explanation = str(q["explanation"])
        
        # 1. Convert markdown **bold** → HTML <b>
        explanation = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", explanation)
        
        # 2. Render bằng HTML engine
        _render_html_text(p_exp, explanation)

    if q.get("quote"):
        p_quote = doc.add_paragraph()
        p_quote.add_run("Thông tin: ").bold = True
        
        quote_html = str(q["quote"])
        _render_html_text(p_quote, quote_html) 
                
    if q.get("translation"):
        p_trans = doc.add_paragraph()
        p_trans.add_run("Tạm dịch: ").bold = True
        p_trans.add_run(str(q['translation']))

# ============================
# CÁC HÀM RENDER CHI TIẾT
# ============================

def _word_count(text):
    if not text:
        return 0
    return len(str(text).strip().split())


def _render_options(doc, q):
    options = {
        "A": q.get("option_a", ""),
        "B": q.get("option_b", ""),
        "C": q.get("option_c", ""),
        "D": q.get("option_d", "")
    }

    # Đếm số từ từng đáp án
    word_counts = {k: _word_count(v) for k, v in options.items()}

    # Kiểm tra điều kiện
    short_options = all(count < 2 for count in word_counts.values())
    long_options = all(count >= 3 for count in word_counts.values())

    # =========================
    # CASE 1: Tất cả ngắn → 1 dòng
    # =========================
    if short_options:
        p = doc.add_paragraph()
        for key, value in options.items():
            run_label = p.add_run(f"{key}. ")
            run_label.bold = True
            p.add_run(f"{value}    ")  # spacing đẹp hơn tab

    # =========================
    # CASE 2: Tất cả dài → mỗi dòng 1 đáp án
    # =========================
    elif long_options:
        for key, value in options.items():
            p = doc.add_paragraph()
            run_label = p.add_run(f"{key}. ")
            run_label.bold = True
            _render_html_text(p, value)

    # =========================
    # CASE 3: Mixed → xử lý linh hoạt (best UX)
    # =========================
    else:
        for key, value in options.items():
            if _word_count(value) >= 3:
                p = doc.add_paragraph()
            else:
                # gộp các option ngắn lại
                if 'p_inline' not in locals():
                    p_inline = doc.add_paragraph()
                p = p_inline

            run_label = p.add_run(f"{key}. ")
            run_label.bold = True
            p.add_run(f"{value}  ")

def _preprocess_question_content(text):
    if not text:
        return text
    
    text = str(text).strip()
    
    # 1. Replace “ ” → "
    text = text.replace("“", '"').replace("”", '"')
    
    # 2. Check nếu toàn bộ string nằm trong dấu "
    if text.startswith('"') and text.endswith('"'):
        inner = text[1:-1].strip()
        
        # Wrap bằng <b> để tận dụng HTML renderer
        return f'<b><u>{inner}</u></b>'
    
    return text

def _render_synonym_antonym_group(doc, group):
    for res in group:
        q_list = res.get("parsed", {}).get("questions", [])
        for q in q_list:
            p = doc.add_paragraph()

            # Question number
            p.add_run(f"Question {q.get('number')}: ").bold = True

            # Nếu có type thì thêm (Synonym / Antonym)
            # if q.get("type"):
            #     p.add_run(f"({q.get('type').capitalize()}) ").italic = True

            # Nội dung câu hỏi
            _render_html_text(p, q.get("question"))

            # Options
            _render_options(doc, q)

            # Solution
            _render_solution_block(doc, q)

def _render_sentence_completion_group(doc, group):
    for res in group:
        q_list = res.get("parsed", {}).get("questions", [])
        for q in q_list:
            p = doc.add_paragraph()

            p.add_run(f"Question {q.get('number')}: ").bold = True

            if q.get("question"):
                p.add_run(q["question"])
            # Options
            _render_options(doc, q)

            # Solution
            _render_solution_block(doc, q)

def _render_transformation_group(doc, group):
    for res in group:
        q_list = res.get("parsed", {}).get("questions", [])
        for q in q_list:
            p = doc.add_paragraph()

            p.add_run(f"Question {q.get('number')}: ").bold = True

            # Loại bài (rewriting / combination)
            if q.get("question"):
                p.add_run(q["question"])

            # Options
            _render_options(doc, q)

            # Solution
            _render_solution_block(doc, q)

            # Dịch câu đúng
            if q.get("correct_translation"):
                p_corr = doc.add_paragraph()
                p_corr.add_run("Dịch câu đúng: ").bold = True
                p_corr.add_run(str(q["correct_translation"]))




def _render_passage_based_group(doc, group):
    """Điền từ, Đọc hiểu, Điền cụm từ - Sử dụng HTML render cho PASSAGE"""
    first_res = group[0]
    parsed_main = first_res.get("parsed", {})
    res_type = group[0].get("type")
    
    # 1. Render Passage Title (Dùng HTML render)
    if parsed_main.get("passage_title"):
        p_t = doc.add_paragraph()
        p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        _render_html_text(p_t, parsed_main["passage_title"])
        
        # Ép tất cả run trong paragraph thành bold
        for run in p_t.runs:
            run.bold = True
    
    # 2. Render Passage (Sử dụng HTML render cho từng đoạn văn)
    if parsed_main.get("passage"):
        passage_content = str(parsed_main["passage"])
        # Chia nhỏ theo dòng để giữ paragraph
        lines = passage_content.split('\n')
        for line in lines:
            if line.strip():
                p_p = doc.add_paragraph()
                p_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _render_html_text(p_p, line.strip())

    # 3. Render Questions
    for res in group:
        q_list = res.get("parsed", {}).get("questions", [])
        for q in q_list:
            p = doc.add_paragraph()
            # p.add_run(f"Question {q.get('number')}: ").bold = True
            # # _render_html_text(p, q.get("question_content"))
            # processed_content = _preprocess_question_content(q.get("question_content"))
            # _render_html_text(p, processed_content)
            content = (q.get("question_content") or "").strip()
            number = q.get("number")

            expected = f"Question {number}:"

            # Normalize nhẹ (tránh lệch do khoảng trắng)
            if content.lower() == expected.lower():
                # Nếu giống hoàn toàn → chỉ render 1 lần và bold
                p.add_run(expected).bold = True
            else:
                # Bình thường
                p.add_run(expected + " ").bold = True
                processed_content = _preprocess_question_content(content)
                _render_html_text(p, processed_content)
            
            # doc.add_paragraph(f"A. {q.get('option_a')}\tB. {q.get('option_b')}\tC. {q.get('option_c')}\tD. {q.get('option_d')}")
            _render_options(doc, q)
            _render_solution_block(doc, q, res_type)

def _render_simple_question_group(doc, group):
    for res in group:
        q_list = res.get("parsed", {}).get("questions", [])
        for q in q_list:
            p = doc.add_paragraph()
            
            p.add_run(f"Question {q.get('number')}: ").bold = True
            _render_html_text(p, q.get("question") or q.get("question_content"))
            doc.add_paragraph(f"A. {q.get('option_a')}\tB. {q.get('option_b')}\tC. {q.get('option_c')}\tD. {q.get('option_d')}")
            _render_solution_block(doc, q)

def _render_pronunciation_group(doc, group):
    for res in group:
        q_list = res.get("parsed", {}).get("questions", [])
        for q in q_list:
            p = doc.add_paragraph()
           
            p.add_run(f"Question {q.get('number')}: ").bold = True
            for label in ['a', 'b', 'c', 'd']:
                p.add_run(f"{label.upper()}. ")
                _render_html_text(p, q.get(f"option_{label}"))
                p.add_run("\t")
            _render_solution_block(doc, q)
            if q.get("details"):
                for d in q["details"]:
                    doc.add_paragraph(f"- {d.get('word')}: {d.get('ipa')} ({d.get('pos')}): {d.get('meaning')}")

def _render_error_group(doc, group):
    for res in group:
        q_list = res.get("parsed", {}).get("questions", [])
        for q in q_list:
            p = doc.add_paragraph()
            p.add_run(f"Question {q.get('number')}: ").bold = True
            _render_html_text(p, q.get("question"))
            doc.add_paragraph(f"A. {q.get('option_a')}\tB. {q.get('option_b')}\tC. {q.get('option_c')}\tD. {q.get('option_d')}")
            _render_solution_block(doc, q)
            if q.get("correction"):
                p_corr = doc.add_paragraph()
                p_corr.add_run("Sửa lại: ").bold = True
                p_corr.add_run(str(q['correction']))

def _render_dialogue_group(doc, group):
    for res in group:
        q_list = res.get("parsed", {}).get("questions", [])
        for q in q_list:
            p = doc.add_paragraph()
            p.add_run(f"Question {q.get('number')}: ").bold = True

            # Speaker A
            if q.get("speaker_a"):
                p = doc.add_paragraph()
                p.add_run("A: ").bold = True
                p.add_run(q["speaker_a"])

            # Speaker B
            if q.get("speaker_b"):
                p = doc.add_paragraph()
                p.add_run("B: ").bold = True
                p.add_run(q["speaker_b"])

            # ✅ Options riêng từng dòng
            if q.get("option_a"):
                doc.add_paragraph(f"A. {q.get('option_a')}")
            if q.get("option_b"):
                doc.add_paragraph(f"B. {q.get('option_b')}")
            if q.get("option_c"):
                doc.add_paragraph(f"C. {q.get('option_c')}")
            if q.get("option_d"):
                doc.add_paragraph(f"D. {q.get('option_d')}")

            _render_solution_block(doc, q)

def _render_logical_group(doc, group):
    for res in group:
        q_list = res.get("parsed", {}).get("questions", [])
        for q in q_list:
            p = doc.add_paragraph()
            p.add_run(f"Question {q.get('number')}: ").bold = True

            # Scenario (in nghiêng)
            if q.get("scenario"):
                doc.add_paragraph(q["scenario"]).italic = True

            # Question
            if q.get("question"):
                doc.add_paragraph(q["question"])

            # ✅ Options riêng từng dòng
            if q.get("option_a"):
                doc.add_paragraph(f"A. {q.get('option_a')}")
            if q.get("option_b"):
                doc.add_paragraph(f"B. {q.get('option_b')}")
            if q.get("option_c"):
                doc.add_paragraph(f"C. {q.get('option_c')}")
            if q.get("option_d"):
                doc.add_paragraph(f"D. {q.get('option_d')}")

            _render_solution_block(doc, q)

def _render_reordering_group(doc, group):
    for res in group:
        q_list = res.get("parsed", {}).get("questions", [])
        for q in q_list:
            p = doc.add_paragraph()
           
            p.add_run(f"Question {q.get('number')}: ").bold = True
            p.add_run(str(q.get("word_list", "")))
            doc.add_paragraph(f"A. {q.get('option_a')}")
            doc.add_paragraph(f"B. {q.get('option_b')}")
            doc.add_paragraph(f"C. {q.get('option_c')}")
            doc.add_paragraph(f"D. {q.get('option_d')}")
            _render_solution_block(doc, q)

def _add_options_line(doc, parsed):
    p = doc.add_paragraph()

    options = [
        ("A", parsed.get("option_a")),
        ("B", parsed.get("option_b")),
        ("C", parsed.get("option_c")),
        ("D", parsed.get("option_d")),
    ]

    for idx, (label, value) in enumerate(options):
        run = p.add_run(f"{label}. ")
        run.bold = True

        p.add_run(str(value) if value else "")

        if idx < len(options) - 1:
            p.add_run("\t")

def _render_arrange_group(doc, group):
    for res in group:
        parsed = res.get("parsed", {})

        p = doc.add_paragraph()
        p.add_run(f"Question {parsed.get('question_number') or parsed.get('number')}: ").bold = True

        if parsed.get("question_content"):
            for line in parsed["question_content"]:
                doc.add_paragraph(str(line))

        # ✅ FIX HERE
        _add_options_line(doc, parsed)

        p_ans = doc.add_paragraph()
        p_ans.add_run("Lời giải").bold = True

        doc.add_paragraph(f"Chọn {parsed.get('answer')}")

        if parsed.get("solution_lines"):
            for line in parsed["solution_lines"]:
                doc.add_paragraph(str(line))

        if parsed.get("translation_lines"):
            p_t_h = doc.add_paragraph()
            p_t_h.add_run("Tạm dịch:").bold = True

            for line in parsed["translation_lines"]:
                doc.add_paragraph(str(line))


def _render_standard_synonym_antonym_group(doc, group):
    for res in group:
        q_list = res.get("parsed", {}).get("questions", [])
        for q in q_list:
            p = doc.add_paragraph()

            # Question number
            p.add_run(f"Question {q.get('number')}: ").bold = True

            # Nội dung câu hỏi
            _render_html_text(p, q.get("question"))

            # Options
            _render_options(doc, q)

            # Solution
            # _render_solution_block(doc, q)

def _render_standard_sentence_completion_group(doc, group):
    for res in group:
        q_list = res.get("parsed", {}).get("questions", [])
        for q in q_list:
            p = doc.add_paragraph()

            p.add_run(f"Question {q.get('number')}: ").bold = True

            if q.get("question"):
                p.add_run(q["question"])
            # Options
            _render_options(doc, q)

            # Solution
            # _render_solution_block(doc, q)

def _render_standard_transformation_group(doc, group):
    for res in group:
        q_list = res.get("parsed", {}).get("questions", [])
        for q in q_list:
            p = doc.add_paragraph()

            p.add_run(f"Question {q.get('number')}: ").bold = True

            # Nội dung question viết tiếp cùng paragraph
            if q.get("question"):
                p.add_run(q["question"])
            # Options
            _render_options(doc, q)





def _render_standard_passage_based_group(doc, group):
    """Điền từ, Đọc hiểu, Điền cụm từ - Sử dụng HTML render cho PASSAGE"""
    first_res = group[0]
    parsed_main = first_res.get("parsed", {})
    
    # 1. Render Passage Title (Dùng HTML render)
    if parsed_main.get("passage_title"):
        p_t = doc.add_paragraph()
        p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        _render_html_text(p_t, parsed_main["passage_title"])
        
        # Ép tất cả run trong paragraph thành bold
        for run in p_t.runs:
            run.bold = True
    
    # 2. Render Passage (Sử dụng HTML render cho từng đoạn văn)
    if parsed_main.get("passage"):
        passage_content = str(parsed_main["passage"])
        # Chia nhỏ theo dòng để giữ paragraph
        lines = passage_content.split('\n')
        for line in lines:
            if line.strip():
                p_p = doc.add_paragraph()
                p_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                _render_html_text(p_p, line.strip())

    # 3. Render Questions
    for res in group:
        q_list = res.get("parsed", {}).get("questions", [])
        for q in q_list:
            p = doc.add_paragraph()
            # p.add_run(f"Question {q.get('number')}: ").bold = True
            # # _render_html_text(p, q.get("question_content"))
            # processed_content = _preprocess_question_content(q.get("question_content"))
            # _render_html_text(p, processed_content)
            content = (q.get("question_content") or "").strip()
            number = q.get("number")

            expected = f"Question {number}:"

            # Normalize nhẹ (tránh lệch do khoảng trắng)
            if content.lower() == expected.lower():
                # Nếu giống hoàn toàn → chỉ render 1 lần và bold
                p.add_run(expected).bold = True
            else:
                # Bình thường
                p.add_run(expected + " ").bold = True
                processed_content = _preprocess_question_content(content)
                _render_html_text(p, processed_content)
            
            # doc.add_paragraph(f"A. {q.get('option_a')}\tB. {q.get('option_b')}\tC. {q.get('option_c')}\tD. {q.get('option_d')}")
            _render_options(doc, q)
            # _render_solution_block(doc, q)

def _render_standard_simple_question_group(doc, group):
    for res in group:
        q_list = res.get("parsed", {}).get("questions", [])
        for q in q_list:
            p = doc.add_paragraph()
            
            p.add_run(f"Question {q.get('number')}: ").bold = True
            _render_html_text(p, q.get("question") or q.get("question_content"))
            doc.add_paragraph(f"A. {q.get('option_a')}\tB. {q.get('option_b')}\tC. {q.get('option_c')}\tD. {q.get('option_d')}")
            # _render_solution_block(doc, q)

def _render_standard_pronunciation_group(doc, group):
    for res in group:
        q_list = res.get("parsed", {}).get("questions", [])
        for q in q_list:
            p = doc.add_paragraph()
           
            p.add_run(f"Question {q.get('number')}: ").bold = True
            for label in ['a', 'b', 'c', 'd']:
                p.add_run(f"{label.upper()}. ")
                _render_html_text(p, q.get(f"option_{label}"))
                p.add_run("\t")
            # _render_solution_block(doc, q)
            # if q.get("details"):
            #     for d in q["details"]:
            #         doc.add_paragraph(f"- {d.get('word')}: {d.get('ipa')} ({d.get('pos')}): {d.get('meaning')}")

def _render_standard_error_group(doc, group):
    for res in group:
        q_list = res.get("parsed", {}).get("questions", [])
        for q in q_list:
            p = doc.add_paragraph()
            p.add_run(f"Question {q.get('number')}: ").bold = True
            _render_html_text(p, q.get("question"))
            doc.add_paragraph(f"A. {q.get('option_a')}\tB. {q.get('option_b')}\tC. {q.get('option_c')}\tD. {q.get('option_d')}")
            # _render_solution_block(doc, q)
            # if q.get("correction"):
            #     p_corr = doc.add_paragraph()
            #     p_corr.add_run("Sửa lại: ").bold = True
            #     p_corr.add_run(str(q['correction']))

def _render_standard_dialogue_group(doc, group):
    for res in group:
        q_list = res.get("parsed", {}).get("questions", [])
        for q in q_list:
            p = doc.add_paragraph()
           
            p.add_run(f"Question {q.get('number')}: ").bold = True

            if q.get("speaker_a"):
                doc.add_paragraph(q["speaker_a"])

            if q.get("speaker_b"):
                doc.add_paragraph(q["speaker_b"])

            # ✅ Mỗi đáp án 1 dòng riêng
            if q.get("option_a"):
                doc.add_paragraph(f"A. {q.get('option_a')}")
            if q.get("option_b"):
                doc.add_paragraph(f"B. {q.get('option_b')}")
            if q.get("option_c"):
                doc.add_paragraph(f"C. {q.get('option_c')}")
            if q.get("option_d"):
                doc.add_paragraph(f"D. {q.get('option_d')}")
            # _render_solution_block(doc, q)

def _render_standard_logical_group(doc, group):
    for res in group:
        q_list = res.get("parsed", {}).get("questions", [])
        for q in q_list:
            p = doc.add_paragraph()
           
            p.add_run(f"Question {q.get('number')}: ").bold = True

            if q.get("scenario"):
                doc.add_paragraph(q["scenario"]).italic = True

            if q.get("question"):
                doc.add_paragraph(q["question"])

            # ✅ Mỗi đáp án là 1 paragraph riêng
            if q.get("option_a"):
                doc.add_paragraph(f"A. {q.get('option_a')}")
            if q.get("option_b"):
                doc.add_paragraph(f"B. {q.get('option_b')}")
            if q.get("option_c"):
                doc.add_paragraph(f"C. {q.get('option_c')}")
            if q.get("option_d"):
                doc.add_paragraph(f"D. {q.get('option_d')}")

            # _render_solution_block(doc, q)

def _render_standard_reordering_group(doc, group):
    for res in group:
        q_list = res.get("parsed", {}).get("questions", [])
        for q in q_list:
            p = doc.add_paragraph()
           
            p.add_run(f"Question {q.get('number')}: ").bold = True
            p.add_run(str(q.get("word_list", "")))
            doc.add_paragraph(f"A. {q.get('option_a')}")
            doc.add_paragraph(f"B. {q.get('option_b')}")
            doc.add_paragraph(f"C. {q.get('option_c')}")
            doc.add_paragraph(f"D. {q.get('option_d')}")
            # _render_solution_block(doc, q)

def _add_options_line(doc, parsed):
    p = doc.add_paragraph()

    options = [
        ("A", parsed.get("option_a")),
        ("B", parsed.get("option_b")),
        ("C", parsed.get("option_c")),
        ("D", parsed.get("option_d")),
    ]

    for idx, (label, value) in enumerate(options):
        run = p.add_run(f"{label}. ")
        run.bold = True

        p.add_run(str(value) if value else "")

        if idx < len(options) - 1:
            p.add_run("\t")

def _render_standard_arrange_group(doc, group):
    for res in group:
        parsed = res.get("parsed", {})

        p = doc.add_paragraph()
        p.add_run(f"Question {parsed.get('question_number') or parsed.get('number')}: ").bold = True

        if parsed.get("question_content"):
            for line in parsed["question_content"]:
                doc.add_paragraph(str(line))

        # ✅ FIX HERE
        _add_options_line(doc, parsed)

        # p_ans = doc.add_paragraph()
        # p_ans.add_run("Lời giải").bold = True

        # doc.add_paragraph(f"Chọn {parsed.get('answer')}")

        # if parsed.get("solution_lines"):
        #     for line in parsed["solution_lines"]:
        #         doc.add_paragraph(str(line))

        # if parsed.get("translation_lines"):
        #     p_t_h = doc.add_paragraph()
        #     p_t_h.add_run("Tạm dịch:").bold = True

        #     for line in parsed["translation_lines"]:
        #         doc.add_paragraph(str(line))