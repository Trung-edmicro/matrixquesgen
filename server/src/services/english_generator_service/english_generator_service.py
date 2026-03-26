import random

from bs4 import BeautifulSoup
from fastapi import HTTPException
import pandas as pd
from collections import defaultdict
from pathlib import Path
import uuid
import os
import json
import requests
import shutil
from api.callApi import get_credentials
# from services.english_generator_service.vertex_async_client import AsyncVertexClient
from .docx_helper_english import add_formatted_paragraph, add_html_formatted_text
from services.english_generator_service.vertex_async_3_1_model import AsyncVertexGemini31
import asyncio
import re
import html
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
import logging
from collections import defaultdict, Counter
from .constants import (CLOZE_EXPLANATION_TEMPLATE, 
                        ARRANGE_JSON_SCHEMA,
     CLOZE_WITH_TITLE_JSON_SCHEMA, CLOZE_JSON_SCHEMA,SENTENCE_COMPLETION_JSON_SCHEMA,SYNONYM_ANTONYM_JSON_SCHEMA,ERROR_IDENTIFICATION_JSON_SCHEMA,SENTENCE_TRANSFORMATION_JSON_SCHEMA,WORD_REORDERING_JSON_SCHEMA,PRONUNCIATION_STRESS_JSON_SCHEMA,DIALOGUE_COMPLETION_JSON_SCHEMA,LOGICAL_THINKING_JSON_SCHEMA,
       READING_COMPREHENSION_EXPLANATION_TEMPLATE, SILENT_PHASE_EXPLANATION_TEMPLATE, PROMPTS)
logger = logging.getLogger(__name__)
from .drive_helper_services import load_vocabulary_from_drive

# ============================
# CONFIG
# ============================

LEVELS = ["Nhận biết", "Thông hiểu", "Vận dụng", "Vận dụng cao"]

APP_DIR = Path(os.environ['APP_DIR']) if os.environ.get('APP_DIR') else Path(__file__).parent.parent.parent.parent.parent
PROMPT_DIR = APP_DIR / "data" / "prompts" / "TIENGANH"
print(f">>>>>> debug APP_DIR: {APP_DIR}")
UPLOAD_DIR = APP_DIR / "data" / "uploads"
OUTPUT_DIR = APP_DIR / "data" / "outputs"

UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

DRIVE_ENGLISH_PROMPT_FOLDER = "https://drive.google.com/drive/folders/1JSFC8FBTY6lA0rlrC7-LAIHU_FjbOK3g"

# DRIVE_VOCABULARY_FOLDER_C10 = "https://drive.google.com/drive/folders/18tVQXctKZdpj8ZrFJhA0cU-xpLBj5Du2"

# DRIVE_VOCABULARY_FOLDER_C11 = "https://drive.google.com/drive/folders/1FqzI2Y-zIWUMnDNCrFCc_Yw-yR0Pm8PT"

# DRIVE_VOCABULARY_FOLDER_C12 = "https://drive.google.com/drive/folders/16Ke0JMipcJbHMIWEWiV1rQh234Mei0pV"

# ============================
# PROMPT LOADER
# ============================



# def load_prompt(filename: str) -> str:
#     path = PROMPT_DIR / filename
#     if not path.exists():
#         raise Exception(f"Không tìm thấy prompt: {filename}")
#     return path.read_text(encoding="utf-8")


DRIVE_FOLDER = "https://drive.google.com/drive/folders/1JSFC8FBTY6lA0rlrC7-LAIHU_FjbOK3g"
PROMPT_DIR = Path("PROMPT_DIR")

_drive_prompts_cache = None


def fetch_drive_md_files():
    try:
        folder_id = re.search(r"folders/([a-zA-Z0-9_-]+)", DRIVE_FOLDER).group(1)

        url = f"https://drive.google.com/drive/folders/{folder_id}"
        html = requests.get(url).text

        file_ids = list(set(re.findall(r'"([a-zA-Z0-9_-]{25,})"', html)))

        prompts = {}

        for fid in file_ids:
            download_url = f"https://drive.google.com/uc?export=download&id={fid}"
            r = requests.get(download_url)

            if r.status_code != 200:
                continue

            disposition = r.headers.get("content-disposition", "")

            if ".md" not in disposition:
                continue

            name_match = re.findall(r'filename="(.+)"', disposition)

            if not name_match:
                continue

            name = name_match[0]
            prompts[name] = r.text

        return prompts

    except Exception:
        return {}


def load_prompt(filename: str) -> str:
    global _drive_prompts_cache

    # Load drive 1 lần
    if _drive_prompts_cache is None:
        print("🔎 Fetching prompts from Drive...")
        _drive_prompts_cache = fetch_drive_md_files()

    # Nếu có trên drive
    if filename in _drive_prompts_cache:
        print(f"✅ Load từ Drive: {filename}")
        return _drive_prompts_cache[filename]

    # fallback local
    path = PROMPT_DIR / filename

    if not path.exists():
        raise Exception(f"Không tìm thấy prompt: {filename}")

    print(f"📂 Load từ LOCAL: {filename}")
    return path.read_text(encoding="utf-8")



# ============================
# EXCEL EXTRACT LOGIC
# ============================

# def detect_type_columns(df):
#     col_map = {}
#     for i, col in enumerate(df.columns):
#         col_lower = str(col).lower()
#         if "điền từ" in col_lower:
#             col_map["Điền từ"] = i
#         elif "sắp xếp" in col_lower:
#             col_map["Sắp xếp"] = i
#         elif "đọc hiểu" in col_lower:
#             col_map["Đọc hiểu"] = i
#         elif "điền cụm" in col_lower:
#             col_map["Điền cụm từ/điền câu"] = i
        
#     return col_map

def detect_type_columns(df):
    col_map = {}
    for i, col in enumerate(df.columns):
        col_lower = str(col).lower()

        if "điền từ" in col_lower:
            col_map["Điền từ"] = i
        elif "sắp xếp từ" in col_lower:
            col_map["Sắp xếp từ"] = i
        elif "sắp xếp" in col_lower:
            col_map["Sắp xếp"] = i
        elif "đọc hiểu" in col_lower:
            col_map["Đọc hiểu"] = i
        elif "điền cụm" in col_lower:
            col_map["Điền cụm từ/điền câu"] = i

        # ==== ADD THÊM ====
        elif "hoàn thành câu" in col_lower:
            col_map["Hoàn thành câu"] = i
        elif "đồng nghĩa" in col_lower or "trái nghĩa" in col_lower:
            col_map["Đồng nghĩa/Trái nghĩa"] = i
        elif "tìm lỗi sai" in col_lower:
            col_map["Tìm lỗi sai"] = i
        elif "kết hợp" in col_lower or "viết lại" in col_lower:
            col_map["Kết hợp/viết lại câu"] = i
        elif "phát âm" in col_lower or "trọng âm" in col_lower:
            col_map["Phát âm/Trọng âm"] = i
        elif "giao tiếp" in col_lower:
            col_map["Câu giao tiếp"] = i
        elif "tình huống" in col_lower or "tư duy" in col_lower:
            col_map["Tư duy/Tình huống"] = i
        # ==================

    return col_map


def detect_all_levels(row, start_index):
    found = []
    for offset in range(4):
        cell = row.iloc[start_index + offset]
        if pd.notna(cell) and str(cell).strip() != "":
            found.append(LEVELS[offset])
    return found


META_COLS = ["STT", "Chủ đề", "Số từ","Từ vựng", "Độ khó", "Dạng thức bài đọc (VI)","Dạng thức bài đọc (EN)","Từ vựng tham khảo","Tài liệu tham khảo"]


def extract_blocks_from_excel(file_path: str):

    df = pd.read_excel(file_path, sheet_name="Ma trận")
    df.columns = [str(col).strip() for col in df.columns]
    df[META_COLS] = df[META_COLS].ffill()

    type_col_map = detect_type_columns(df)
    blocks = []
    total_counts = Counter()
    grouped = df.groupby("STT")

    for stt, group in grouped:
        topic = group.iloc[0]["Chủ đề"]
        vocabulary_example = group.iloc[0]["Từ vựng"]
        word_count = str(group.iloc[0]["Số từ"])
        difficulty = group.iloc[0]["Độ khó"]
        text_type = group.iloc[0]["Dạng thức bài đọc (VI)"]
        text_type_en = group.iloc[0]["Dạng thức bài đọc (EN)"]
        vocabulary = group.iloc[0]["Từ vựng tham khảo"]
        document_sample = group.iloc[0]["Tài liệu tham khảo"]
        question_types = defaultdict(list)

        for _, row in group.iterrows():
            spec = row.get("Đặc tả ma trận")
            if pd.isna(spec):
                continue
            spec = str(spec).strip()
            for q_type, start_col in type_col_map.items():
                levels = detect_all_levels(row, start_col)
                for lv in levels:
                    question_types[q_type].append({"spec": spec, "level": lv})

        print(f"\n=== STT {stt} ===")
        for q_type, questions in question_types.items():
            count = len(questions)
            print(f"{q_type}: {count} questions")
            total_counts[q_type] += count
            blocks.append({
                "type": q_type,
                "topic": topic,
                "difficulty": difficulty,
                "text_type": text_type,
                "text_type_en": text_type_en,
                "word_count": word_count,
                "question_count": count,
                "questions": questions,
                "vocabulary": vocabulary,
                "document_sample": document_sample,
                "vocabulary_example": vocabulary_example

            })

    print("\n=== TOTAL QUESTIONS PER COLUMN ===")
    for q_type, count in total_counts.items():
        print(f"{q_type}: {count}")

    return blocks


# ============================
# JSON PARSE HELPERS
# ============================

def _safe_parse_json(raw: str) -> dict | None:
    """
    Cố gắng parse JSON từ response của AI.
    Xử lý trường hợp AI trả về markdown fences hoặc text thừa.
    """
    if not raw:
        return None

    # Strip markdown code fences nếu có
    text = raw.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```$", "", text)
    text = text.strip()

    # Tìm JSON object đầu tiên
    start = text.find("{")
    if start == -1:
        logger.warning("No JSON object found in AI response")
        return None

    # Tìm closing brace tương ứng (depth counting)
    depth = 0
    end = -1
    in_string = False
    escape_next = False

    for i, ch in enumerate(text[start:], start=start):
        if escape_next:
            escape_next = False
            continue
        if ch == "\\" and in_string:
            escape_next = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                end = i
                break

    if end == -1:
        logger.warning("Could not find closing brace in AI response")
        return None

    json_str = text[start:end + 1]

    try:
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        logger.warning(f"JSON parse error: {e}. Attempting cleanup...")

        # Cleanup: loại bỏ trailing commas trước } hoặc ]
        json_str = re.sub(r",\s*([\]}])", r"\1", json_str)
        try:
            return json.loads(json_str)
        except json.JSONDecodeError as e2:
            logger.error(f"JSON parse failed after cleanup: {e2}")
            return None


# ============================
# GENERATE DOCX LOGIC
# ============================


def safe_str(val):
    if val is None:
        return ""
    if isinstance(val, float) and pd.isna(val):
        return ""
    s = str(val).strip()
    return "" if s.lower() == "nan" else s

SEM = asyncio.Semaphore(3)   # có thể chỉnh 2–5 tùy quota

async def generate_with_retry(client, prompt, max_retries=5):
    for attempt in range(max_retries):
        try:
            return await client.generate(prompt=prompt)

        except Exception as e:
            if "429" in str(e) or "RESOURCE_EXHAUSTED" in str(e):
                wait_time = (2 ** attempt) + random.uniform(0, 1)
                logger.warning(f"⚠️ 429 hit. Retry after {wait_time:.2f}s")
                await asyncio.sleep(wait_time)
            else:
                raise e

    raise Exception("❌ Max retries exceeded")


async def limited_generate(client, prompt):
    async with SEM:
        result = await generate_with_retry(client, prompt)

        # throttle nhẹ để tránh spike
        await asyncio.sleep(2)

        return result


async def run_in_batches(tasks, batch_size=3):
    results = []
    for i in range(0, len(tasks), batch_size):
        batch = tasks[i:i + batch_size]
        res = await asyncio.gather(*batch, return_exceptions=True)
        results.extend(res)
    return results

async def generate_exam_docx(blocks, output_path):

    credentials, project_id = get_credentials()

    # client = AsyncVertexClient(
    #     project_id=project_id,
    #     creds=credentials,
    #     model="gemini-2.5-pro"
    # )

    client = AsyncVertexGemini31(
    project_id="onluyen-media",
    model="gemini-3.1-pro-preview",
    thinking_level="HIGH"
)
    # doc.add_heading("ĐỀ THI TIẾNG ANH", level=1)

    q_count = 1
    tasks = []
    block_meta = []

    for block in blocks:

        topic     = block["topic"]
        vocabulary = block["vocabulary"]
        vocabulary_example = block["vocabulary_example"]
        if pd.isna(vocabulary) or str(vocabulary).strip() == "":
            # vocabulary = load_vocabulary_from_drive(topic) or ""
            vocabulary = load_vocabulary_from_drive(vocabulary_example) or ""

        vocabulary = str(vocabulary)
        q_type    = block["type"]
        text_type = block["text_type"]
        text_type_en = block["text_type_en"]
        diff      = block["difficulty"]
        so_tu     = block.get("word_count", "")
        questions = block["questions"]
        document_sample = block["document_sample"]
       
        n_q       = len(questions)

        try:
            prompt_template = load_prompt(PROMPTS[q_type])
        except Exception:
            prompt_template = ""

        specs_list = [
            f"Câu {i+1}: {q['spec']} (Cấp độ: {q['level']})"
            for i, q in enumerate(questions)
        ]

        # ===============================
        # 1️⃣ CLOZE (Điền từ)
        # ===============================

        if q_type == "Điền từ":

            formatted_cloze_prompt = (
                prompt_template
                    .replace("{TOPIC_NAME}", safe_str(topic))
                    .replace("{TEXT_TYPE}", safe_str(text_type))
                    .replace("{CEFR_LEVEL}", safe_str(diff))
                    .replace("{VOCABULARY_LIST}", safe_str(vocabulary))
                    .replace("{TARGET_WORD_COUNT}", safe_str(so_tu))
                    .replace("{SOURCE_TEXT}", safe_str(document_sample))
            )

            # Chọn schema có/không có tiêu đề tùy text_type
            has_title = text_type in ("Quảng cáo", "Thông báo", "Advertisement", "Announcement")
            if has_title:
                output_rule_cloze = CLOZE_WITH_TITLE_JSON_SCHEMA.format(
                    TEXT_TYPE=text_type,
                    N_Q=n_q,
                    START_NUM=q_count
                )
            else:
                output_rule_cloze = CLOZE_JSON_SCHEMA.format(
                    N_Q=n_q,
                    START_NUM=q_count
                )

            ai_input_cloze = (
                f"{formatted_cloze_prompt}\n\n"
                f"Chủ đề: {topic}\n"
                f"Từ vựng tham khảo: {vocabulary}\n"
                f"## EXPLANATION MICRO-FORMAT (STRICT)\n"
                f"{CLOZE_EXPLANATION_TEMPLATE}\n\n"
                f"Số từ: {so_tu}\n"
                f"Độ khó: {diff}\n"
                f"Dạng thức: {text_type}\n"
                f"Số câu: {n_q}\n"
                + "\n".join(specs_list)
                + "\n\n## OUTPUT FORMAT\n"
                + output_rule_cloze
            )

            # task = client.generate(prompt=ai_input)
            task = limited_generate(client, ai_input_cloze)
            tasks.append(task)
            block_meta.append(("CLOZE", topic, n_q, q_count,text_type_en))
            q_count += n_q

        # ===============================
        # 2️⃣ ARRANGE (Sắp xếp)
        # ===============================

        elif q_type == "Sắp xếp":

            spec_item = questions[0]
            output_rule_arranage = ARRANGE_JSON_SCHEMA.format(START_NUM=q_count)
            formatted_arrange_prompt = (
                            prompt_template
                                .replace("{TOPIC_NAME}", safe_str(topic))
                                .replace("{TEXT_TYPE}", safe_str(text_type))
                                .replace("{CEFR_LEVEL}", safe_str(diff))
                                .replace("{VOCABULARY_LIST}", safe_str(vocabulary))
                                .replace("{TARGET_WORD_COUNT}", safe_str(so_tu))
                        )

            ai_input_arrange = (
                f"{formatted_arrange_prompt}\n\n"
                f"Chủ đề: {topic}\n"
                f"Từ vựng tham khảo:{vocabulary}\n"
                f"Độ khó: {diff}\n"
                f"Dạng thức: {text_type}\n"
                f"Đặc tả ma trận: {spec_item['spec']}\n"
                f"Mức độ: {spec_item['level']}\n"
                f"Đánh số câu: Question {q_count}\n\n"
                f"## OUTPUT FORMAT\n"
                + output_rule_arranage
            )

            # task = client.generate(prompt=ai_input)
            task = limited_generate(client, ai_input_arrange)
            tasks.append(task)
            block_meta.append(("ARRANGE", topic, 1, q_count, text_type_en))
            q_count += 1

        # ===============================
        # 3️⃣ READING COMPREHENSION (Đọc hiểu)
        # ===============================

        elif q_type == "Đọc hiểu":

            output_rule_reading_comprehensive = CLOZE_WITH_TITLE_JSON_SCHEMA.format(
                TEXT_TYPE=text_type,
                N_Q=n_q,
                START_NUM=q_count
            )

            formatted_reading_prompt = (
                prompt_template
                    .replace("{TOPIC_NAME}", safe_str(topic))
                    .replace("{TEXT_TYPE}", safe_str(text_type))
                    .replace("{CEFR_LEVEL}", safe_str(diff))
                    .replace("{VOCABULARY_LIST}", safe_str(vocabulary))
                    .replace("{TARGET_WORD_COUNT}", safe_str(so_tu))
                    .replace("{SOURCE_TEXT}", safe_str(document_sample))
            )

            ai_input_reading_comprehensive = (
                f"{formatted_reading_prompt}\n\n"
                f"Chủ đề: {topic}\n"
                f"Từ vựng tham khảo: {vocabulary}"
                f"Tài liệu tham khảo: {document_sample}"
                f"Số từ: {so_tu}\n"
                f"Độ khó: {diff}\n"
                f"## EXPLANATION MICRO-FORMAT (STRICT)\n"
                f"{READING_COMPREHENSION_EXPLANATION_TEMPLATE}"
                f"Dạng thức: {text_type}\n"
                f"Số câu: {n_q}\n"
                + "\n".join(specs_list)
                + "\n\n## OUTPUT FORMAT\n"
                + output_rule_reading_comprehensive
            )

            # task = client.generate(prompt=ai_input)
            task = limited_generate(client, ai_input_reading_comprehensive)
            tasks.append(task)
            block_meta.append(("RC", topic, n_q, q_count,text_type_en))
            q_count += n_q

        # ===============================
        # 4️⃣ GAP FILL (Điền cụm từ/điền câu)
        # ===============================

        elif q_type == "Điền cụm từ/điền câu":

            output_rule_fill_in_phrase = CLOZE_WITH_TITLE_JSON_SCHEMA.format(
                TEXT_TYPE=text_type,
                N_Q=n_q,
                START_NUM=q_count
            )

            formatted_silent_prompt = (
                        prompt_template
                    .replace("{TOPIC_NAME}", safe_str(topic))
                    .replace("{TEXT_TYPE}", safe_str(text_type))
                    .replace("{CEFR_LEVEL}", safe_str(diff))
                    .replace("{VOCABULARY_LIST}", safe_str(vocabulary))
                    .replace("{TARGET_WORD_COUNT}", safe_str(so_tu))
                    .replace("{SOURCE_TEXT}", safe_str(document_sample))
            )


            ai_input_silent = (
                f"{formatted_silent_prompt}\n\n"
                f"Chủ đề: {topic}\n"
                f"Từ vựng tham khảo: {vocabulary}"
                f"Tài liệu tham khảo: {document_sample}"
                f"## EXPLANATION MICRO-FORMAT (STRICT)\n"
                f"Cấm tuyệt đối không viết lại câu hỏi trong phần giải thích"
                f"{SILENT_PHASE_EXPLANATION_TEMPLATE}\n\n"
                f"Số từ: {so_tu}\n"
                f"Độ khó: {diff}\n"
                f"Dạng thức: {text_type}\n"
                f"Số câu: {n_q}\n"
                + "\n".join(specs_list)
                + "\n\n## OUTPUT FORMAT\n"
                + output_rule_fill_in_phrase
            )
           

            # task = client.generate(prompt=ai_input)
            task = limited_generate(client, ai_input_silent)
            tasks.append(task)
            block_meta.append(("GAP", topic, n_q, q_count,text_type_en))
            q_count += n_q
        
        elif q_type == "Hoàn thành câu":
            spec_item = questions[0]
            output_rule_sentence_completion = SENTENCE_COMPLETION_JSON_SCHEMA.format(
                N_Q = n_q,
                START_NUM = q_count
            )

            formatted_sentence_prompt = (prompt_template
                    .replace("{TOPIC_NAME}", safe_str(topic))
                    .replace("{TEXT_TYPE}", safe_str(text_type))
                    .replace("{CEFR_LEVEL}", safe_str(diff))
                    .replace("{VOCABULARY_LIST}", safe_str(vocabulary))
                    .replace("{TARGET_WORD_COUNT}", safe_str(so_tu))
                    .replace("{SOURCE_TEXT}", safe_str(document_sample)))
            ai_input_sentence_completion = (
                f"{formatted_sentence_prompt}\n\n"
                f"Từ vựng tham khảo: {vocabulary}"
                f"Tài liệu tham khảo: {document_sample}"
                f"Đặc tả ma trận: {spec_item['spec']}\n"
                f"Mức độ: {spec_item['level']}\n"
                f"Cấm tuyệt đối không viết lại câu hỏi trong phần giải thích"
                f"Dạng thức: {text_type}\n"
                f"Số câu: {n_q}\n"
                + "\n".join(specs_list)
                + "\n\n## OUTPUT FORMAT\n"
                + output_rule_sentence_completion
            )

            task = limited_generate(client, ai_input_sentence_completion)
            tasks.append(task)
            block_meta.append(("SENTENCE_COMPLETION", topic, n_q, q_count,text_type_en))
            q_count += n_q

        elif q_type == "Đồng nghĩa/Trái nghĩa":
            spec_item = questions[0]
            output_rule_synonum_antonym = SYNONYM_ANTONYM_JSON_SCHEMA.format(
                N_Q = n_q,
                START_NUM = q_count
            )

            formatted_synonum_antonym_prompt = (prompt_template
                    .replace("{TOPIC_NAME}", safe_str(topic))
                    .replace("{TEXT_TYPE}", safe_str(text_type))
                    .replace("{CEFR_LEVEL}", safe_str(diff))
                    .replace("{VOCABULARY_LIST}", safe_str(vocabulary))
                    .replace("{TARGET_WORD_COUNT}", safe_str(so_tu))
                    .replace("{SOURCE_TEXT}", safe_str(document_sample)))
            
            ai_input_synonym_antonym = (
                f"{formatted_synonum_antonym_prompt}\n\n"
                f"Từ vựng tham khảo: {vocabulary}"
                f"Tài liệu tham khảo: {document_sample}"
                f"Đặc tả ma trận: {spec_item['spec']}\n"
                f"Mức độ: {spec_item['level']}\n"
                f"Cấm tuyệt đối không viết lại câu hỏi trong phần giải thích"
                f"Dạng thức: {text_type}\n"
                f"Số câu: {n_q}\n"
                + "\n".join(specs_list)
                + "\n\n## OUTPUT FORMAT\n"
                + output_rule_synonum_antonym
            )

            task = limited_generate(client, ai_input_synonym_antonym)
            tasks.append(task)
            block_meta.append(("SYNONYM_ANTONYM", topic, n_q, q_count,text_type_en))
            q_count += n_q

        elif q_type == "Tìm lỗi sai":
            spec_item = questions[0]
            output_rule_error_identification = ERROR_IDENTIFICATION_JSON_SCHEMA.format(
                N_Q = n_q,
                START_NUM = q_count
            )

            formatted_error_identification_prompt = (prompt_template
                    .replace("{TOPIC_NAME}", safe_str(topic))
                    .replace("{TEXT_TYPE}", safe_str(text_type))
                    .replace("{CEFR_LEVEL}", safe_str(diff))
                    .replace("{VOCABULARY_LIST}", safe_str(vocabulary))
                    .replace("{TARGET_WORD_COUNT}", safe_str(so_tu))
                    .replace("{SOURCE_TEXT}", safe_str(document_sample)))
            ai_input_error_identification = (
                f"{formatted_error_identification_prompt}\n\n"
                f"Từ vựng tham khảo: {vocabulary}"
                f"Tài liệu tham khảo: {document_sample}"
                f"Đặc tả ma trận: {spec_item['spec']}\n"
                f"Mức độ: {spec_item['level']}\n"
                f"Cấm tuyệt đối không viết lại câu hỏi trong phần giải thích"
                f"Dạng thức: {text_type}\n"
                f"Số câu: {n_q}\n"
                + "\n".join(specs_list)
                + "\n\n## OUTPUT FORMAT\n"
                + output_rule_error_identification
            )

            task = limited_generate(client, ai_input_error_identification)
            tasks.append(task)
            block_meta.append(("ERROR_IDENTIFICATION", topic, n_q, q_count,text_type_en))
            q_count += n_q
        
        elif q_type == "Kết hợp/viết lại câu":
            spec_item = questions[0]
            output_rule_sentence_transformation = SENTENCE_TRANSFORMATION_JSON_SCHEMA.format(
                N_Q = n_q,
                START_NUM = q_count
            )

            formatted_sentence_transformation_prompt = (prompt_template
                    .replace("{TOPIC_NAME}", safe_str(topic))
                    .replace("{TEXT_TYPE}", safe_str(text_type))
                    .replace("{CEFR_LEVEL}", safe_str(diff))
                    .replace("{VOCABULARY_LIST}", safe_str(vocabulary))
                    .replace("{TARGET_WORD_COUNT}", safe_str(so_tu))
                    .replace("{SOURCE_TEXT}", safe_str(document_sample)))
            
            ai_input_sentence_transformation = (
                f"{formatted_sentence_transformation_prompt}\n\n"
                f"Từ vựng tham khảo: {vocabulary}"
                f"Tài liệu tham khảo: {document_sample}"
                f"Đặc tả ma trận: {spec_item['spec']}\n"
                f"Mức độ: {spec_item['level']}\n"
                f"Cấm tuyệt đối không viết lại câu hỏi trong phần giải thích"
                f"Dạng thức: {text_type}\n"
                f"Số câu: {n_q}\n"
                + "\n".join(specs_list)
                + "\n\n## OUTPUT FORMAT\n"
                + output_rule_sentence_transformation
            )

            print(f">>>>> debug ai_input_sentence_transformation: {ai_input_sentence_transformation}")

            task = limited_generate(client, ai_input_sentence_transformation)
            tasks.append(task)
            block_meta.append(("SENTENCE_TRANSFORMATION", topic, n_q, q_count,text_type_en))
            q_count += n_q

        elif q_type == "Sắp xếp từ":
            spec_item = questions[0]
            output_rule_word_reordering = WORD_REORDERING_JSON_SCHEMA.format(
                N_Q = n_q,
                START_NUM = q_count
            )

            formatted_word_reordering_prompt = (prompt_template
                    .replace("{TOPIC_NAME}", safe_str(topic))
                    .replace("{TEXT_TYPE}", safe_str(text_type))
                    .replace("{CEFR_LEVEL}", safe_str(diff))
                    .replace("{VOCABULARY_LIST}", safe_str(vocabulary))
                    .replace("{TARGET_WORD_COUNT}", safe_str(so_tu))
                    .replace("{SOURCE_TEXT}", safe_str(document_sample)))
            
            ai_input_word_reordering = (
                f"{formatted_word_reordering_prompt}\n\n"
                f"Từ vựng tham khảo: {vocabulary}"
                f"Tài liệu tham khảo: {document_sample}"
                f"Đặc tả ma trận: {spec_item['spec']}\n"
                f"Mức độ: {spec_item['level']}\n"
                f"Cấm tuyệt đối không viết lại câu hỏi trong phần giải thích"
                f"Dạng thức: {text_type}\n"
                f"Số câu: {n_q}\n"
                + "\n".join(specs_list)
                + "\n\n## OUTPUT FORMAT\n"
                + output_rule_word_reordering
            )

            task = limited_generate(client, ai_input_word_reordering)
            tasks.append(task)
            block_meta.append(("WORD_REORDERING", topic, n_q, q_count,text_type_en))
            q_count += n_q
        
        elif q_type == "Phát âm/Trọng âm":
            spec_item = questions[0]
            output_rule_pronounciation_stress = PRONUNCIATION_STRESS_JSON_SCHEMA.format(
                N_Q = n_q,
                START_NUM = q_count
            )

            formatted_word_reordering_prompt = (prompt_template
                    .replace("{TOPIC_NAME}", safe_str(topic))
                    .replace("{TEXT_TYPE}", safe_str(text_type))
                    .replace("{CEFR_LEVEL}", safe_str(diff))
                    .replace("{VOCABULARY_LIST}", safe_str(vocabulary))
                    .replace("{TARGET_WORD_COUNT}", safe_str(so_tu))
                    .replace("{SOURCE_TEXT}", safe_str(document_sample)))
            ai_input_pronounciation_stress = (
                f"{formatted_word_reordering_prompt}\n\n"
                f"Từ vựng tham khảo: {vocabulary}"
                f"Tài liệu tham khảo: {document_sample}"
                f"Đặc tả ma trận: {spec_item['spec']}\n"
                f"Mức độ: {spec_item['level']}\n"
                f"Cấm tuyệt đối không viết lại câu hỏi trong phần giải thích"
                f"Dạng thức: {text_type}\n"
                f"Số câu: {n_q}\n"
                + "\n".join(specs_list)
                + "\n\n## OUTPUT FORMAT\n"
                + output_rule_pronounciation_stress
            )

            task = limited_generate(client, ai_input_pronounciation_stress)
            tasks.append(task)
            block_meta.append(("PRONUNCIATION_STRESS", topic, n_q, q_count,text_type_en))
            q_count += n_q

        elif q_type == "Câu giao tiếp":
            spec_item = questions[0]
            output_rule_dialouge_completion = DIALOGUE_COMPLETION_JSON_SCHEMA.format(
                N_Q = n_q,
                START_NUM = q_count
            )

            formatted_dialouge_completion_prompt = (prompt_template
                    .replace("{TOPIC_NAME}", safe_str(topic))
                    .replace("{TEXT_TYPE}", safe_str(text_type))
                    .replace("{CEFR_LEVEL}", safe_str(diff))
                    .replace("{VOCABULARY_LIST}", safe_str(vocabulary))
                    .replace("{TARGET_WORD_COUNT}", safe_str(so_tu))
                    .replace("{SOURCE_TEXT}", safe_str(document_sample)))
            
            ai_input_pronounciation_stress = (
                f"{formatted_dialouge_completion_prompt}\n\n"
                f"Từ vựng tham khảo: {vocabulary}"
                f"Tài liệu tham khảo: {document_sample}"
                f"Đặc tả ma trận: {spec_item['spec']}\n"
                f"Mức độ: {spec_item['level']}\n"
                f"Cấm tuyệt đối không viết lại câu hỏi trong phần giải thích"
                f"Dạng thức: {text_type}\n"
                f"Số câu: {n_q}\n"
                + "\n".join(specs_list)
                + "\n\n## OUTPUT FORMAT\n"
                + output_rule_dialouge_completion
            )

            task = limited_generate(client, ai_input_pronounciation_stress)
            tasks.append(task)
            block_meta.append(("DIALOUGE_COMPLETION", topic, n_q, q_count,text_type_en))
            q_count += n_q

        elif q_type == "Tư duy/Tình huống":
            spec_item = questions[0]

            output_rule_logical_thinking = LOGICAL_THINKING_JSON_SCHEMA.format(
                N_Q = n_q,
                START_NUM = q_count
            )

            formatted_logical_thinking_prompt = (prompt_template
                    .replace("{TOPIC_NAME}", safe_str(topic))
                    .replace("{TEXT_TYPE}", safe_str(text_type))
                    .replace("{CEFR_LEVEL}", safe_str(diff))
                    .replace("{VOCABULARY_LIST}", safe_str(vocabulary))
                    .replace("{TARGET_WORD_COUNT}", safe_str(so_tu))
                    .replace("{SOURCE_TEXT}", safe_str(document_sample)))
            
            ai_input_logical_thinking = (
                f"{formatted_logical_thinking_prompt}\n\n"
                f"Từ vựng tham khảo: {vocabulary}"
                f"Tài liệu tham khảo: {document_sample}"
                f"Đặc tả ma trận: {spec_item['spec']}\n"
                f"Mức độ: {spec_item['level']}\n"
                f"Cấm tuyệt đối không viết lại câu hỏi trong phần giải thích"
                f"Dạng thức: {text_type}\n"
                f"Số câu: {n_q}\n"
                + "\n".join(specs_list)
                + "\n\n## OUTPUT FORMAT\n"
                + output_rule_logical_thinking
            )

            task = limited_generate(client, ai_input_logical_thinking)
            tasks.append(task)
            block_meta.append(("LOGICAL_THINKING", topic, n_q, q_count,text_type_en))
            q_count += n_q

    responses = await asyncio.gather(*tasks, return_exceptions=True)

    results = []
    for (block_type, topic, n_q, start_num, text_type_en), response_text in zip(block_meta, responses):
        if isinstance(response_text, Exception):
            logger.error(f"Block {block_type}/{topic} failed: {response_text}")
            response_text = ""

        parsed = _safe_parse_json(response_text)

        results.append({
            "type": block_type,
            "title": topic,    # raw text (giữ lại để debug)
            "parsed": parsed,            # dict đã parse
            "question_count": n_q,
            "start_num": start_num,
            "text_type_en": text_type_en.lower()
        })
        # print(f">>>>>> debug results {results}")

    # generate_docx_from_ai_results(results, output_path)
    return results


# ============================
# DOCX GENERATION FROM RESULTS
# ============================

def generate_docx_from_ai_results(results, output_path):
    doc = Document()
    _apply_default_style(doc)

    for res in results:
        _add_instruction(doc, res)
        parsed = res.get("parsed")

        if parsed is None:
            # Fallback: dùng raw text parser cũ nếu JSON parse thất bại
            logger.warning(f"Falling back to text parser for block {res['type']}/{res['title']}")
            raw = res.get("data") or ""
            _render_fallback(doc, res["type"], raw)
        else:
            if res["type"] in ("CLOZE", "GAP", "RC"):
                _render_cloze_from_json(doc, parsed, merge_options=(res["type"] == "CLOZE"),qtype=res["type"])
            # if res["type"] == "SENTENCE_COMPLETION":

            elif res["type"] == "ARRANGE":
                _render_arrange_from_json(doc, parsed)


    doc.save(output_path)


def export_standard_docx_from_data(json_data, output_path):
    """
    Nhận trực tiếp JSON object (dict), lấy key 'results', render docx.
    Hỗ trợ cả format mới (có 'parsed') và format cũ (chỉ có 'data' text).
    """
    results = json_data.get("results")
    if not results:
        raise Exception("No 'results' found in JSON")

    # Nếu result chưa có 'parsed', thử parse từ 'data'
    for res in results:
        if "parsed" not in res or res["parsed"] is None:
            res["parsed"] = _safe_parse_json(res.get("data") or "")

    doc = Document()
    _apply_default_style(doc)

    for res in results:
        _add_instruction(doc, res)
        parsed = res.get("parsed")
        res_type = res.get("type")

        if parsed is None:
            raw = res.get("data") or ""
            _render_fallback(doc, res_type, raw)
        else:
            if res_type in ("CLOZE", "GAP", "RC"):
                _render_standard_cloze_from_json(doc, parsed, merge_options=(res_type == "CLOZE"))
            elif res_type == "ARRANGE":
                _render_standard_arrange_from_json(doc, parsed)
            else:
                logger.warning(f"Unknown type: {res_type}")


    doc.save(output_path)
    return output_path



# def export_docx_from_data(json_data, output_path):
#     """
#     Nhận trực tiếp JSON object (dict), lấy key 'results', render docx.
#     Hỗ trợ cả format mới (có 'parsed') và format cũ (chỉ có 'data' text).
#     """
#     results = json_data.get("results")
#     if not results:
#         raise Exception("No 'results' found in JSON")

#     # Nếu result chưa có 'parsed', thử parse từ 'data'
#     for res in results:
#         if "parsed" not in res or res["parsed"] is None:
#             res["parsed"] = _safe_parse_json(res.get("data") or "")

#     doc = Document()
#     _apply_default_style(doc)

#     for res in results:
#         _add_instruction(doc, res)
#         parsed = res.get("parsed")
#         res_type = res.get("type")

#         if parsed is None:
#             raw = res.get("data") or ""
#             _render_fallback(doc, res_type, raw)
#         else:
#             if res_type in ("CLOZE", "GAP", "RC"):
#                 _render_cloze_from_json(doc, parsed, merge_options=(res_type == "CLOZE"))
#             elif res_type == "ARRANGE":
#                 _render_arrange_from_json(doc, parsed)
#             else:
#                 logger.warning(f"Unknown type: {res_type}")

#         _add_separator(doc)

#     doc.save(output_path)
#     return output_path

def export_docx_from_data(json_data, output_path):
    """
    Render docx từ JSON.
    - Hỗ trợ parsed + raw data
    - Tối ưu: group SYNONYM / ANTONYM liên tiếp
    """
    results = json_data.get("results")
    if not results:
        raise Exception("No 'results' found in JSON")

    # ===== Ensure parsed =====
    for res in results:
        if "parsed" not in res or res["parsed"] is None:
            res["parsed"] = _safe_parse_json(res.get("data") or "")

    doc = Document()
    _apply_default_style(doc)

    i = 0
    n = len(results)

    while i < n:
        res = results[i]
        parsed = res.get("parsed")
        res_type = res.get("type")

        # ===== FALLBACK =====
        if parsed is None:
            _add_instruction(doc, res)
            raw = res.get("data") or ""
            _render_fallback(doc, res_type, raw)
            _add_separator(doc)
            i += 1
            continue

        # ===== CLOZE / GAP / RC =====
        if res_type in ("CLOZE", "GAP", "RC"):
            _add_instruction(doc, res)
            _render_cloze_from_json(doc, parsed, merge_options=(res_type == "CLOZE"))
           
            i += 1

        # ===== ARRANGE =====
        elif res_type == "ARRANGE":
            _add_instruction(doc, res)
            _render_arrange_from_json(doc, parsed)
        
            i += 1

        # ===== SYNONYM / ANTONYM (GROUP) =====
        elif res_type == "SYNONYM_ANTONYM":
            i = render_synonym_antonym_group(doc, results, i)
           
        
        elif res_type == "SENTENCE_COMPLETION":
            i = render_sentence_completion_group(doc, results, i)
           
        elif res_type == "ERROR_IDENTIFICATION":
            i = render_error_identification_group(doc, results, i)
            
        elif res_type == "SENTENCE_TRANSFORMATION":
            i = render_sentence_transformation_group(doc, results, i)
            
        elif res_type == "PRONUNCIATION_STRESS":
            i = render_pronunciation_stress_group(doc, results, i)
            
        elif res_type == "LOGICAL_THINKING":
            i = render_logical_thinking_group(doc, results, i)
        elif res_type == "WORD_REORDERING":
            i = render_word_reordering_group(doc, results, i)
           
        # ===== UNKNOWN =====
        else:
            logger.warning(f"Unknown type: {res_type}")
            i += 1

    doc.save(output_path)
    return output_path


# ============================
# JSON → DOCX RENDERERS
# ============================


def _render_standard_cloze_from_json(doc: Document, parsed: dict, merge_options: bool = False):
    """
    Render CLOZE / GAP / RC từ JSON đã parse.

    JSON structure:
    {
      "passage_title": "...",
      "passage": "...",
      "questions": [
        {
          "number": 1,
          "question_content": "...",   ← nội dung câu hỏi (có thể rỗng với CLOZE)
          "option_a": "...", "option_b": "...", "option_c": "...", "option_d": "...",
          "answer": "A",
          "explanation": "...",
          "quote": "...",
          "translation": "..."
        }, ...
      ]
    }
    """
    # ── Tiêu đề ──
    title = (parsed.get("passage_title") or "").strip()
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Passage ──
    _render_passage(doc, parsed.get("passage") or "")

    # ── Questions ──
    questions = parsed.get("questions") or []
    for q in questions:
        num              = q.get("number", "?")
        question_content = (q.get("question_content") or "").strip()
        question_content = question_content.replace("'", '"')
        opt_a = q.get("option_a", "")
        opt_b = q.get("option_b", "")
        opt_c = q.get("option_c", "")
        opt_d = q.get("option_d", "")

        # Header: "Question N:" + nội dung câu hỏi (nếu có) trên cùng đoạn
        p = doc.add_paragraph()
        p.add_run(f"Question {num}: ").bold = True
        if question_content:
            question_content = question_content.replace("“", '"').replace("”", '"')
            parts = re.split(r'(".*?")', question_content)

            for part in parts:
                if part.startswith('"') and part.endswith('"'):
                    word = part[1:-1]
                    p.add_run('"')
                    run = p.add_run(word)
                    run.bold = True
                    run.underline = True
                    p.add_run('"')
                else:
                    words = re.split(r'(OPPOSITE)', part)
                    for w in words:
                        run = p.add_run(w)
                        if w == "OPPOSITE":
                            run.bold = True

        if merge_options:
            # Ghép A B C D trên 1 dòng (dạng CLOZE)
            opts_line = f"A. {opt_a}\t\tB. {opt_b}\t\tC. {opt_c}\t\tD. {opt_d}"
            doc.add_paragraph(opts_line)
        else:
            doc.add_paragraph(f"A. {opt_a}")
            doc.add_paragraph(f"B. {opt_b}")
            doc.add_paragraph(f"C. {opt_c}")
            doc.add_paragraph(f"D. {opt_d}")




def add_html_paragraph(doc, text, prefix=None):
    p = doc.add_paragraph()

    # prefix (in đậm)
    if prefix:
        run = p.add_run(prefix)
        run.bold = True

    soup = BeautifulSoup(text, "html.parser")

    for element in soup.descendants:
        if element.name is None and element.string:
            parent = element.parent

            run = p.add_run(element.string)

            run.bold = any(tag.name in ["strong", "b"] for tag in parent.parents) or parent.name in ["strong", "b"]
            run.underline = parent.find_parent("u") is not None or parent.name == "u"
            run.italic = parent.find_parent("i") is not None or parent.name == "i"

def add_multiline_text(doc, text: str):
    p = doc.add_paragraph()
    lines = text.split("\n")
    
    for i, line in enumerate(lines):
        run = p.add_run(line)
        if i < len(lines) - 1:
            run.add_break()  # xuống dòng trong cùng paragraph

def render_sentence_completion_group(doc, results, start_index):
    title_added = False
    i = start_index
    n = len(results)

    while i < n:
        res = results[i]

        if res.get("type") != "SENTENCE_COMPLETION":
            break

        parsed = res.get("parsed")
        if not parsed:
            break

        # ===== Add title (chỉ 1 lần) =====
        if not title_added:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(
                "Sentence completion: Choose A, B, C or D to complete each sentence."
            )
            title_run.bold = True
            title_added = True

        questions = parsed.get("questions", [])

        for q in questions:
            # Question
            p = doc.add_paragraph()
            run_bold = p.add_run(f"Question {q['number']}: ")
            run_bold.bold = True
            run_normal = p.add_run(q['question'])

            # Options (1 dòng)
            p = doc.add_paragraph()
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(1.5))
            tab_stops.add_tab_stop(Inches(2.5))
            tab_stops.add_tab_stop(Inches(2.5))

            p.add_run(
                f"A. {q['option_a']}\t"
                f"B. {q['option_b']}\t"
                f"C. {q['option_c']}\t"
                f"D. {q['option_d']}"
            )
            
            p = doc.add_paragraph()
            p.add_run("Lời giải").bold = True

            # Answer
            doc.add_paragraph(f"Chọn {q['answer']}")

            # Explanation
            add_multiline_text(doc, q["explanation"])

            # Translation
            doc.add_paragraph(f"Tạm dich: {q['translation']}")


        i += 1

    return i


def render_logical_thinking_group(doc, results, start_index):
    i = start_index
    n = len(results)
    title_added = False

    while i < n:
        res = results[i]

        if res.get("type") != "LOGICAL_THINKING":
            break

        parsed = res.get("parsed")
        if not parsed:
            break

        # ===== Title (bold, 1 lần) =====
        if not title_added:
            p = doc.add_paragraph()
            run = p.add_run(
                "Logical thinking and problem solving: Choose A, B, C or D to answer each question."
            )
            run.bold = True
            title_added = True

        questions = parsed.get("questions", [])

        for q in questions:
            # ===== Question number =====
            # doc.add_paragraph(f"Question {q['number']}:")
            p = doc.add_paragraph()  # only 1 paragraph for this question
    
            # Bold part: "Question 1:"
            run_bold = p.add_run(f"Question {q['number']}: ")
            run_bold.bold = True

            # ===== Scenario =====
            if q.get("scenario"):
                doc.add_paragraph(q["scenario"])

            # ===== Dialogue =====
            if q.get("speaker_a"):
                doc.add_paragraph(f"{q['speaker_a']}")
            if q.get("speaker_b"):
                doc.add_paragraph(f"{q['speaker_b']}")

            # ===== Question text =====
            if q.get("question"):
                doc.add_paragraph(q["question"])

            # ===== Options (mỗi dòng riêng) =====
            doc.add_paragraph(f"A. {q['option_a']}")
            doc.add_paragraph(f"B. {q['option_b']}")
            doc.add_paragraph(f"C. {q['option_c']}")
            doc.add_paragraph(f"D. {q['option_d']}")

            p = doc.add_paragraph()
            p.add_run("Lời giải").bold = True
            # ===== Answer =====
            doc.add_paragraph(f"Chọn {q['answer']}")

            # ===== Explanation =====
            add_multiline_text(doc, q["explanation"])

            # ===== Translation =====
            trans = q.get("translation", {})
            if trans:
                if trans.get("scenario"):
                    doc.add_paragraph(f"Tình huống: {trans['scenario']}")
                if trans.get("question"):
                    doc.add_paragraph(f"Câu hỏi: {trans['question']}")
                if trans.get("speaker_a"):
                    doc.add_paragraph(f"{trans['speaker_a']}")
                if trans.get("speaker_b"):
                    doc.add_paragraph(f"{trans['speaker_b']}")


        i += 1

    return i

def render_sentence_transformation_group(doc, results, start_index):
    i = start_index
    n = len(results)
    title_added = False
    current_group_type = None  # rewriting / combination

    while i < n:
        res = results[i]

        if res.get("type") != "SENTENCE_TRANSFORMATION":
            break

        parsed = res.get("parsed")
        if not parsed:
            break

        questions = parsed.get("questions", [])
        if not questions:
            break

        q_type = questions[0].get("type", "").lower()

        # ===== Nếu khác group → break =====
        if current_group_type is None:
            current_group_type = q_type
        elif q_type != current_group_type:
            break

        # ===== Add title (1 lần) =====
        if not title_added:
            p = doc.add_paragraph()
            run = p.add_run()

            if q_type == "rewriting":
                run.text = ("Sentence rewriting: Choose A, B, C or D that has the CLOSEST "
                            "meaning to the given sentence in each question.")
            else:  # combination
                run.text = ("Sentence combination: Choose A, B, C or D that has the CLOSEST "
                            "meaning to the given pair of sentences in each question.")

            run.bold = True
            title_added = True

        # ===== Render câu hỏi =====
        for q in questions:
            # doc.add_paragraph(f"Question {q['number']}: {q['question']}")

            p = doc.add_paragraph()  # only 1 paragraph for this question
    
            # Bold part: "Question 1:"
            run_bold = p.add_run(f"Question {q['number']}: ")
            run_bold.bold = True
            
            # Normal part: the actual question text
            run_normal = p.add_run(q['question'])
            # ===== Options (mỗi đáp án 1 dòng) =====
            doc.add_paragraph(f"A. {q['option_a']}")
            doc.add_paragraph(f"B. {q['option_b']}")
            doc.add_paragraph(f"C. {q['option_c']}")
            doc.add_paragraph(f"D. {q['option_d']}")

            p = doc.add_paragraph()
            p.add_run("Lời giải").bold = True

            # ===== Answer =====
            # doc.add_paragraph(f"Chọn {q['answer']}")

            # ===== Explanation =====
            add_multiline_text(doc, q["explanation"])

            # ===== Translation =====
            doc.add_paragraph(f"Tạm dịch: {q['translation']}")

            # ===== Correct translation (nếu có) =====
            if q.get("correct_translation"):
                doc.add_paragraph(f"→ {q['correct_translation']}")
        i += 1

    return i

def render_word_reordering_group(doc, results, start_index):
    i = start_index
    n = len(results)
    title_added = False

    while i < n:
        res = results[i]
        if res.get("type") != "WORD_REORDERING":
            break

        parsed = res.get("parsed")
        if not parsed:
            break

        # ===== Title (bold + italic, 1 lần) =====
        if not title_added:
            p_title = doc.add_paragraph()
            run = p_title.add_run("Reorder the words given to make correct sentences.")
            run.bold = True
            run.italic = True
            title_added = True

        questions = parsed.get("questions", [])

        for q in questions:
            # ===== Question =====
            p_question = doc.add_paragraph()
            # Phần bold
            run_bold = p_question.add_run(f"Question {q['number']}: ")
            run_bold.bold = True
            # Phần bình thường
            run_normal = p_question.add_run(q.get('word_list', ''))

            # ===== Options ===== (mỗi option 1 paragraph riêng)
            for label in ["option_a", "option_b", "option_c", "option_d"]:
                option_text = q.get(label)
                if option_text:
                    doc.add_paragraph(f"{label[-1].upper()}. {option_text}")

            # ===== Lời giải (bold) =====
            p_ex = doc.add_paragraph()
            run_ex = p_ex.add_run("Lời giải")
            run_ex.bold = True

            # ===== Answer =====
            doc.add_paragraph(f"Chọn {q['answer']}")

            # ===== Explanation =====
            add_multiline_text(doc, q.get("explanation", ""))

            # ===== Translation =====
            translation = q.get("translation")
            if translation:
                doc.add_paragraph(translation)

         

        i += 1

    return i

def render_error_identification_group(doc, results, start_index):
    i = start_index
    n = len(results)
    title_added = False

    while i < n:
        res = results[i]

        if res.get("type") != "ERROR_IDENTIFICATION":
            break

        parsed = res.get("parsed")
        if not parsed:
            break

        # ===== Title (bold + italic, chỉ 1 lần) =====
        if not title_added:
            p = doc.add_paragraph()
            run = p.add_run(
                "Mark the letter A, B, C, or D to indicate the underlined part that needs correction in the following questions."
            )
            run.bold = True
            run.italic = True

            title_added = True

        questions = parsed.get("questions", [])

        for q in questions:
            # ===== Question (có format HTML) =====

            p = doc.add_paragraph()
            run_bold = p.add_run(f"Question {q['number']}: ")
            run_bold.bold = True
            add_html_formatted_text(p, q["question"])

            # ===== Options =====
            p_opt = doc.add_paragraph()
            tab_stops = p_opt.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(1.5))
            tab_stops.add_tab_stop(Inches(2.5))
            tab_stops.add_tab_stop(Inches(2.5))

            p_opt.add_run(
                f"A. {q['option_a']}\t"
                f"B. {q['option_b']}\t"
                f"C. {q['option_c']}\t"
                f"D. {q['option_d']}"
            )

            p = doc.add_paragraph()
            p.add_run("Lời giải").bold = True

            # ===== Answer =====
            doc.add_paragraph(f"Chọn {q['answer']}")

            # ===== Explanation =====
            add_multiline_text(doc, q["explanation"])

            # ===== Correction =====
            doc.add_paragraph(f"Sửa: {q['correction']}")

            # ===== Translation =====
            doc.add_paragraph(f"Tạm dịch: {q['translation']}")

        i += 1

    return i

def render_pronunciation_stress_group(doc, results, start_index):
    i = start_index
    n = len(results)
    title_added = False

    while i < n:
        res = results[i]

        if res.get("type") != "PRONUNCIATION_STRESS":
            break

        parsed = res.get("parsed")
        if not parsed:
            break

        # ===== Title (bold + italic, chỉ 1 lần) =====
        if not title_added:
            p = doc.add_paragraph()
            run = p.add_run(
                "Mark the letter A, B, C, or D on your answer sheet to indicate the word whose underlined part differs from the other three in pronunciation in each of the following questions."
            )
            run.bold = True
            run.italic = True
            title_added = True

        questions = parsed.get("questions", [])

        for q in questions:
            # ===== Question + options cùng dòng =====
            # 
            p = doc.add_paragraph()
            run_bold = p.add_run(f"Question {q['number']}: ")
            run_bold.bold = True

            # A
            run_a = p.add_run("A. ")
            run_a.bold = True
            add_html_formatted_text(p, q['option_a'])
            p.add_run("\t")

            # B
            run_b = p.add_run("B. ")
            run_b.bold = True
            add_html_formatted_text(p, q['option_b'])
            p.add_run("\t")

            # C
            run_c = p.add_run("C. ")
            run_c.bold = True
            add_html_formatted_text(p, q['option_c'])
            p.add_run("\t")

            # D
            run_d = p.add_run("D. ")
            run_d.bold = True
            add_html_formatted_text(p, q['option_d'])

            p = doc.add_paragraph()
            p.add_run("Lời giải").bold = True

            # ===== Answer =====
            doc.add_paragraph(f"Chọn {q['answer']}")

            # ===== Details =====
            for d in q.get("details", []):
                doc.add_paragraph(
                    f"{d['word']} /{d['ipa']}/ ({d['pos']}): {d['meaning']}"
                )

            # ===== Explanation =====
            add_multiline_text(doc, q["explanation"])

        i += 1

    return i


def render_synonym_antonym_group(doc, results, start_index):
    first_res = results[start_index]
    questions = first_res.get("parsed", {}).get("questions", [])

    if not questions:
        return start_index + 1

    group_type = questions[0].get("type", "").lower()

    # ===== Title =====
    title_para = doc.add_paragraph()
    title_run = title_para.add_run()

    if group_type == "synonym":
        title_run.text = (
            "Synonyms: Choose A, B, C or D that has the CLOSEST "
            "meaning to the underlined word/phrase in each question."
        )
    else:
        title_run.text = (
            "Antonyms: Choose A, B, C or D that has the OPPOSITE "
            "meaning to the underlined word/phrase in each question."
        )

    title_run.bold = True

    i = start_index
    n = len(results)

    while i < n:
        res = results[i]

        # ❌ dừng nếu khác loại block
        if res.get("type") != "SYNONYM_ANTONYM":
            break

        parsed = res.get("parsed") or {}
        questions = parsed.get("questions") or []

        if not questions:
            break

        current_type = questions[0].get("type", "").lower()

        # ❌ dừng nếu khác synonym / antonym
        if current_type != group_type:
            break

        # ===== render từng câu =====
        for q in questions:
            # đảm bảo question có HTML vẫn render đúng
            if "question" in q and q["question"]:
                # dùng hàm đã sửa trước đó
                render_single_synonym_question(doc, q)
            else:
                # fallback nếu thiếu data
                doc.add_paragraph(f"Question {q.get('number', '')}: [Missing question]")

        i += 1

    return i

def render_single_synonym_question(doc, q):
    p = doc.add_paragraph()
    
    # tránh crash nếu thiếu number
    question_number = q.get("number", "")
    run_bold = p.add_run(f"Question {question_number}: ")
    run_bold.bold = True

    # render HTML cho question
    add_html_formatted_text(p, q.get("question", ""))

    # options 1 dòng
    p = doc.add_paragraph()
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(1.5))
    tab_stops.add_tab_stop(Inches(2.5))
    tab_stops.add_tab_stop(Inches(2.5))

    p.add_run(
        f"A. {q.get('option_a', '')}\t"
        f"B. {q.get('option_b', '')}\t"
        f"C. {q.get('option_c', '')}\t"
        f"D. {q.get('option_d', '')}"
    )

    p = doc.add_paragraph()
    p.add_run("Lời giải").bold = True

    doc.add_paragraph(f"Chọn {q.get('answer', '')}")

    add_multiline_text(doc, q.get("explanation", ""))

    doc.add_paragraph(f"Tạm dịch: {q.get('translation', '')}")



def render_dialogue_completion_group(doc, results, start_index):
    i = start_index
    n = len(results)
    title_added = False

    while i < n:
        res = results[i]

        if res.get("type") != "DIALOUGE_COMPLETION":
            break

        parsed = res.get("parsed")
        if not parsed:
            break

        # ===== Title (bold, 1 lần) =====
        if not title_added:
            instruction = parsed.get("instruction", "")
            p = doc.add_paragraph()
            run = p.add_run(instruction)
            run.bold = True
            title_added = True

        questions = parsed.get("questions", [])

        for q in questions:
            # ===== Question =====
            # doc.add_paragraph(f"Question {q['number']}:")
            p_question = doc.add_paragraph()
            run_bold = p_question.add_run(f"Question {q['number']}: ")
            run_bold.bold = True

            if q.get("speaker_a"):
                doc.add_paragraph(q["speaker_a"])
            if q.get("speaker_b"):
                doc.add_paragraph(q["speaker_b"])

            # ===== Options dòng 1 (A B) =====
            p1 = doc.add_paragraph()
            tab_stops = p1.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(3))

            p1.add_run(f"A. {q['option_a']}\tB. {q['option_b']}")

            # ===== Options dòng 2 (C D) =====
            p2 = doc.add_paragraph()
            tab_stops = p2.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(3))

            p2.add_run(f"C. {q['option_c']}\tD. {q['option_d']}")

            # ===== Lời giải (bold) =====
            p = doc.add_paragraph()
            p.add_run("Lời giải").bold = True

            # ===== Answer =====
            doc.add_paragraph(f"Chọn {q['answer']}")

            # ===== Explanation =====
            add_multiline_text(doc, q["explanation"])

            # ===== Translation =====
            trans = q.get("translation", {})
            if trans:
                if trans.get("speaker_a"):
                    doc.add_paragraph(trans["speaker_a"])
                if trans.get("speaker_b"):
                    doc.add_paragraph(trans["speaker_b"])


        i += 1

    return i

def _render_cloze_from_json(doc: Document, parsed: dict, merge_options: bool = False, qtype: str = ""):
    """
    Render CLOZE / GAP / RC từ JSON đã parse.
    Question -> Options -> Lời giải -> Explanation -> Quote -> Translation
    """

    # ── Title ─────────────────────────────
    title = (parsed.get("passage_title") or "").strip()
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # ── Passage ───────────────────────────
    _render_passage(doc, parsed.get("passage") or "")

    questions = parsed.get("questions") or []

    # ── Questions + Explanations ──────────
    for q in questions:

        num = q.get("number", "?")
        question_content = (q.get("question_content") or "").strip()
        question_content = question_content.replace("'", '"')
        opt_a = q.get("option_a", "")
        opt_b = q.get("option_b", "")
        opt_c = q.get("option_c", "")
        opt_d = q.get("option_d", "")

        answer = q.get("answer", "?")
        explanation = q.get("explanation", "")
        quote = q.get("quote", "")
        translation = q.get("translation", "")

        # Question header
        p = doc.add_paragraph()
        p.add_run(f"Question {num}: ").bold = True
        if question_content:
            question_content = question_content.replace("“", '"').replace("”", '"').replace("**", '"')
            parts = re.split(r'(".*?")', question_content)

            for part in parts:
                if part.startswith('"') and part.endswith('"'):
                    word = part[1:-1]
                    p.add_run('"')
                    run = p.add_run(word)
                    run.bold = True
                    run.underline = True
                    p.add_run('"')
                else:
                    words = re.split(r'(OPPOSITE)', part)
                    for w in words:
                        run = p.add_run(w)
                        if w == "OPPOSITE":
                            run.bold = True


        # Options
        if merge_options:
            doc.add_paragraph(
                f"A. {opt_a}\t\tB. {opt_b}\t\tC. {opt_c}\t\tD. {opt_d}"
            )
        else:
            doc.add_paragraph(f"A. {opt_a}")
            doc.add_paragraph(f"B. {opt_b}")
            doc.add_paragraph(f"C. {opt_c}")
            doc.add_paragraph(f"D. {opt_d}")


        # ── Lời giải ──
        p = doc.add_paragraph()
        p.add_run("Lời giải").bold = True

        p = doc.add_paragraph()
        p.add_run(f"Chọn {answer}").bold = True

        # Explanation
        if explanation:
            for line in explanation.splitlines():
                line = line.strip()
                if line:
                    p = doc.add_paragraph()
                    add_text_with_markdown_bold(p, line)

        if qtype not in ["RC", "GAP"]:

            if quote:
                # p = doc.add_paragraph()
                # p.add_run("Trích bài: ").bold = True
                # p.add_run(quote)
                add_formatted_paragraph(doc, "Trích bài: ", quote)



            if translation:
                # p = doc.add_paragraph()
                # p.add_run("Tạm dịch: ").bold = True
                # p.add_run(translation)
                add_formatted_paragraph(doc, "Tạm dịch: ", translation)

        # khoảng cách giữa các câu


def _render_standard_arrange_from_json(doc: Document, parsed: dict):
    """
    Render ARRANGE từ JSON đã parse.

    JSON structure:
    {
      "question_number": 1,
      "question_stem": "...",
      "option_a": "...", "option_b": "...", "option_c": "...", "option_d": "...",
      "answer": "A",
      "solution_lines": ["line1", "line2", ...],
      "translation_lines": ["trans1", "trans2", ...]
    }
    """
    num        = parsed.get("question_number", "?")
    stem       = parsed.get("question_stem", "")
    opt_a      = parsed.get("option_a", "")
    opt_b      = parsed.get("option_b", "")
    opt_c      = parsed.get("option_c", "")
    opt_d      = parsed.get("option_d", "")

    # Question stem
    p = doc.add_paragraph()
    p.add_run(f"Question {num}:").bold = True
    if stem:
        doc.add_paragraph(stem)

    # Options (ghép 1 dòng như chuẩn cũ)
    opts_line = f"A. {opt_a}\t\tB. {opt_b}\t\tC. {opt_c}\t\tD. {opt_d}"
    doc.add_paragraph(opts_line)


def _render_arrange_from_json(doc: Document, parsed: dict):
    """
    Render ARRANGE từ JSON đã parse.

    JSON structure:
    {
      "question_number": 1,
      "question_stem": "...",
      "option_a": "...", "option_b": "...", "option_c": "...", "option_d": "...",
      "answer": "A",
      "solution_lines": ["line1", "line2", ...],
      "translation_lines": ["trans1", "trans2", ...]
    }
    """
    num        = parsed.get("question_number", "?")
    stem       = parsed.get("question_stem", "title")
    opt_a      = parsed.get("option_a", "")
    opt_b      = parsed.get("option_b", "")
    opt_c      = parsed.get("option_c", "")
    opt_d      = parsed.get("option_d", "")
    answer     = parsed.get("answer", "?")
    sol_lines  = parsed.get("solution_lines") or []
    trans_lines = parsed.get("translation_lines") or []

    # Question stem
    p = doc.add_paragraph()
    p.add_run(f"Question {num}: ").bold = True
    if stem:
        doc.add_paragraph(stem)

    # Options (ghép 1 dòng như chuẩn cũ)
    opts_line = f"A. {opt_a}\t\tB. {opt_b}\t\tC. {opt_c}\t\tD. {opt_d}"
    doc.add_paragraph(opts_line)

    # Lời giải
    p = doc.add_paragraph()
    p.add_run("Lời giải").bold = True
    doc.add_paragraph(f"Chọn {answer}")

    for line in sol_lines:
        if line.strip():
            doc.add_paragraph(line.strip())

    # Tạm dịch
    if trans_lines:
        p = doc.add_paragraph()
        p.add_run("Tạm dịch:").bold = True
        for line in trans_lines:
            if line.strip():
                # doc.add_paragraph(line.strip())
                add_html_paragraph(doc, line.strip())


# ============================
# FALLBACK: TEXT PARSERS (giữ lại phòng khi JSON parse thất bại)
# ============================

def _render_fallback(doc: Document, block_type: str, raw: str):
    """Fallback về text parser cũ khi JSON parse thất bại."""
    logger.warning(f"Using text fallback renderer for type={block_type}")
    if block_type in ("CLOZE", "GAP", "RC"):
        _render_cloze_or_gap(doc, raw, merge_options=(block_type == "CLOZE"))
    elif block_type == "ARRANGE":
        _render_arrange(doc, raw)
    else:
        doc.add_paragraph(raw or "(No content)")


def _render_cloze_or_gap(doc, raw_text, merge_options=False):
    blocks = extract_cloze_sections(raw_text)
    _render_passage(doc, blocks["passage"])
    q_lines = normalize_question_lines(blocks["questions"])
    if merge_options:
        q_lines = merge_options_single_line(q_lines)
    _render_questions(doc, q_lines)
    _render_answer_key(doc, blocks["answer_key"])
    _render_explanation(doc, blocks["explanation"])


def _render_arrange(doc, raw_text):
    parts = extract_arrange_sections(raw_text)
    q_lines = normalize_question_lines(parts["question_block"])
    q_lines = merge_options_single_line(q_lines)
    _render_questions(doc, q_lines)
    doc.add_paragraph().add_run("Lời giải").bold = True
    if parts["answer_letter"]:
        doc.add_paragraph(f"Chọn {parts['answer_letter']}")
    if parts["solution_text"]:
        for ln in parts["solution_text"].splitlines():
            p = doc.add_paragraph()
            add_text_with_markdown_bold(p, ln.strip())
    if parts["translation"]:
        doc.add_paragraph().add_run("Tạm dịch:").bold = True
        for ln in parts["translation"].splitlines():
            doc.add_paragraph(ln.strip())


# ============================
# STYLE & COMMON DOCX HELPERS
# ============================

def _apply_default_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)


def _add_instruction(doc, res):
    instruction = doc.add_paragraph()
    if res['type'] == "CLOZE":
        text = (
            f"Read the following {res.get('text_type_en').lower()} and mark the letter A, B, C or D "
            "on your answer sheet to indicate the option that best fits each of the numbered blanks."
        )
    elif res['type'] == "ARRANGE":
        text = (
            "Mark the letter A, B, C or D on your answer sheet to indicate the best arrangement "
            "of utterances or sentences to make a meaningful exchange or text."
        )
    elif res['type'] == "RC":
        text = (
            f"Read the following passage and mark the letter A, B, C or D "
            "on your answer sheet to indicate the correct answer to each of the following questions."
        )
    elif res['type'] == "GAP":  # GAP
        text = (
            f"Read the following passage and mark the letter A, B, C or D "
            "on your answer sheet to indicate the option that best fits each of the numbered blank."
        )
    elif res['type'] == "DIALOUGE_COMPLETION":
        text = (
            f"Dialogue completion: Choose A, B, C or D to complete each dialogue."
        )
    elif res['type'] == "SENTENCE_COMPLETION":
        text = (
            f"Sentence completion: Choose A, B, C or D to complete each sentence."
        )
    elif res['type'] == "LOGICAL_THINKING":
        text = (f"Logical thinking and problem-solving: Choose A, B, C or D to answer each question.")
    elif res['type'] == "EERROR_IDENTIFICATION":
        text = (f"Mark the letter A, B, C, or D on your answer sheet to indicate the underlined part that needs correction in the following question.")
    elif res['type'] == "WORD_REORDERING":
        text = (f"Reorder the words given to make a correct sentence.")
    run = instruction.add_run(text)
    run.italic = True


def _render_passage(doc, passage):
    bold_ul_pattern = r"\*\*(.*?)\*\*"
    underline_pattern = r"<u>(.*?)</u>"
    roman_pattern = r"\[(?:I|II|III|IV|V|VI|VII|VIII|IX|X)\]"
    # FIX: thêm \s* để cho phép khoảng trắng giữa (1) và ______
    question_blank_pattern = r"\(\d+\)\s*_{2,}"

    combined = f"{bold_ul_pattern}|{underline_pattern}|{roman_pattern}|{question_blank_pattern}"

    for line in (passage or "").split("\n"):
        line = line.strip()
        if not line:
            continue

        p = doc.add_paragraph()
        pos = 0

        for m in re.finditer(combined, line):
            if m.start() > pos:
                p.add_run(line[pos:m.start()])

            if m.group(1):  # **text**
                run = p.add_run(m.group(1))
                run.bold = True
                run.underline = True

            elif m.group(2):  # <u>text</u>
                run = p.add_run(m.group(2))
                run.bold = True
                run.underline = True

            # FIX: dùng re.search thay vì re.match
            elif re.search(question_blank_pattern, m.group()):
                run = p.add_run(m.group())
                run.bold = True

            else:  # roman [I] [II]
                run = p.add_run(m.group())
                run.bold = True

            pos = m.end()

        if pos < len(line):
            p.add_run(line[pos:])

        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _render_questions(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        if re.match(r"^Question\s+\d+[:\.]", line, re.I):
            r = p.add_run(line)
            r.bold = True
        else:
            p.add_run(line)


def _render_answer_key(doc, key_block):
    if not key_block:
        return
    doc.add_paragraph().add_run("ANSWER KEY").bold = True
    pairs = parse_answer_key_block(key_block)
    if pairs:
        ak_text = ", ".join([f"{i}-{k}" for i, k in pairs])
        doc.add_paragraph(ak_text)
    else:
        doc.add_paragraph(key_block)


def _render_explanation(doc, explanation):
    if not explanation:
        return
    doc.add_paragraph().add_run("HƯỚNG DẪN GIẢI CHI TIẾT").bold = True
    for line in explanation.split('\n'):
        text = line.strip()
        if not text:
            continue
        p = doc.add_paragraph()
        add_text_with_markdown_bold(p, text)


def _add_separator(doc):
    doc.add_paragraph("\n" + "=" * 20 + "\n")


def add_text_with_markdown_bold(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


# ============================
# LEGACY TEXT EXTRACT (FALLBACK)
# ============================

def extract_cloze_sections(raw_text):
    print(f">>>>>> debug raw_text (fallback): {raw_text[:200]}")
    s = html.unescape(raw_text or "").replace("\r\n", "\n").replace("\r", "\n")
    PASS_HDR = "BÀI ĐỌC (PASSAGE)"
    KEY1 = "ANSWER KEY"
    EXPL1 = "HƯỚNG DẪN GIẢI CHI TIẾT"
    EXPL2 = "HƯỚNG DẪN GIẢI"

    pos_pass = s.find(PASS_HDR)
    if pos_pass < 0:
        return {"passage": s[:200], "questions": s[200:], "answer_key": "", "explanation": ""}

    start_pass = pos_pass + len(PASS_HDR)
    pos_first_q = s.find("Question", start_pass)
    if pos_first_q < 0:
        return {"passage": s[start_pass:].strip(), "questions": "", "answer_key": "", "explanation": ""}

    passage = s[start_pass:pos_first_q].strip()
    pos_key = s.find(KEY1, pos_first_q)
    pos_expl = -1
    for h in [EXPL1, EXPL2]:
        p = s.find(h, pos_first_q)
        if p >= 0:
            pos_expl = p
            break

    cut = min([p for p in [pos_key, pos_expl] if p >= 0], default=len(s))
    questions = s[pos_first_q:cut].strip()

    answer_key = ""
    explanation = ""
    if pos_key >= 0:
        key_end = pos_expl if pos_expl > pos_key else len(s)
        answer_key = s[pos_key + len(KEY1):key_end].strip()
    if pos_expl >= 0:
        expl_hdr = EXPL1 if s.find(EXPL1) == pos_expl else EXPL2
        explanation = s[pos_expl + len(expl_hdr):].strip()

    return {"passage": passage, "questions": questions,
            "answer_key": answer_key, "explanation": explanation}


def extract_arrange_sections(raw_text):
    s = html.unescape(raw_text or "").replace("\r\n", "\n").replace("\r", "\n")
    result = {"question_block": "", "answer_letter": "", "solution_text": "", "translation": ""}
    pos_loi = re.search(r"^Lời giải", s, re.MULTILINE)
    if not pos_loi:
        result["question_block"] = s
        return result
    result["question_block"] = s[:pos_loi.start()].strip()
    after = s[pos_loi.end():]
    m_chon = re.search(r"Chọn\s+([ABCD])", after)
    if m_chon:
        result["answer_letter"] = m_chon.group(1)
        after2 = after[m_chon.end():]
    else:
        after2 = after
    pos_tam = after2.find("Tạm dịch:")
    if pos_tam >= 0:
        result["solution_text"] = after2[:pos_tam].strip()
        result["translation"] = after2[pos_tam + len("Tạm dịch:"):].strip()
    else:
        result["solution_text"] = after2.strip()
    return result


def normalize_question_lines(qtext):
    lines = []
    for raw in (qtext or "").splitlines():
        line = html.unescape(raw).strip()
        if not line:
            continue
        line = re.sub(r"\sA\.\s", "\tA. ", line)
        line = re.sub(r"\sB\.\s", "\tB. ", line)
        line = re.sub(r"\sC\.\s", "\tC. ", line)
        line = re.sub(r"\sD\.\s", "\tD. ", line)
        lines.append(line)
    return lines


def parse_answer_key_block(key_block):
    s = (key_block or "").replace("\r\n", "\n")
    pairs = []
    for m in re.finditer(r"(\d+)\s*[-=:]\s*([ABCD])", s, re.I):
        pairs.append((int(m.group(1)), m.group(2).upper()))
    if pairs:
        return sorted(pairs, key=lambda x: x[0])
    for m in re.finditer(r"(\d+)\s+([ABCD])\b", s, re.I):
        pairs.append((int(m.group(1)), m.group(2).upper()))
    return sorted(pairs, key=lambda x: x[0])


def merge_options_single_line(lines):
    merged = []
    buffer = []
    current_question = None

    for line in lines:
        if re.match(r"^Question\s+\d+[:\.]", line, re.I):
            if current_question:
                if buffer:
                    merged.append("\t".join(buffer))
                    buffer = []
            merged.append(line)
            current_question = line
        elif re.match(r"^[ABCD]\.\s", line):
            buffer.append(line.strip())
        else:
            if buffer:
                merged.append("\t".join(buffer))
                buffer = []
            merged.append(line)

    if buffer:
        merged.append("\t".join(buffer))

    return merged


# ============================
# MAIN FLOW
# ============================

async def generate_english_flow(file):
    session_id = str(uuid.uuid4())
    excel_path = UPLOAD_DIR / f"{session_id}_{file.filename}"
    output_path = OUTPUT_DIR / f"ENGLISH_EXAM_{session_id}.docx"

    try:
        with open(excel_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        blocks = extract_blocks_from_excel(str(excel_path))
        results = await generate_exam_docx(blocks, str(output_path))

        # FileResponse(
        #     path=str(output_path),
        #     filename=f"ENGLISH_EXAM_{session_id}.docx",
        #     media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        # )

        return {
            "session_id": session_id,
            "status": "success",
            "message": "Generate English exam successfully",
            "results": results
        }

    except Exception as e:
        logger.exception("Error while generating English exam")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate English exam: {str(e)}"
        )