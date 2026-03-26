
import re
from bs4 import BeautifulSoup,NavigableString
from docx import Document
from docx.shared import Inches

def add_html_formatted_text(paragraph, html_text: str):
    """
    Parse đơn giản các tag: <u>, <i>, <strong>, <b> 
    Hỗ trợ nested tag, nhiều tag cùng lúc.
    """
    # mapping tag -> run attribute
    tag_map = {
        "b": "bold",
        "strong": "bold",
        "i": "italic",
        "u": "underline"
    }

    # regex để tách tag và text
    pattern = r'(<(/?)(strong|b|i|u)>)'
    tags_stack = []
    pos = 0

    for match in re.finditer(pattern, html_text):
        start, end = match.span()
        tag_full, closing, tag_name = match.groups()

        # text giữa các tag
        if start > pos:
            text_segment = html_text[pos:start]
            if text_segment.strip():
                run = paragraph.add_run(text_segment)
                # apply tất cả tag đang active trong stack
                for active_tag in tags_stack:
                    attr = tag_map[active_tag]
                    setattr(run, attr, True)
        pos = end

        # update stack
        if closing:  # </tag>
            if tag_name in tags_stack:
                tags_stack.remove(tag_name)
        else:  # <tag>
            tags_stack.append(tag_name)

    # phần text còn lại sau tag cuối
    if pos < len(html_text):
        text_segment = html_text[pos:]
        if text_segment.strip():
            run = paragraph.add_run(text_segment)
            for active_tag in tags_stack:
                attr = tag_map[active_tag]
                setattr(run, attr, True)


def render_formatted_paragraph(doc, text, prefix=None):
    """
    Render text có format Markdown + HTML vào docx.

    Hỗ trợ:
    - **bold**
    - *italic*
    - __underline__
    - <b>, <strong>, <u>, <i>
    """

    if not text:
        return

    # ── 1. Convert Markdown → HTML ─────────────────
    text = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", text)   # bold
    text = re.sub(r"\*(.*?)\*", r"<i>\1</i>", text)                 # italic
    text = re.sub(r"__(.*?)__", r"<u>\1</u>", text)                 # underline

    # ── 2. Parse HTML ──────────────────────────────
    soup = BeautifulSoup(text, "html.parser")

    p = doc.add_paragraph()

    # prefix (in đậm)
    if prefix:
        run = p.add_run(prefix)
        run.bold = True

    # ── 3. Render từng đoạn text ───────────────────
    for element in soup.descendants:
        if isinstance(element, NavigableString):
            content = str(element)

            if not content.strip():
                continue

            run = p.add_run(content)

            parent = element.parent

            # check toàn bộ ancestor để support nested
            run.bold = any(tag.name in ["strong", "b"] for tag in parent.parents) or parent.name in ["strong", "b"]
            run.italic = any(tag.name == "i" for tag in parent.parents) or parent.name == "i"
            run.underline = any(tag.name == "u" for tag in parent.parents) or parent.name == "u"

def parse_html_like(text):
    # pattern = re.compile(
    #     r'(<strong>)?(<u>)?(<i>)?(.*?)(</i>)?(</u>)?(</strong>)?',
    #     re.DOTALL
    # )
    pattern = re.compile(
        r'(<strong>|<b>)?(<u>)?(<i>)?(.*?)(</i>)?(</u>)?(</strong>|</b>)?',
        re.DOTALL
    )
    results = []

    for match in pattern.finditer(text):
        content = match.group(4)
        if not content:
            continue

        bold = bool(match.group(1) or match.group(7))
        underline = bool(match.group(2) or match.group(6))
        italic = bool(match.group(3) or match.group(5))

        results.append({
            "text": content,
            "bold": bold,
            "underline": underline,
            "italic": italic
        })

    return results

def add_formatted_paragraph(doc, prefix, text):
    p = doc.add_paragraph()
    p.add_run(prefix).bold = True

    soup = BeautifulSoup(text, "html.parser")

    for element in soup.descendants:
        if isinstance(element, NavigableString):
            content = str(element)
            if not content.strip():
                continue

            run = p.add_run(content)

            parent = element.parent

            # check toàn bộ ancestor
            run.bold = any(tag.name in ["strong", "b"] for tag in parent.parents) or parent.name in ["strong", "b"]
            run.underline = any(tag.name == "u" for tag in parent.parents) or parent.name == "u"
            run.italic = any(tag.name == "i" for tag in parent.parents) or parent.name == "i"



def render_standard_sentence_completion_group(doc, results, start_index):
    title_added = False
    i = start_index
    n = len(results)

    while i < n:
        res = results[i]

        if res.get("type") != "SENTENCE_COMPLETION":
            break

        parsed = res.get("parsed")
        if not parsed:
            break

        # ===== Add title (chỉ 1 lần) =====
        if not title_added:
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(
                "Sentence completion: Choose A, B, C or D to complete each sentence."
            )
            title_run.bold = True
            title_added = True

        questions = parsed.get("questions", [])

        for q in questions:
            # Question
            p = doc.add_paragraph()
            run_bold = p.add_run(f"Question {q['number']}: ")
            run_bold.bold = True
            run_normal = p.add_run(q['question'])

            # Options (1 dòng)
            p = doc.add_paragraph()
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(1.5))
            tab_stops.add_tab_stop(Inches(2.5))
            tab_stops.add_tab_stop(Inches(2.5))

            p.add_run(
                f"A. {q['option_a']}\t"
                f"B. {q['option_b']}\t"
                f"C. {q['option_c']}\t"
                f"D. {q['option_d']}"
            )
            

        i += 1

    return i


def render_standard_logical_thinking_group(doc, results, start_index):
    i = start_index
    n = len(results)
    title_added = False

    while i < n:
        res = results[i]

        if res.get("type") != "LOGICAL_THINKING":
            break

        parsed = res.get("parsed")
        if not parsed:
            break

        # ===== Title (bold, 1 lần) =====
        if not title_added:
            p = doc.add_paragraph()
            run = p.add_run(
                "Logical thinking and problem solving: Choose A, B, C or D to answer each question."
            )
            run.bold = True
            title_added = True

        questions = parsed.get("questions", [])

        for q in questions:
            # ===== Question number =====
            # doc.add_paragraph(f"Question {q['number']}:")
            p = doc.add_paragraph()  # only 1 paragraph for this question
    
            # Bold part: "Question 1:"
            run_bold = p.add_run(f"Question {q['number']}: ")
            run_bold.bold = True

            # ===== Scenario =====
            if q.get("scenario"):
                doc.add_paragraph(q["scenario"])

            # ===== Dialogue =====
            if q.get("speaker_a"):
                doc.add_paragraph(f"{q['speaker_a']}")
            if q.get("speaker_b"):
                doc.add_paragraph(f"{q['speaker_b']}")

            # ===== Question text =====
            if q.get("question"):
                doc.add_paragraph(q["question"])

            # ===== Options (mỗi dòng riêng) =====
            doc.add_paragraph(f"A. {q['option_a']}")
            doc.add_paragraph(f"B. {q['option_b']}")
            doc.add_paragraph(f"C. {q['option_c']}")
            doc.add_paragraph(f"D. {q['option_d']}")
        i += 1

    return i

def render_standard_sentence_transformation_group(doc, results, start_index):
    i = start_index
    n = len(results)
    title_added = False
    current_group_type = None  # rewriting / combination

    while i < n:
        res = results[i]

        if res.get("type") != "SENTENCE_TRANSFORMATION":
            break

        parsed = res.get("parsed")
        if not parsed:
            break

        questions = parsed.get("questions", [])
        if not questions:
            break

        q_type = questions[0].get("type", "").lower()

        # ===== Nếu khác group → break =====
        if current_group_type is None:
            current_group_type = q_type
        elif q_type != current_group_type:
            break

        # ===== Add title (1 lần) =====
        if not title_added:
            p = doc.add_paragraph()
            run = p.add_run()

            if q_type == "rewriting":
                run.text = ("Sentence rewriting: Choose A, B, C or D that has the CLOSEST "
                            "meaning to the given sentence in each question.")
            else:  # combination
                run.text = ("Sentence combination: Choose A, B, C or D that has the CLOSEST "
                            "meaning to the given pair of sentences in each question.")

            run.bold = True
            title_added = True

        # ===== Render câu hỏi =====
        for q in questions:
            # doc.add_paragraph(f"Question {q['number']}: {q['question']}")

            p = doc.add_paragraph()  # only 1 paragraph for this question
    
            # Bold part: "Question 1:"
            run_bold = p.add_run(f"Question {q['number']}: ")
            run_bold.bold = True
            
            # Normal part: the actual question text
            run_normal = p.add_run(q['question'])
            # ===== Options (mỗi đáp án 1 dòng) =====
            doc.add_paragraph(f"A. {q['option_a']}")
            doc.add_paragraph(f"B. {q['option_b']}")
            doc.add_paragraph(f"C. {q['option_c']}")
            doc.add_paragraph(f"D. {q['option_d']}")

        i += 1

    return i

def render_standard_word_reordering_group(doc, results, start_index):
    i = start_index
    n = len(results)
    title_added = False

    while i < n:
        res = results[i]
        if res.get("type") != "WORD_REORDERING":
            break

        parsed = res.get("parsed")
        if not parsed:
            break

        # ===== Title (bold + italic, 1 lần) =====
        if not title_added:
            p_title = doc.add_paragraph()
            run = p_title.add_run("Reorder the words given to make correct sentences.")
            run.bold = True
            run.italic = True
            title_added = True

        questions = parsed.get("questions", [])

        for q in questions:
            # ===== Question =====
            p_question = doc.add_paragraph()
            # Phần bold
            run_bold = p_question.add_run(f"Question {q['number']}: ")
            run_bold.bold = True
            # Phần bình thường
            run_normal = p_question.add_run(q.get('word_list', ''))

            # ===== Options ===== (mỗi option 1 paragraph riêng)
            for label in ["option_a", "option_b", "option_c", "option_d"]:
                option_text = q.get(label)
                if option_text:
                    doc.add_paragraph(f"{label[-1].upper()}. {option_text}")


        
        i += 1

    return i

def render_standard_error_identification_group(doc, results, start_index):
    i = start_index
    n = len(results)
    title_added = False

    while i < n:
        res = results[i]

        if res.get("type") != "ERROR_IDENTIFICATION":
            break

        parsed = res.get("parsed")
        if not parsed:
            break

        # ===== Title (bold + italic, chỉ 1 lần) =====
        if not title_added:
            p = doc.add_paragraph()
            run = p.add_run(
                "Mark the letter A, B, C, or D to indicate the underlined part that needs correction in the following questions."
            )
            run.bold = True
            run.italic = True

            title_added = True

        questions = parsed.get("questions", [])

        for q in questions:
            # ===== Question (có format HTML) =====

            p = doc.add_paragraph()
            run_bold = p.add_run(f"Question {q['number']}: ")
            run_bold.bold = True
            add_html_formatted_text(p, q["question"])

            # ===== Options =====
            p_opt = doc.add_paragraph()
            tab_stops = p_opt.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(1.5))
            tab_stops.add_tab_stop(Inches(2.5))
            tab_stops.add_tab_stop(Inches(2.5))

            p_opt.add_run(
                f"A. {q['option_a']}\t"
                f"B. {q['option_b']}\t"
                f"C. {q['option_c']}\t"
                f"D. {q['option_d']}"
            )

        i += 1

    return i

def render_standard_pronunciation_stress_group(doc, results, start_index):
    i = start_index
    n = len(results)
    title_added = False

    while i < n:
        res = results[i]

        if res.get("type") != "PRONUNCIATION_STRESS":
            break

        parsed = res.get("parsed")
        if not parsed:
            break

        # ===== Title (bold + italic, chỉ 1 lần) =====
        if not title_added:
            p = doc.add_paragraph()
            run = p.add_run(
                "Mark the letter A, B, C, or D on your answer sheet to indicate the word whose underlined part differs from the other three in pronunciation in each of the following questions."
            )
            run.bold = True
            run.italic = True
            title_added = True

        questions = parsed.get("questions", [])

        for q in questions:
            # ===== Question + options cùng dòng =====
            # 
            p = doc.add_paragraph()
            run_bold = p.add_run(f"Question {q['number']}: ")
            run_bold.bold = True

            # A
            run_a = p.add_run("A. ")
            run_a.bold = True
            add_html_formatted_text(p, q['option_a'])
            p.add_run("\t")

            # B
            run_b = p.add_run("B. ")
            run_b.bold = True
            add_html_formatted_text(p, q['option_b'])
            p.add_run("\t")

            # C
            run_c = p.add_run("C. ")
            run_c.bold = True
            add_html_formatted_text(p, q['option_c'])
            p.add_run("\t")

            # D
            run_d = p.add_run("D. ")
            run_d.bold = True
            add_html_formatted_text(p, q['option_d'])


        i += 1

    return i


def render_standard_synonym_antonym_group(doc, results, start_index):
    first_res = results[start_index]
    questions = first_res.get("parsed", {}).get("questions", [])

    if not questions:
        return start_index + 1

    group_type = questions[0].get("type", "").lower()

    # ===== Title =====
    title_para = doc.add_paragraph()
    title_run = title_para.add_run()

    if group_type == "synonym":
        title_run.text = (
            "Synonyms: Choose A, B, C or D that has the CLOSEST "
            "meaning to the underlined word/phrase in each question."
        )
    else:
        title_run.text = (
            "Antonyms: Choose A, B, C or D that has the OPPOSITE "
            "meaning to the underlined word/phrase in each question."
        )

    title_run.bold = True

    i = start_index
    n = len(results)

    while i < n:
        res = results[i]

        # ❌ dừng nếu khác loại block
        if res.get("type") != "SYNONYM_ANTONYM":
            break

        parsed = res.get("parsed") or {}
        questions = parsed.get("questions") or []

        if not questions:
            break

        current_type = questions[0].get("type", "").lower()

        # ❌ dừng nếu khác synonym / antonym
        if current_type != group_type:
            break

        # ===== render từng câu =====
        for q in questions:
            # đảm bảo question có HTML vẫn render đúng
            if "question" in q and q["question"]:
                # dùng hàm đã sửa trước đó
                render_standard_single_synonym_question(doc, q)
            else:
                # fallback nếu thiếu data
                doc.add_paragraph(f"Question {q.get('number', '')}: [Missing question]")

        i += 1

    return i

def render_standard_single_synonym_question(doc, q):
    p = doc.add_paragraph()
    
    # tránh crash nếu thiếu number
    question_number = q.get("number", "")
    run_bold = p.add_run(f"Question {question_number}: ")
    run_bold.bold = True

    # render HTML cho question
    add_html_formatted_text(p, q.get("question", ""))

    # options 1 dòng
    p = doc.add_paragraph()
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(1.5))
    tab_stops.add_tab_stop(Inches(2.5))
    tab_stops.add_tab_stop(Inches(2.5))

    p.add_run(
        f"A. {q.get('option_a', '')}\t"
        f"B. {q.get('option_b', '')}\t"
        f"C. {q.get('option_c', '')}\t"
        f"D. {q.get('option_d', '')}"
    )



def render_standard_dialogue_completion_group(doc, results, start_index):
    i = start_index
    n = len(results)
    title_added = False

    while i < n:
        res = results[i]

        if res.get("type") != "DIALOUGE_COMPLETION":
            break

        parsed = res.get("parsed")
        if not parsed:
            break

        # ===== Title (bold, 1 lần) =====
        if not title_added:
            instruction = parsed.get("instruction", "")
            p = doc.add_paragraph()
            run = p.add_run(instruction)
            run.bold = True
            title_added = True

        questions = parsed.get("questions", [])

        for q in questions:
            # ===== Question =====
            # doc.add_paragraph(f"Question {q['number']}:")
            p_question = doc.add_paragraph()
            run_bold = p_question.add_run(f"Question {q['number']}: ")
            run_bold.bold = True

            if q.get("speaker_a"):
                doc.add_paragraph(q["speaker_a"])
            if q.get("speaker_b"):
                doc.add_paragraph(q["speaker_b"])

            # ===== Options dòng 1 (A B) =====
            p1 = doc.add_paragraph()
            tab_stops = p1.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(3))

            p1.add_run(f"A. {q['option_a']}\tB. {q['option_b']}")

            # ===== Options dòng 2 (C D) =====
            p2 = doc.add_paragraph()
            tab_stops = p2.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(3))

            p2.add_run(f"C. {q['option_c']}\tD. {q['option_d']}")


        i += 1

    return i
